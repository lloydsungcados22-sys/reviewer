"""
PH Criminology Exam Reviewer - Streamlit Web App
==================================================
A comprehensive exam reviewer with PDF upload, question generation,
free limits, paywall, premium codes, and admin panel.

Features:
- PDF text extraction and question generation (LLM or rule-based)
- Free limit: 15 questions
- Paywall: ₱50 via GCash
- Premium codes: 100 questions
- Admin panel for code generation and receipt validation
"""

import os
import hashlib
import secrets
import string
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple
import json
import io

import streamlit as st

# Detect if running on Streamlit Cloud
IS_STREAMLIT_CLOUD = os.environ.get("STREAMLIT_SHARING", "").lower() == "true" or \
                     os.environ.get("STREAMLIT_SERVER_ENVIRONMENT", "").lower() == "cloud" or \
                     "streamlit.app" in os.environ.get("_", "")

# Base directory for file storage (works on both local and cloud)
# On Streamlit Cloud, files are stored in the app's working directory
# Note: Files uploaded to Streamlit Cloud are ephemeral and may be lost on redeploy
# For persistent storage, consider using cloud storage (S3, GCS) or database
BASE_DIR = os.getcwd()
UPLOADS_DIR = os.path.join(BASE_DIR, "uploads")
ADMIN_DOCS_DIR = os.path.join(BASE_DIR, "admin_docs")
RECEIPTS_DIR = os.path.join(BASE_DIR, "data", "receipts")

# Ensure directories exist
for directory in [UPLOADS_DIR, ADMIN_DOCS_DIR, RECEIPTS_DIR]:
    try:
        os.makedirs(directory, exist_ok=True)
    except Exception as e:
        print(f"[WARNING] Could not create directory {directory}: {e}")

# Log cloud environment info for debugging (only in development)
if IS_STREAMLIT_CLOUD:
    print(f"🌐 Running on Streamlit Cloud")
    print(f"📁 Working directory: {BASE_DIR}")
    print(f"📁 Uploads: {UPLOADS_DIR} (exists: {os.path.exists(UPLOADS_DIR)})")
    print(f"📁 Admin docs: {ADMIN_DOCS_DIR} (exists: {os.path.exists(ADMIN_DOCS_DIR)})")

# Snowflake connector
try:
    import snowflake.connector
    SNOWFLAKE_AVAILABLE = True
except ImportError:
    SNOWFLAKE_AVAILABLE = False
    st.warning("⚠️ Snowflake connector not installed. Install with: pip install snowflake-connector-python")
import pandas as pd

# PDF processing
fitz_module = None
pdfplumber_module = None
PDF_AVAILABLE = False
PDF_LIB = None
OCR_AVAILABLE = False

try:
    import fitz as fitz_module  # PyMuPDF
    PDF_AVAILABLE = True
    PDF_LIB = "fitz"
    print("[OK] PyMuPDF (fitz) imported successfully")
except ImportError as e:
    fitz_module = None
    print(f"[WARNING] PyMuPDF import failed: {e}")
    try:
        import pdfplumber as pdfplumber_module
        PDF_AVAILABLE = True
        PDF_LIB = "pdfplumber"
        print("[OK] pdfplumber imported successfully")
    except ImportError as e2:
        PDF_AVAILABLE = False
        PDF_LIB = None
        pdfplumber_module = None
        print(f"[WARNING] pdfplumber import failed: {e2}")
        print("[ERROR] No PDF library available - PDF processing will not work")

# Optional OCR support for scanned/image-based PDFs and documents
EASYOCR_AVAILABLE = False
TESSERACT_AVAILABLE = False
easyocr_reader = None

try:
    import easyocr
    EASYOCR_AVAILABLE = True
    print("[OK] EasyOCR imported successfully")
    # Initialize EasyOCR reader lazily (on first use to avoid slow startup)
    # We'll initialize it when needed in the extraction function
except Exception as e:
    EASYOCR_AVAILABLE = False
    print(f"[WARNING] EasyOCR import failed: {e}")
    import traceback
    print(traceback.format_exc())

def get_easyocr_reader(progress_callback=None, timeout=300):
    """Get or initialize EasyOCR reader (lazy initialization with timeout for cloud)"""
    global easyocr_reader
    
    # Check if EasyOCR is available (read-only check, no global needed)
    if not EASYOCR_AVAILABLE:
        if progress_callback:
            progress_callback("⚠ EasyOCR library not available")
        return None
    
    # Check if we've disabled EasyOCR due to too many failures
    if hasattr(get_easyocr_reader, 'disabled') and get_easyocr_reader.disabled:
        if progress_callback:
            progress_callback("⚠ EasyOCR disabled due to repeated initialization failures")
        return None
    
    if easyocr_reader is None:
        try:
            import easyocr
            
            # Initialize EasyOCR reader (English only for now, can add more languages)
            # gpu=False for CPU, set gpu=True if CUDA is available
            # Note: First initialization downloads models (~500MB) and may take 2-5 minutes
            if progress_callback:
                progress_callback("📥 Initializing EasyOCR (downloading models on first use - this may take 2-5 minutes)...")
            
            # For Streamlit Cloud, use a simpler initialization that's more resilient
            # Try to initialize with timeout handling
            # For Streamlit Cloud, be more careful with initialization
            if IS_STREAMLIT_CLOUD:
                # On cloud, try to use cached models first to avoid timeout
                try:
                    easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False, download_enabled=False)
                    if progress_callback:
                        progress_callback("✓ EasyOCR reader initialized (using cached models)")
                    print("[OK] EasyOCR reader initialized with cached models (cloud mode)")
                except Exception as cache_error:
                    # If cached models don't work, try downloading (but this might timeout)
                    print(f"[WARNING] Cached models not available, attempting download (may timeout on cloud)...")
                    if progress_callback:
                        progress_callback("📥 Downloading EasyOCR models (this may take 2-5 minutes, may timeout on cloud)...")
                    try:
                        easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False, download_enabled=True)
                        if progress_callback:
                            progress_callback("✓ EasyOCR reader initialized successfully")
                        print("[OK] EasyOCR reader initialized successfully (cloud mode)")
                    except Exception as download_error:
                        print(f"[ERROR] EasyOCR download failed on cloud: {download_error}")
                        raise download_error
            else:
                # Local environment - try download first
                try:
                    easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False, download_enabled=True)
                    if progress_callback:
                        progress_callback("✓ EasyOCR reader initialized successfully")
                    print("[OK] EasyOCR reader initialized successfully")
                except Exception as init_error:
                    # If initialization fails, try without download (models might already be cached)
                    print(f"[WARNING] First EasyOCR init failed: {init_error}, trying cached models...")
                    try:
                        easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False, download_enabled=False)
                        if progress_callback:
                            progress_callback("✓ EasyOCR reader initialized (using cached models)")
                        print("[OK] EasyOCR reader initialized with cached models")
                    except Exception as cache_error:
                        raise init_error  # Raise original error
                    
        except Exception as e:
            import traceback
            error_msg = f"EasyOCR initialization error: {e}"
            print(f"[ERROR] {error_msg}")
            print(traceback.format_exc())
            if progress_callback:
                progress_callback(f"❌ {error_msg}")
            # Don't set to None permanently - allow retry on next call
            # But limit retries to avoid infinite loops
            if not hasattr(get_easyocr_reader, 'retry_count'):
                get_easyocr_reader.retry_count = 0
            get_easyocr_reader.retry_count += 1
            if get_easyocr_reader.retry_count > 3:
                print("[ERROR] EasyOCR initialization failed after 3 attempts, disabling OCR for this session")
                get_easyocr_reader.disabled = True
            return None
    return easyocr_reader

try:
    from PIL import Image
    import pytesseract
    TESSERACT_AVAILABLE = True
    print("[OK] Tesseract OCR (pytesseract) imported successfully")
except Exception as e:
    TESSERACT_AVAILABLE = False
    print(f"[WARNING] Tesseract OCR import failed: {e}")

OCR_AVAILABLE = EASYOCR_AVAILABLE or TESSERACT_AVAILABLE
if OCR_AVAILABLE:
    print(f"[OK] OCR available: EasyOCR={EASYOCR_AVAILABLE}, Tesseract={TESSERACT_AVAILABLE}")
else:
    print("[ERROR] No OCR library available - scanned document processing will not work")

# Word document processing
Document_class = None
DOCX_AVAILABLE = False
try:
    from docx import Document as Document_class
    DOCX_AVAILABLE = True
    print("[OK] python-docx imported successfully")
except ImportError as e:
    DOCX_AVAILABLE = False
    Document_class = None
    print(f"[WARNING] python-docx import failed: {e}")
    print("[ERROR] Word document processing will not work")


def extract_text(uploaded_file) -> str:
    """
    Lightweight text extractor for Streamlit uploaded files.
    - Reads bytes once to avoid "empty extraction" issues.
    - Uses python-docx for DOCX.
    - Uses pypdf first, then pdfplumber for PDFs.
    - Returns empty string on failure.
    """
    if not uploaded_file:
        return ""

    # Get raw bytes from the uploader
    try:
        file_bytes = uploaded_file.getvalue()
    except Exception:
        try:
            uploaded_file.seek(0)
            file_bytes = uploaded_file.read()
        except Exception:
            return ""

    filename = getattr(uploaded_file, "name", "") or ""
    ext = os.path.splitext(filename.lower())[1]

    # ---- DOCX ----
    if ext == ".docx":
        try:
            from docx import Document
        except Exception:
            return ""

        try:
            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs if p.text and p.text.strip()])
            return text.strip()
        except Exception:
            return ""

    # ---- DOC (old Word format) ----
    if ext == ".doc":
        # python-docx cannot reliably read legacy .doc files
        return ""

    # ---- Images (OCR) ----
    if ext in ('.png', '.jpg', '.jpeg'):
        try:
            uploaded_file.seek(0)
            text, _ = extract_text_from_image(uploaded_file)
            return text or ""
        except Exception:
            return ""

    # ---- PDF ----
    if ext == ".pdf":
        # Try pypdf first (fast, pure Python)
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            page_texts = []
            for page in reader.pages:
                try:
                    page_texts.append((page.extract_text() or ""))
                except Exception:
                    page_texts.append("")
            text = "\n".join(page_texts).strip()
            if len(text) >= 50:
                return text
        except Exception:
            # Fall through to pdfplumber
            pass

        # Fallback: pdfplumber (already used elsewhere in the app)
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_texts = []
                for page in pdf.pages:
                    try:
                        page_texts.append((page.extract_text() or ""))
                    except Exception:
                        page_texts.append("")
                text = "\n".join(page_texts).strip()
            return text
        except Exception:
            return ""

    # Unsupported extension
    return ""

# LLM for question generation (optional)
try:
    import openai
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False

# ReportLab for PDF export (optional)
try:
    import reportlab
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# ============================================================================
# CONFIGURATION
# ============================================================================

# Page config
st.set_page_config(
    page_title="PH Criminology Exam Reviewer",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Constants - Revised Tiering System
FREE_QUESTION_LIMIT = 15
ADVANCE_QUESTION_LIMIT = 90  # 15 free + 75 additional = 90 total
ADVANCE_PAYMENT_AMOUNT = 50  # ₱50 for Advance (75 additional questions)
PREMIUM_PAYMENT_AMOUNT = 299  # ₱299 for Premium (unlimited questions for 1 month)
PAYMENT_AMOUNT = ADVANCE_PAYMENT_AMOUNT  # Default for backward compatibility
GCASH_NUMBER = "0927 159 5709"
GCASH_NAME = "M**K L***D S."
RECEIPT_EMAIL = "criminologysupp@gmail.com"

# ============================================================================
# DATABASE SETUP - SNOWFLAKE
# ============================================================================

def get_snowflake_config():
    """Get Snowflake configuration from secrets.toml - password authentication only"""
    try:
        snowflake_config = st.secrets["SNOWFLAKE"]
        # Get password - required for password authentication
        password = snowflake_config.get("password", "")
        
        # Only use authenticator if explicitly set (but prefer password auth)
        authenticator = snowflake_config.get("authenticator", "").strip()
        # If authenticator is explicitly set to something other than password, use it
        # Otherwise, use password authentication (don't set authenticator parameter)
        
        return {
            "account": snowflake_config.get("account", ""),
            "user": snowflake_config.get("user", ""),
            "password": password,
            "authenticator": authenticator if authenticator and authenticator.lower() != "password" else "",  # Empty means password auth
            "role": snowflake_config.get("role", ""),
            "warehouse": snowflake_config.get("warehouse", ""),
            "database": snowflake_config.get("database", ""),
            "schema": snowflake_config.get("schema", "")
        }
    except Exception as e:
        st.error(f"❌ Error loading Snowflake configuration: {str(e)}")
        return None

@st.cache_resource
def get_snowflake_connection():
    """Get cached Snowflake database connection"""
    if not SNOWFLAKE_AVAILABLE:
        st.error("❌ snowflake-connector-python is not installed. Please install it: pip install snowflake-connector-python")
        st.stop()
        return None
    
    config = get_snowflake_config()
    if not config:
        st.error("❌ Snowflake configuration not found in secrets.toml. Please configure [SNOWFLAKE] section.")
        st.stop()
        return None
    
    try:
        # Extract account locator from full account identifier
        account = config["account"].strip()
        # Remove any domain suffix if present (e.g., .snowflakecomputing.com)
        if '.' in account:
            account = account.split('.')[0]
        
        # For SSO/SAML accounts, ensure account format is correct
        # Account should be in format: orgname-accountname (dash-separated, not dot-separated)
        # The account "TOJPYBV-JF89768" looks correct already
        
        # Build connection parameters
        conn_params = {
            "account": account,
            "user": config["user"].strip(),
        }
        
        # Require password for authentication (password-only, no external browser)
        password = config.get("password", "").strip()
        if not password:
            st.error("❌ **Password is required** for Snowflake connection. Please set 'password' in secrets.toml [SNOWFLAKE] section.")
            st.error("⚠️ External browser authentication is not supported. Password authentication is required.")
            st.stop()
            return None
        
        conn_params["password"] = password
        # Only include authenticator if explicitly set to something other than password
        # If authenticator is empty or "password", omit it (defaults to password auth)
        authenticator = config.get("authenticator", "").strip()
        if authenticator and authenticator.lower() not in ["", "password"]:
            conn_params["authenticator"] = authenticator
        # Otherwise, don't set authenticator - Snowflake will use password authentication by default
        
        # Add optional parameters only if they are provided and not empty
        if config.get("role") and config["role"].strip() and config["role"] != "<none selected>":
            conn_params["role"] = config["role"].strip()
        if config.get("warehouse") and config["warehouse"].strip() and config["warehouse"] != "<none selected>":
            conn_params["warehouse"] = config["warehouse"].strip()
        if config.get("database") and config["database"].strip():
            conn_params["database"] = config["database"].strip()
        if config.get("schema") and config["schema"].strip():
            conn_params["schema"] = config["schema"].strip()
        
        # Clear any proxy-related environment variables that might interfere
        import os
        proxy_vars = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 'ALL_PROXY', 'all_proxy']
        for var in proxy_vars:
            if var in os.environ:
                del os.environ[var]
        
        conn = snowflake.connector.connect(**conn_params)
        return conn
    except Exception as e:
        error_msg = str(e)
        st.error(f"❌ Error connecting to Snowflake: {error_msg}")
        st.error("⚠️ **Critical:** Snowflake connection is required. SQLite fallback is disabled.")
        
        # Provide specific troubleshooting based on error
        if "SAML" in error_msg or "Identity Provider" in error_msg:
            st.error("""
            **SAML/Identity Provider Error Detected:**
            
            This error typically occurs when:
            1. The account format is incorrect for SSO/SAML authentication
            2. SSO/SAML authentication is required but not properly configured
            3. The authenticator type doesn't match your account setup
            
            **Solutions:**
            1. **Verify Account Format**: The account in secrets.toml should be in format `orgname-accountname` (dash-separated).
               Current value: `{account}`
            
            2. **Password Authentication Required**:
               - This application requires password authentication (external browser is not supported)
               - Ensure `password = "your_password"` is set in secrets.toml [SNOWFLAKE] section
               - Verify your password is correct and has not expired
            
            3. **If Your Account Uses SSO/SAML**:
               - Contact your Snowflake administrator to set up password authentication
               - Ask your admin to enable password-based login for your account
               - Some accounts may require a service account with password authentication
            
            4. **Account Identifier Format**:
               - Verify with your Snowflake admin the correct account identifier for programmatic access
               - Ensure the account format matches your organization's requirements
            """.format(account=config.get("account", "N/A")))
        elif "Failed to connect" in error_msg:
            st.error("""
            **Connection Failed:**
            
            Possible causes:
            1. Network/firewall blocking Snowflake access
            2. Incorrect account identifier
            3. Snowflake service unavailable
            
            **Solutions:**
            - Verify account identifier format in secrets.toml
            - Check network connectivity to Snowflake
            - Try connecting from Snowflake web interface to verify credentials
            """)
        
        st.stop()
        return None

def init_snowflake_tables():
    """Initialize Snowflake tables if they don't exist"""
    # Use the global connection instead of creating a new one
    if not snowflake_conn:
        return False
    conn = snowflake_conn
    
    try:
        cursor = conn.cursor()
        
        # Use the configured schema
        config = get_snowflake_config()
        database = config["database"] if config else "REVIEWER"
        schema = config["schema"] if config else "PUBLIC"
        
        # Try to use the database and schema
        try:
            cursor.execute(f"USE DATABASE {database}")
        except Exception as e:
            st.warning(f"⚠️ Could not switch to database {database}: {str(e)}")
        
        # NOTE: We intentionally do NOT try to CREATE SCHEMA here.
        # Many managed Snowflake roles (like learning/demo roles) do not have
        # permission to create schemas, only to use existing ones.
        # We assume the target schema already exists and simply try to use it.

        # Try to use the schema directly
        try:
            cursor.execute(f"USE SCHEMA {database}.{schema}")
        except Exception as use_error:
            st.error(f"❌ Cannot use schema {schema} in database {database}")
            st.error(f"Error: {str(use_error)}")
            st.error("""
            **Schema Access Error:**
            
            Your role `SNOWFLAKE_LEARNING_ROLE` does not have permissions to use the schema `{schema}`.
            
            **Solutions:**
            1. **Ask your Snowflake administrator to grant permissions:**
               ```sql
               GRANT USAGE ON SCHEMA {database}.{schema} TO ROLE SNOWFLAKE_LEARNING_ROLE;
               GRANT CREATE TABLE ON SCHEMA {database}.{schema} TO ROLE SNOWFLAKE_LEARNING_ROLE;
               ```
            
            2. **Or use a different schema** that you have access to:
               - Update `schema = "YOUR_SCHEMA_NAME"` in `.streamlit/secrets.toml`
               - Common options: `PUBLIC`, or a custom schema name
            
            3. **Or create a new schema** (if you have permission):
               ```sql
               CREATE SCHEMA {database}.REVIEWER_SCHEMA;
               GRANT USAGE ON SCHEMA {database}.REVIEWER_SCHEMA TO ROLE SNOWFLAKE_LEARNING_ROLE;
               GRANT CREATE TABLE ON SCHEMA {database}.REVIEWER_SCHEMA TO ROLE SNOWFLAKE_LEARNING_ROLE;
               ```
               Then update `secrets.toml` with `schema = "REVIEWER_SCHEMA"`
            """.format(schema=schema, database=database))
            cursor.close()
            return False
        
        # After USE SCHEMA, we can reference tables without schema prefix
        # Premium codes table
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS premium_codes (
                code VARCHAR(255) PRIMARY KEY,
                status VARCHAR(50) DEFAULT 'active',
                expiry_date VARCHAR(50),
                max_uses INTEGER DEFAULT 1,
                uses_left INTEGER DEFAULT 1,
                created_at VARCHAR(50),
                created_by VARCHAR(255),
                access_level VARCHAR(50) DEFAULT 'Premium'
            )
        """)
        
        # Code usage tracking
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS code_usage (
                id INTEGER AUTOINCREMENT PRIMARY KEY,
                code VARCHAR(255),
                used_at VARCHAR(50),
                session_id VARCHAR(255),
                FOREIGN KEY (code) REFERENCES premium_codes(code)
            )
        """)
        
        # Payment receipts table
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS payment_receipts (
                id INTEGER AUTOINCREMENT PRIMARY KEY,
                full_name VARCHAR(255),
                email VARCHAR(255),
                gcash_reference VARCHAR(255),
                receipt_filename VARCHAR(500),
                status VARCHAR(50) DEFAULT 'Pending',
                submitted_at VARCHAR(50),
                reviewed_at VARCHAR(50),
                reviewed_by VARCHAR(255),
                notes VARCHAR(2000)
            )
        """)
        
        # Users table
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS users (
                email VARCHAR(255) PRIMARY KEY,
                access_level VARCHAR(50) DEFAULT 'Free',
                questions_answered INTEGER DEFAULT 0,
                created_at VARCHAR(50),
                last_login VARCHAR(50),
                premium_code_used VARCHAR(255),
                is_admin INTEGER DEFAULT 0,
                last_ip VARCHAR(45)
            )
        """)
        
        # Migration: add last_ip to users if table existed without it
        try:
            user_cols = get_table_columns("users")
            if user_cols and "last_ip" not in [c.lower() for c in user_cols]:
                cursor.execute(f"ALTER TABLE {get_table_name('users')} ADD COLUMN last_ip VARCHAR(45)")
                if is_snowflake():
                    db_conn.commit()
        except Exception as _:
            pass  # Column may already exist or DB doesn't support ALTER
        
        # PDF management table
        cursor.execute(f"""
            CREATE TABLE IF NOT EXISTS pdf_resources (
                id INTEGER AUTOINCREMENT PRIMARY KEY,
                filename VARCHAR(500),
                filepath VARCHAR(1000),
                is_premium_only INTEGER DEFAULT 0,
                use_for_ai_generation INTEGER DEFAULT 1,
                uploaded_at VARCHAR(50),
                uploaded_by VARCHAR(255),
                description VARCHAR(2000),
                is_downloadable INTEGER DEFAULT 1
            )
        """)
        
        cursor.close()
        return True
    except Exception as e:
        st.error(f"❌ Error initializing Snowflake tables: {str(e)}")
        return False

# Helper functions for database operations
def is_snowflake():
    """Check if using Snowflake database"""
    return snowflake_conn is not None

def get_table_name(table: str) -> str:
    """Get table name with schema prefix if using Snowflake"""
    if is_snowflake():
        config = get_snowflake_config()
        schema = config["schema"] if config else "PUBLIC"
        return f"{schema}.{table}"
    return table

def execute_query(query: str, params: tuple = None):
    """Execute query with proper handling for Snowflake"""
    if not db_conn:
        st.error("❌ Database connection not available. Snowflake connection required.")
        st.stop()
        return None
    
    cursor = db_conn.cursor()
    try:
        if params:
            cursor.execute(query, params)
        else:
            cursor.execute(query)
        # Snowflake doesn't need explicit commit for DDL, but we'll commit for DML
        if query.strip().upper().startswith(('INSERT', 'UPDATE', 'DELETE')):
            db_conn.commit()
        return cursor
    except Exception as e:
        # Log error but don't rollback (Snowflake handles transactions differently)
        raise e

def get_table_columns(table: str) -> List[str]:
    """Get column names for a table"""
    if is_snowflake():
        config = get_snowflake_config()
        schema = config["schema"] if config else "PUBLIC"
        cursor = db_conn.cursor()
        cursor.execute(f"SHOW COLUMNS IN TABLE {schema}.{table}")
        columns = [row[2] for row in cursor.fetchall()]  # Column name is in position 2
        cursor.close()
        return columns
    else:
        # Should never reach here since SQLite fallback is removed
        st.error("❌ Database connection error: SQLite fallback is disabled. Snowflake connection required.")
        return []

# Initialize Snowflake connection and tables - REQUIRED, NO FALLBACK
snowflake_conn = get_snowflake_connection()
if not snowflake_conn:
    st.error("❌ **CRITICAL ERROR:** Cannot connect to Snowflake database. The application requires Snowflake to function.")
    st.error("Please fix the Snowflake connection configuration in secrets.toml and restart the app.")
    st.stop()

# Initialize tables
if not init_snowflake_tables():
    st.error("❌ **CRITICAL ERROR:** Failed to initialize Snowflake tables. Please check your database permissions.")
    st.stop()

# Set database connection - Snowflake only, no fallback
db_conn = snowflake_conn

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def init_session_state():
    """Initialize session state variables (lazy initialization)"""
    # Only initialize if not already set to avoid unnecessary work
    defaults = {
        "user_email": None,
        "user_logged_in": False,
        "user_access_level": "Free",
        "questions_answered": 0,
        "premium_active": False,
        "premium_code_used": None,
        "current_questions": [],
        "current_question_index": 0,
        "answers": {},
        "score": 0,
        "pdf_text": "",
        "pdf_name": "",
        "admin_logged_in": False,
        "uploaded_pdfs": [],
        "selected_pdf": None,
        "selected_documents": [],  # Document library selections
        "openai_api_key": None,
        "openai_model": "gpt-4",
        "openai_temperature": 0.3,
        "generation_progress": "",
        "exam_paused": False,
        "paused_at_index": 0,
        "_css_injected": False  # CSS injection flag
    }
    
    # Only set missing keys (more efficient than checking each)
    missing_keys = set(defaults.keys()) - set(st.session_state.keys())
    for key in missing_keys:
        st.session_state[key] = defaults[key]

# Initialize session state (runs on every rerun but optimized)
init_session_state()

# ============================================================================
# PDF PROCESSING
# ============================================================================

def extract_text_from_docx(docx_file) -> Tuple[str, str]:
    """Extract text from Word document - comprehensive extraction including tables, headers, footers, with OCR fallback"""
    filename = "unknown.docx"
    try:
        # Get filename first
        if hasattr(docx_file, 'name'):
            filename = docx_file.name
        elif isinstance(docx_file, str):
            filename = os.path.basename(docx_file)
        
        if not DOCX_AVAILABLE or Document_class is None:
            print(f"[WARNING] Word library not available for {filename}")
            return "", filename
        
        # Open document with better error handling
        doc = None
        file_path_for_ocr = None
        try:
            if isinstance(docx_file, str):
                # File path
                if not os.path.exists(docx_file):
                    print(f"[WARNING] File not found: {docx_file}")
                    return "", filename
                file_path_for_ocr = docx_file
                try:
                    doc = Document_class(docx_file)
                except Exception as open_error:
                    print(f"[WARNING] Error opening Word document {filename} with python-docx: {open_error}")
                    # Try OCR as fallback if document can't be opened normally
                    if OCR_AVAILABLE:
                        print(f"🔄 Attempting OCR extraction for {filename}...")
                        return extract_text_from_docx_with_ocr(docx_file, filename)
                    return "", filename
            else:
                # File object - save temporarily for OCR if needed
                docx_file.seek(0)
                file_bytes = docx_file.read()
                if not file_bytes:
                    print(f"[WARNING] Empty file: {filename}")
                    return "", filename
                try:
                    doc = Document_class(io.BytesIO(file_bytes))
                except Exception as open_error:
                    print(f"[WARNING] Error opening Word document {filename} from file object: {open_error}")
                    # Save to temp file for OCR
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                        tmp_file.write(file_bytes)
                        file_path_for_ocr = tmp_file.name
                    if OCR_AVAILABLE:
                        print(f"🔄 Attempting OCR extraction for {filename}...")
                        result = extract_text_from_docx_with_ocr(file_path_for_ocr, filename)
                        try:
                            os.unlink(file_path_for_ocr)
                        except:
                            pass
                        return result
                    return "", filename
        except Exception as e:
            print(f"[ERROR] Unexpected error opening Word document {filename}: {str(e)}")
            import traceback
            traceback.print_exc()
            # Try OCR as last resort
            if OCR_AVAILABLE and file_path_for_ocr and os.path.exists(file_path_for_ocr):
                try:
                    return extract_text_from_docx_with_ocr(file_path_for_ocr, filename)
                except:
                    pass
            return "", filename
        
        if doc is None:
            return "", filename
        
        text_parts = []
        
        # 1. Extract text from all paragraphs
        try:
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip() if paragraph.text else ""
                if para_text:
                    text_parts.append(para_text)
        except Exception as e:
            print(f"[WARNING] Error extracting paragraphs: {e}")
        
        # 2. Extract text from tables
        try:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip() if cell.text else ""
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
        except Exception as e:
            print(f"[WARNING] Error extracting tables: {e}")
        
        # 3. Extract text from headers
        try:
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        para_text = paragraph.text.strip() if paragraph.text else ""
                        if para_text:
                            text_parts.append(para_text)
        except Exception as e:
            print(f"[WARNING] Error extracting headers: {e}")
        
        # 4. Extract text from footers
        try:
            for section in doc.sections:
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        para_text = paragraph.text.strip() if paragraph.text else ""
                        if para_text:
                            text_parts.append(para_text)
        except Exception as e:
            print(f"[WARNING] Error extracting footers: {e}")
        
        # Combine all text
        text = "\n".join(text_parts)
        
        # Log extraction results
        if text.strip():
            print(f"[OK] Extracted {len(text)} characters from {filename} (paragraphs: {len([p for p in text_parts if p])})")
        else:
            print(f"[WARNING] No text extracted from {filename} - document may be empty, corrupted, or image-based")
            # Check if document has any structure
            try:
                if doc:
                    para_count = len(doc.paragraphs)
                    table_count = len(doc.tables)
                    print(f"   Document structure: {para_count} paragraphs, {table_count} tables")
            except:
                pass
        
        # If no text extracted, try OCR as fallback
        if not text.strip() and OCR_AVAILABLE:
            print(f"[WARNING] No text extracted from {filename}, attempting OCR...")
            ocr_file_path = None
            try:
                if isinstance(docx_file, str):
                    ocr_file_path = docx_file
                else:
                    # Save file object to temp file for OCR
                    import tempfile
                    docx_file.seek(0)
                    file_bytes = docx_file.read()
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                        tmp_file.write(file_bytes)
                        ocr_file_path = tmp_file.name
                
                if ocr_file_path and os.path.exists(ocr_file_path):
                    ocr_text, _ = extract_text_from_docx_with_ocr(ocr_file_path, filename)
                    # Clean up temp file if we created it
                    if not isinstance(docx_file, str) and ocr_file_path != docx_file:
                        try:
                            os.unlink(ocr_file_path)
                        except:
                            pass
                    if ocr_text and ocr_text.strip():
                        print(f"[OK] OCR extracted {len(ocr_text)} chars from {filename}")
                        return ocr_text, filename
            except Exception as e:
                print(f"[ERROR] OCR fallback failed for {filename}: {e}")
                # Clean up temp file if we created it
                if not isinstance(docx_file, str) and ocr_file_path and ocr_file_path != docx_file:
                    try:
                        os.unlink(ocr_file_path)
                    except:
                        pass
        
        return text, filename
    except ImportError:
        print(f"[ERROR] python-docx not available for {filename}")
        return "", filename
    except Exception as e:
        # Log error with full traceback
        import traceback
        error_msg = f"Error extracting Word document {filename}: {str(e)}"
        print(f"[ERROR] {error_msg}")
        print(traceback.format_exc())
        # Try OCR as last resort if file path is available
        if OCR_AVAILABLE and isinstance(docx_file, str) and os.path.exists(docx_file):
            try:
                return extract_text_from_docx_with_ocr(docx_file, filename)
            except:
                pass
        return "", filename

def extract_text_from_docx_with_ocr(docx_file_path: str, filename: str) -> Tuple[str, str]:
    """Extract text from Word document using OCR (extract images and OCR them, or convert to PDF)"""
    try:
        if not OCR_AVAILABLE:
            print(f"[WARNING] OCR not available for {filename}")
            return "", filename
        
        ocr_texts = []
        
        # Method 1: Extract images from Word document and OCR them
        if EASYOCR_AVAILABLE:
            reader = get_easyocr_reader()
            if reader:
                try:
                    import zipfile
                    # Word documents are ZIP files
                    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
                        # Find image files
                        image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                        print(f"📷 Found {len(image_files)} images in {filename}")
                        for img_file in image_files[:20]:  # Limit to first 20 images
                            try:
                                img_data = docx_zip.read(img_file)
                                img = Image.open(io.BytesIO(img_data))
                                import numpy as np
                                img_array = np.array(img)
                                results = reader.readtext(img_array, paragraph=True, detail=0)
                                if results:
                                    page_text = "\n".join(results)
                                    if page_text.strip():
                                        ocr_texts.append(page_text)
                                        print(f"[OK] OCR extracted text from image {img_file}")
                            except Exception as img_error:
                                print(f"[WARNING] Error processing image {img_file}: {img_error}")
                                continue
                except Exception as zip_error:
                    print(f"[WARNING] Error reading Word document as ZIP: {zip_error}")
        
        # Method 3: Try Tesseract OCR on images if EasyOCR didn't work
        if not ocr_texts and TESSERACT_AVAILABLE:
            try:
                import zipfile
                with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                    for img_file in image_files[:10]:
                        try:
                            img_data = docx_zip.read(img_file)
                            img = Image.open(io.BytesIO(img_data))
                            ocr_text = pytesseract.image_to_string(img)
                            if ocr_text and ocr_text.strip():
                                ocr_texts.append(ocr_text)
                        except Exception:
                            continue
            except Exception:
                pass
        
        if ocr_texts:
            combined_text = "\n".join(ocr_texts)
            print(f"[OK] OCR extracted {len(combined_text)} total characters from {filename}")
            return combined_text, filename
        
        print(f"[WARNING] No text extracted via OCR from {filename}")
        return "", filename
    except Exception as e:
        import traceback
        print(f"[ERROR] Error in OCR extraction from Word {filename}: {e}")
        print(traceback.format_exc())
        return "", filename

# Supported document and image extensions (documents for library + question generation)
DOC_EXTENSIONS = ('.pdf', '.docx', '.doc')
IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg')
ALL_DOC_AND_IMAGE_EXTENSIONS = DOC_EXTENSIONS + IMAGE_EXTENSIONS

def extract_text_from_image(image_file, progress_callback=None) -> Tuple[str, str]:
    """
    Extract text from image file (PNG, JPG, JPEG) using OCR.
    Uses EasyOCR first, then Tesseract fallback. Returns (text, filename).
    """
    filename = "unknown.png"
    try:
        if not OCR_AVAILABLE:
            if progress_callback:
                progress_callback("⚠ OCR not available for image files")
            return "", filename

        # Resolve path or bytes
        if isinstance(image_file, str):
            if not os.path.exists(image_file):
                return "", os.path.basename(image_file)
            filename = os.path.basename(image_file)
            try:
                img = Image.open(image_file)
            except Exception as e:
                print(f"[WARNING] Cannot open image {filename}: {e}")
                return "", filename
        else:
            # File-like object (uploaded file)
            filename = getattr(image_file, 'name', 'unknown.png') or 'unknown.png'
            try:
                image_file.seek(0)
                img = Image.open(io.BytesIO(image_file.read()))
            except Exception as e:
                print(f"[WARNING] Cannot open image {filename}: {e}")
                return "", filename

        # Convert to RGB if necessary (e.g. RGBA, P mode)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        import numpy as np
        img_array = np.array(img)

        ocr_text = ""
        # Try EasyOCR first (better for scanned documents and varied fonts)
        if EASYOCR_AVAILABLE:
            if progress_callback:
                progress_callback(f"🔍 Reading image with OCR: {filename}...")
            reader = get_easyocr_reader(progress_callback=progress_callback)
            if reader:
                try:
                    results = reader.readtext(img_array, paragraph=True, detail=0)
                    if results:
                        ocr_text = "\n".join(results).strip()
                        if progress_callback and ocr_text:
                            progress_callback(f"✓ Extracted {len(ocr_text)} chars from {filename}")
                except Exception as e:
                    print(f"[WARNING] EasyOCR failed for {filename}: {e}")

        # Fallback to Tesseract
        if not ocr_text and TESSERACT_AVAILABLE:
            if progress_callback:
                progress_callback(f"🔍 Trying Tesseract OCR on {filename}...")
            try:
                ocr_text = pytesseract.image_to_string(img, config='--psm 6').strip()
                if ocr_text and progress_callback:
                    progress_callback(f"✓ Extracted {len(ocr_text)} chars from {filename}")
            except Exception as e:
                print(f"[WARNING] Tesseract failed for {filename}: {e}")

        if ocr_text:
            print(f"[OK] OCR extracted {len(ocr_text)} characters from {filename}")
        return ocr_text, filename
    except Exception as e:
        import traceback
        print(f"[ERROR] extract_text_from_image {filename}: {e}")
        print(traceback.format_exc())
        return "", filename

def extract_text_from_pdf(pdf_file) -> Tuple[str, str]:
    """Extract text from PDF file with improved error handling"""
    try:
        if not PDF_AVAILABLE:
            return "", ""
        
        # Get filename
        if hasattr(pdf_file, 'name'):
            filename = pdf_file.name
        elif isinstance(pdf_file, str):
            filename = os.path.basename(pdf_file)
        else:
            filename = "unknown.pdf"
        
        text = ""
        extraction_success = False
        
        # 1) Try native text extraction first (pdfplumber or PyMuPDF)
        if PDF_LIB == "pdfplumber" and pdfplumber_module:
            # pdfplumber can handle file objects or paths
            try:
                if isinstance(pdf_file, str):
                    with pdfplumber_module.open(pdf_file) as pdf:
                        # Try multiple extraction methods
                        page_texts = []
                        for page in pdf.pages:
                            # Try extract_text first
                            page_text = page.extract_text()
                            if not page_text or len(page_text.strip()) < 10:
                                # Try extracting tables and text separately
                                page_text = page.extract_text(layout=True) or page.extract_text() or ""
                            if page_text:
                                page_texts.append(page_text)
                        text = "\n".join(page_texts)
                else:
                    # File object - need to reset position
                    pdf_file.seek(0)
                    with pdfplumber_module.open(pdf_file) as pdf:
                        page_texts = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if not page_text or len(page_text.strip()) < 10:
                                page_text = page.extract_text(layout=True) or page.extract_text() or ""
                            if page_text:
                                page_texts.append(page_text)
                        text = "\n".join(page_texts)
                if text and text.strip():
                    extraction_success = True
            except Exception as e:
                # Try fallback method
                pass
        
        if not extraction_success and PDF_LIB == "fitz" and fitz_module:
            # PyMuPDF (fitz) as fallback text extractor - try multiple methods
            try:
                if isinstance(pdf_file, str):
                    # File path
                    pdf_doc = fitz_module.open(pdf_file)
                else:
                    # File object
                    pdf_file.seek(0)
                    pdf_doc = fitz_module.open(stream=pdf_file.read(), filetype="pdf")
                
                # Try different text extraction methods
                page_texts = []
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    # Try standard text extraction
                    page_text = page.get_text()
                    if not page_text or len(page_text.strip()) < 10:
                        # Try with different options
                        page_text = page.get_text("text") or page.get_text("dict") or ""
                        if isinstance(page_text, dict):
                            # Extract text from dict format
                            page_text = " ".join([block.get("text", "") for block in page_text.get("blocks", [])])
                    
                    if page_text and page_text.strip():
                        page_texts.append(page_text)
                
                text = "\n".join(page_texts)
                pdf_doc.close()
                if text and text.strip():
                    extraction_success = True
            except Exception as e:
                pass

        # 2) If still no text and OCR is available, try EasyOCR first, then Tesseract
        if (not extraction_success or not text.strip()) and OCR_AVAILABLE and fitz_module:
            try:
                if isinstance(pdf_file, str):
                    doc = fitz_module.open(pdf_file)
                else:
                    pdf_file.seek(0)
                    doc = fitz_module.open(stream=pdf_file.read(), filetype="pdf")

                ocr_chunks = []
                # Process more pages for OCR (up to 30 pages for better coverage)
                max_ocr_pages = min(30, len(doc))
                
                # Try EasyOCR first (better for scanned documents)
                if EASYOCR_AVAILABLE:
                    # Pass a simple progress callback if available
                    reader = get_easyocr_reader()
                    if reader:
                        try:
                            for page_index in range(max_ocr_pages):
                                try:
                                    page = doc[page_index]
                                    # Get higher resolution for better OCR
                                    zoom = 2.0  # 2x zoom for better OCR accuracy
                                    mat = fitz_module.Matrix(zoom, zoom)
                                    pix = page.get_pixmap(matrix=mat)
                                    mode = "RGB" if pix.alpha == 0 else "RGBA"
                                    img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                                    
                                    # Convert to numpy array for EasyOCR
                                    import numpy as np
                                    img_array = np.array(img)
                                    
                                    # Use EasyOCR to read text with paragraph grouping
                                    results = reader.readtext(img_array, paragraph=True, detail=0)
                                    # results is a list of strings when detail=0
                                    page_text = "\n".join(results) if results else ""
                                    
                                    # Alternative: if detail=0 doesn't work, use detail=1 and extract text
                                    if not page_text:
                                        results_detailed = reader.readtext(img_array, paragraph=True)
                                        page_text = "\n".join([result[1] for result in results_detailed if len(result) > 1 and result[2] > 0.3])
                                    
                                    if page_text and page_text.strip():
                                        ocr_chunks.append(page_text)
                                except Exception as e:
                                    continue
                            
                            if ocr_chunks:
                                text = "\n".join(ocr_chunks)
                                extraction_success = bool(text.strip())
                        except Exception as e:
                            import traceback
                            print(f"EasyOCR error: {str(e)}")
                            print(traceback.format_exc())
                
                # Fallback to Tesseract if EasyOCR didn't work or isn't available
                if not extraction_success and TESSERACT_AVAILABLE:
                    try:
                        ocr_chunks = []
                        for page_index in range(max_ocr_pages):
                            try:
                                page = doc[page_index]
                                zoom = 2.0
                                mat = fitz_module.Matrix(zoom, zoom)
                                pix = page.get_pixmap(matrix=mat)
                                mode = "RGB" if pix.alpha == 0 else "RGBA"
                                img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                                ocr_text = pytesseract.image_to_string(img, config='--psm 6')
                                if ocr_text and ocr_text.strip():
                                    ocr_chunks.append(ocr_text)
                            except Exception:
                                continue
                        if ocr_chunks:
                            text = "\n".join(ocr_chunks)
                            extraction_success = bool(text.strip())
                    except Exception as e:
                        import traceback
                        print(f"Tesseract OCR error: {str(e)}")
                        print(traceback.format_exc())
                
                doc.close()
            except Exception as e:
                import traceback
                print(f"OCR error: {str(e)}")
                print(traceback.format_exc())
                pass
        
        # If still no text extracted, try one more time with more aggressive OCR
        if not text or not text.strip():
            # Last resort: try OCR on all pages if available
            if OCR_AVAILABLE and fitz_module:
                try:
                    if isinstance(pdf_file, str):
                        doc = fitz_module.open(pdf_file)
                    else:
                        pdf_file.seek(0)
                        doc = fitz_module.open(stream=pdf_file.read(), filetype="pdf")
                    
                    # Try EasyOCR on first 10 pages as last resort
                    if EASYOCR_AVAILABLE:
                        reader = get_easyocr_reader()
                        if reader:
                            ocr_texts = []
                            max_pages = min(10, len(doc))
                            for page_idx in range(max_pages):
                                try:
                                    page = doc[page_idx]
                                    zoom = 2.0
                                    mat = fitz_module.Matrix(zoom, zoom)
                                    pix = page.get_pixmap(matrix=mat)
                                    img = Image.frombytes("RGB" if pix.alpha == 0 else "RGBA", 
                                                         [pix.width, pix.height], pix.samples)
                                    import numpy as np
                                    img_array = np.array(img)
                                    results = reader.readtext(img_array, paragraph=True, detail=0)
                                    if results:
                                        ocr_texts.append("\n".join(results))
                                except Exception:
                                    continue
                            if ocr_texts:
                                text = "\n".join(ocr_texts)
                    doc.close()
                except Exception:
                    pass
        
        # If still no text, return empty but keep filename
        if not text or not text.strip():
            return "", filename
        
        # Clean text and be less strict about length
        text = " ".join(text.split())
        
        return text, filename
    except Exception as e:
        # Return empty text but keep filename for preview purposes
        return "", filename if isinstance(pdf_file, str) else (getattr(pdf_file, 'name', 'unknown.pdf'))

def chunk_text(text: str, chunk_size: int = 1000) -> List[str]:
    """Split text into chunks for processing"""
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    
    for word in words:
        if current_length + len(word) + 1 > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
        else:
            current_chunk.append(word)
            current_length += len(word) + 1
    
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks

# ============================================================================
# QUESTION GENERATION
# ============================================================================

def deduplicate_questions(questions: List[Dict]) -> List[Dict]:
    """Remove duplicate or near-duplicate questions"""
    unique_questions = []
    seen_texts = set()
    
    for q in questions:
        # Normalize question text for comparison
        q_text = q.get('question', '').lower().strip()
        # Remove extra whitespace and special chars
        q_text_normalized = ' '.join(q_text.split())
        
        # Check if similar question already exists
        is_duplicate = False
        for seen in seen_texts:
            # Simple similarity check - if 80% of words match, consider duplicate
            q_words = set(q_text_normalized.split())
            seen_words = set(seen.split())
            if len(q_words) > 0 and len(seen_words) > 0:
                similarity = len(q_words & seen_words) / max(len(q_words), len(seen_words))
                if similarity > 0.8:
                    is_duplicate = True
                    break
        
        if not is_duplicate:
            unique_questions.append(q)
            seen_texts.add(q_text_normalized)
    
    return unique_questions

def generate_questions_llm(text: str, difficulty: str, num_questions: int, question_types: List[str], progress_callback=None) -> List[Dict]:
    """Generate questions using OpenAI API if available"""
    if not LLM_AVAILABLE:
        return []
    
    api_key = get_openai_api_key()
    if not api_key:
        return []
    
    try:
        if progress_callback:
            progress_callback("Connecting to OpenAI API...")
        
        # Initialize OpenAI client - handle 'proxies' error
        # This error can occur even with correct versions due to internal library issues
        import os
        import sys
        
        # Get the actual OpenAI version being used
        try:
            openai_version = openai.__version__
        except AttributeError:
            openai_version = "unknown"
        
        # Temporarily remove ALL proxy-related environment variables
        saved_env = {}
        proxy_keys = ['HTTP_PROXY', 'HTTPS_PROXY', 'http_proxy', 'https_proxy', 
                     'ALL_PROXY', 'all_proxy', 'NO_PROXY', 'no_proxy',
                     'REQUESTS_CA_BUNDLE', 'CURL_CA_BUNDLE']
        for key in proxy_keys:
            if key in os.environ:
                saved_env[key] = os.environ.pop(key)
        
        # Also clear any proxy settings from httpx (used by OpenAI)
        try:
            import httpx
            # Clear any default proxy settings
            if hasattr(httpx, '_default_proxy'):
                httpx._default_proxy = None
        except:
            pass
        
        try:
            # Method 1: Try the simplest possible initialization
            # Don't pass http_client at all - let OpenAI create its own
            # This avoids any httpx version compatibility issues
            client = openai.OpenAI(api_key=api_key)
            
        except (TypeError, ValueError, AttributeError) as init_error:
            error_msg = str(init_error)
            if "proxies" in error_msg.lower() or "unexpected keyword" in error_msg.lower():
                # The 'proxies' error suggests httpx or OpenAI internal issue
                # Method 2: Try downgrading httpx or using a workaround
                try:
                    # Check httpx version and try to work around
                    import httpx
                    httpx_version = getattr(httpx, '__version__', 'unknown')
                    
                    # Try creating httpx client with explicit no-proxy config
                    # Some versions of httpx might have proxy detection that conflicts
                    http_client = httpx.Client(
                        timeout=httpx.Timeout(60.0),
                        follow_redirects=True
                    )
                    client = openai.OpenAI(api_key=api_key, http_client=http_client)
                except Exception as e2:
                    # Method 3: The issue is likely in OpenAI library itself
                    # Suggest downgrading OpenAI to a version that works
                    try:
                        # Last attempt: try with minimal httpx client
                        import httpx
                        # Create client with absolute minimum
                        http_client = httpx.Client(timeout=60.0)
                        client = openai.OpenAI(api_key=api_key, http_client=http_client)
                    except Exception as e3:
                        # All methods failed - the issue is likely in the library itself
                        if progress_callback:
                            progress_callback(f"❌ All client initialization methods failed")
                        
                        # The issue is httpx/OpenAI version incompatibility
                        # Provide clear solution
                        import httpx
                        httpx_version = getattr(httpx, '__version__', 'unknown')
                        
                        raise Exception(
                            f"❌ OpenAI client initialization failed: {error_msg}\n\n"
                            f"**Detected versions:**\n"
                            f"- OpenAI: {openai_version}\n"
                            f"- httpx: {httpx_version}\n\n"
                            f"**The 'proxies' error is a version compatibility issue.**\n\n"
                            f"**SOLUTION: Downgrade httpx to a compatible version:**\n\n"
                            f"```bash\n"
                            f"pip uninstall httpx -y\n"
                            f"pip install httpx==0.25.2\n"
                            f"```\n\n"
                            f"**OR downgrade OpenAI to a stable version:**\n\n"
                            f"```bash\n"
                            f"pip uninstall openai -y\n"
                            f"pip install openai==1.3.0\n"
                            f"```\n\n"
                            f"**After changing versions, restart Streamlit completely.**\n\n"
                            f"Secondary errors:\n- Method 2: {str(e2)}\n- Method 3: {str(e3)}"
                        )
            else:
                raise
        finally:
            # Restore environment variables
            for key, value in saved_env.items():
                os.environ[key] = value
        model = get_openai_model()
        temperature = get_openai_temperature()
        
        # Ensure we have enough text for generation
        if not text or len(text.strip()) < 100:
            if progress_callback:
                progress_callback("❌ Document text too short for AI generation")
            st.error("❌ Document text is too short. Please upload documents with more content.")
            return []
        
        # Prioritize situational questions
        situational_priority = "situational" in [qt.lower() for qt in question_types]
        if situational_priority:
            type_instruction = "Focus heavily on situational/scenario-based questions (at least 60%). Also include: " + ", ".join([qt for qt in question_types if qt.lower() != "situational"])
        else:
            type_instruction = ", ".join(question_types)
        
        prompt = f"""You are generating questions for the Philippine National Police (PNP) Criminology Licensure Examination. Generate EXACTLY {num_questions} {difficulty} level questions in proper PNP exam format based EXCLUSIVELY on the following review material from the document library. DO NOT generate more than {num_questions} questions.

CRITICAL: ALL questions MUST be based ONLY on the provided review material. Do NOT use general knowledge or information not found in the material below.

REVIEW MATERIAL FROM DOCUMENT LIBRARY:
{text[:4000]}

EXAM FORMAT REQUIREMENTS:
1. PRIORITY: Generate at least 60% situational/scenario-based questions that test practical application of criminology knowledge in real-world PNP scenarios.
2. Each question must be in proper PNP exam format with clear, complete answer choices.
3. Questions should test knowledge relevant to Philippine laws, PNP procedures, and criminology practice.

QUESTION TYPES (in priority order):
{type_instruction}

IMPORTANT: ALL questions MUST be Multiple Choice (MCQ) format ONLY. NO True/False, NO Identification, NO fill-in-the-blank questions.

QUESTION TYPE GUIDELINES (ALL must be MCQ format):
- Situational/Scenario-based (PRIORITY): Present a realistic scenario involving PNP operations, criminal investigation, or law enforcement. Ask what action should be taken, what law applies, what procedure to follow, or what the correct response is. Provide 4 multiple choice options. Example: "Officer Juan responds to a domestic violence call. The victim is bleeding but refuses medical treatment. What should Officer Juan do first?" with 4 MCQ options.
- Recall/Definition: Test knowledge of criminology terms, Philippine laws, or PNP procedures with 4 MCQ options
- Decision-making/Application: Ask how to apply specific laws, procedures, or protocols in given contexts with 4 MCQ options
- Ethics/Procedure: Focus on PNP Code of Conduct, ethical decision-making, and proper police procedures with 4 MCQ options

REQUIRED FORMAT FOR EACH QUESTION:
{{
  "question": "Complete question text with scenario if situational",
  "type": "One of: {', '.join(question_types)}",
  "options": ["First complete answer choice (full sentence or phrase)", "Second complete answer choice (full sentence or phrase)", "Third complete answer choice (full sentence or phrase)", "Fourth complete answer choice (full sentence or phrase)"],
  "correct_answer": "Exact text matching one of the four options above",
  "explanation": "2-3 sentence explanation of why this is correct, referencing relevant law or procedure"
}}

CRITICAL REQUIREMENTS - YOU MUST FOLLOW THESE EXACTLY:
1. ALL questions MUST be Multiple Choice (MCQ) format ONLY - NO True/False, NO Identification, NO fill-in-the-blank
2. ALL MCQ questions MUST have exactly 4 complete, meaningful answer options
3. NEVER EVER use generic placeholders like "Option A", "Option B", "Option C", "Option D" - these are STRICTLY FORBIDDEN
4. NEVER use just "A", "B", "C", "D" as options
5. NEVER ask questions that require user input like "Enter your answer:" - ALL questions must be multiple choice
6. Each option must be a complete, meaningful answer that directly relates to the question (minimum 10 characters, preferably a full sentence)
7. Options should be realistic, plausible alternatives that test actual knowledge
8. Example of CORRECT format:
   "options": ["Secure and preserve the crime scene to prevent contamination of evidence", "Arrest all persons present at the scene immediately without investigation", "Remove evidence for safekeeping in the police station before documentation", "Interview witnesses before securing the scene to gather information"]
9. Example of INCORRECT format (DO NOT USE - STRICTLY FORBIDDEN):
   "options": ["Option A", "Option B", "Option C", "Option D"]
   OR: ["HANDOUT", "Option C", "Option B", "Option D"]
   OR: ["A", "B", "C", "D"]
10. The correct_answer must exactly match one of the four option texts (word for word)
11. Questions must be based EXCLUSIVELY on the provided review material from the document library
12. Focus on Philippine laws, PNP procedures, and criminology practice

VALIDATION CHECKLIST - Verify each question before including:
✓ Does it have exactly 4 options?
✓ Are all options complete sentences/phrases with at least 10 characters?
✓ Are all options meaningful and related to the question?
✓ Are there NO generic placeholders like "Option A/B/C/D"?
✓ Does the correct_answer match one of the options exactly?
✓ If any question fails these checks, DO NOT include it - regenerate or exclude it.

Return ONLY a valid JSON array with EXACTLY {num_questions} question objects. Do not include any markdown formatting, code blocks, or explanations outside the JSON. IMPORTANT: The array must contain exactly {num_questions} questions, no more, no less.
"""
        
        if progress_callback:
            progress_callback("AI agent is generating questions based on your uploaded files. Please wait 30–60 seconds.")
        
        try:
            if progress_callback:
                progress_callback("AI agent is generating questions based on your uploaded files. Please wait 30–60 seconds.")
            # Make API call with explicit parameters
            # Note: timeout is set at client level, not in create() call
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature,
                max_tokens=5000  # Safe limit for gpt-4 (8192 total context - ~3000 for prompt = ~5000 for completion)
            )
            if progress_callback:
                progress_callback("✓ Received response from OpenAI API")
        except Exception as api_error:
            if progress_callback:
                progress_callback(f"❌ API request failed: {str(api_error)}")
            raise api_error
        
        if progress_callback:
            progress_callback("Parsing AI-generated questions...")
        
        # Parse JSON response
        import json
        content = response.choices[0].message.content.strip()
        
        # Try to extract JSON from markdown code blocks
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].split("```")[0].strip()
        
        try:
            questions = json.loads(content)
            if not isinstance(questions, list):
                questions = [questions]
            
            # Immediately limit to requested number before validation
            questions = questions[:num_questions]
            
            # Ensure all questions have required fields and validate format
            validated_questions = []
            for q in questions:
                # Ensure required fields exist
                if 'question' not in q or not q['question']:
                    continue
                if 'options' not in q or not q['options']:
                    # If MCQ but no options, skip or add default
                    if q.get('type', 'MCQ') == 'MCQ':
                        continue
                    q['options'] = []
                if 'type' not in q:
                    q['type'] = 'MCQ'
                if 'explanation' not in q:
                    q['explanation'] = ''
                if 'correct_answer' not in q:
                    continue
                
                # Validate MCQ options - must have 4 complete options
                if q['type'] == 'MCQ' and len(q['options']) < 4:
                    continue
                
                # Ensure options are not generic placeholders - STRICT VALIDATION
                if q['type'] == 'MCQ':
                    # Check for generic patterns: "Option A", "Option B", "Option C", "Option D"
                    generic_patterns = ['option a', 'option b', 'option c', 'option d', 'optiona', 'optionb', 'optionc', 'optiond']
                    has_generic = False
                    for opt in q['options']:
                        opt_lower = opt.lower().strip()
                        # Check if option is too short or matches generic pattern
                        if len(opt_lower.split()) <= 2 and any(pattern in opt_lower for pattern in generic_patterns):
                            has_generic = True
                            break
                        # Check if option is just "A", "B", "C", "D" or similar
                        if opt_lower in ['a', 'b', 'c', 'd', 'a.', 'b.', 'c.', 'd.']:
                            has_generic = True
                            break
                        # Check if option contains "Option" followed by letter and is very short
                        if 'option' in opt_lower and len(opt_lower.split()) <= 3:
                            has_generic = True
                            break
                        # Check if option is just a single letter or very short
                        if len(opt_lower.strip()) <= 3 and opt_lower.strip() in ['a', 'b', 'c', 'd', 'a.', 'b.', 'c.', 'd.', 'option a', 'option b', 'option c', 'option d']:
                            has_generic = True
                            break
                    
                    if has_generic:
                        continue  # Skip questions with generic options
                    
                    # Additional check: ensure all options have meaningful content (at least 5 characters)
                    if any(len(opt.strip()) < 5 for opt in q['options']):
                        continue
                    
                    # Check if options are too similar (might be duplicates) - but allow some similarity
                    unique_options = set(opt.lower().strip() for opt in q['options'])
                    if len(unique_options) < 2:  # At least 2 unique options
                        continue
                
                # Ensure correct_answer matches one of the options
                if q['type'] == 'MCQ' and q['correct_answer'] not in q['options']:
                    # Try to find a match (case-insensitive)
                    matched = False
                    for opt in q['options']:
                        if opt.strip().lower() == q['correct_answer'].strip().lower():
                            q['correct_answer'] = opt  # Use the exact option text
                            matched = True
                            break
                    if not matched:
                        continue  # Skip if no match found
                
                validated_questions.append(q)
            
            if not validated_questions:
                if progress_callback:
                    progress_callback("⚠️ AI generated invalid questions (generic options detected), retrying...")
                # Show debug info
                st.warning(f"⚠️ First AI attempt generated {len(questions)} questions but all were rejected due to generic options or validation failures.")
                st.info("💡 Retrying with stricter prompt...")
                # Retry with an even more explicit prompt
                retry_prompt = f"""Generate EXACTLY {num_questions} {difficulty} level PNP Criminology exam questions based on this material. DO NOT generate more than {num_questions} questions:

{text[:3000]}

CRITICAL RULES:
1. Each question MUST have exactly 4 answer options
2. NEVER use "Option A", "Option B", "Option C", "Option D" - these are FORBIDDEN
3. Each option must be a complete sentence or phrase (minimum 5 words)
4. Options must be meaningful answers related to the question

CORRECT EXAMPLE:
{{
  "question": "What is the first step when securing a crime scene?",
  "type": "Situational",
  "options": ["Establish a perimeter to prevent unauthorized access", "Collect all evidence immediately", "Interview witnesses on the spot", "Allow media to enter and document"],
  "correct_answer": "Establish a perimeter to prevent unauthorized access",
  "explanation": "Securing the perimeter is the first priority to preserve evidence integrity."
}}

Return ONLY a JSON array with EXACTLY {num_questions} questions. Each option must be a real answer, not a placeholder. DO NOT generate more than {num_questions} questions.
"""
                try:
                    retry_response = client.chat.completions.create(
                        model=model,
                        messages=[{"role": "user", "content": retry_prompt}],
                        temperature=temperature,
                        max_tokens=5000  # Safe limit for gpt-4 (8192 total context - ~3000 for prompt = ~5000 for completion)
                    )
                    retry_content = retry_response.choices[0].message.content.strip()
                    if "```json" in retry_content:
                        retry_content = retry_content.split("```json")[1].split("```")[0].strip()
                    elif "```" in retry_content:
                        retry_content = retry_content.split("```")[1].split("```")[0].strip()
                    
                    retry_questions = json.loads(retry_content)
                    if not isinstance(retry_questions, list):
                        retry_questions = [retry_questions]
                    
                    # Immediately limit to requested number before validation
                    retry_questions = retry_questions[:num_questions]
                    
                    # Validate retry questions with same strict rules
                    validated_retry = []
                    for q in retry_questions:
                        if 'question' not in q or not q['question']:
                            continue
                        if 'options' not in q or len(q.get('options', [])) < 4:
                            continue
                        if 'correct_answer' not in q:
                            continue
                        if 'type' not in q:
                            q['type'] = 'MCQ'
                        if 'explanation' not in q:
                            q['explanation'] = ''
                        
                        # Strict validation for generic options
                        has_generic = False
                        for opt in q['options']:
                            opt_lower = opt.lower().strip()
                            if len(opt_lower.split()) <= 2 and any(pattern in opt_lower for pattern in ['option a', 'option b', 'option c', 'option d']):
                                has_generic = True
                                break
                            if opt_lower in ['a', 'b', 'c', 'd', 'a.', 'b.', 'c.', 'd.']:
                                has_generic = True
                                break
                            if 'option' in opt_lower and len(opt_lower.split()) <= 3:
                                has_generic = True
                                break
                            if len(opt.strip()) < 5:
                                has_generic = True
                                break
                        
                        if has_generic:
                            continue
                        
                        # Ensure correct_answer matches
                        if q['correct_answer'] not in q['options']:
                            matched = False
                            for opt in q['options']:
                                if opt.strip().lower() == q['correct_answer'].strip().lower():
                                    q['correct_answer'] = opt
                                    matched = True
                                    break
                            if not matched:
                                continue
                        
                        validated_retry.append(q)
                    
                    if validated_retry:
                        validated_retry = deduplicate_questions(validated_retry)
                        # Strictly limit to requested number
                        validated_retry = validated_retry[:num_questions]
                        if progress_callback:
                            progress_callback(f"✅ Generated {len(validated_retry)} valid questions on retry!")
                        return validated_retry
                    else:
                        if progress_callback:
                            progress_callback("❌ Retry also produced invalid questions. Falling back to rule-based...")
                        st.warning("⚠️ AI retry also failed. Using rule-based generation from your documents...")
                except Exception as retry_e:
                    if progress_callback:
                        progress_callback(f"❌ Retry also failed: {str(retry_e)}")
                    st.warning(f"Retry failed: {str(retry_e)}. Falling back to rule-based generation...")
                
                # Don't return empty - let it fall through to rule-based
                return []
            
            # Deduplicate
            validated_questions = deduplicate_questions(validated_questions)
            
            if progress_callback:
                progress_callback(f"✅ Generated {len(validated_questions)} valid AI questions in PNP format!")
            
            return validated_questions[:num_questions]
        except json.JSONDecodeError as e:
            if progress_callback:
                progress_callback(f"⚠️ JSON parsing error: {str(e)}")
            st.warning(f"Failed to parse AI response as JSON: {str(e)}")
            return []
    except Exception as e:
        if progress_callback:
            progress_callback(f"❌ Error: {str(e)}")
        error_msg = str(e)
        st.error(f"❌ AI generation failed: {error_msg}")
        # Don't show API key errors to users
        pass
        return []

def generate_questions_rule_based(text: str, difficulty: str, num_questions: int, question_types: List[str]) -> List[Dict]:
    """Generate questions using rule-based approach"""
    chunks = chunk_text(text, chunk_size=500)
    if not chunks:
        return []
    
    questions = []
    words = text.split()
    
    # Simple question generation based on text patterns
    for i in range(min(num_questions, len(chunks) * 2)):
        chunk = chunks[i % len(chunks)]
        chunk_words = chunk.split()
        
        if len(chunk_words) < 5:
            continue
        
        q_type = secrets.choice(question_types) if question_types else "MCQ"
        
        if q_type == "True/False":
            # Create True/False question
            question_text = f"True or False: {chunk[:200]}"
            questions.append({
                "question": question_text,
                "type": "True/False",
                "options": ["True", "False"],
                "correct_answer": "True",  # Simplified
                "explanation": chunk[:150]
            })
        
        elif q_type == "Identification":
            # Create fill-in-the-blank
            if len(chunk_words) > 3:
                blank_word = chunk_words[2]
                question_text = chunk.replace(blank_word, "_____", 1)
                questions.append({
                    "question": question_text[:300],
                    "type": "Identification",
                    "options": [],
                    "correct_answer": blank_word,
                    "explanation": chunk[:150]
                })
        
        else:  # MCQ
            # Create multiple choice with meaningful options
            question_text = f"What is described in the following: {chunk[:150]}?"
            # Extract key terms from chunk for better options
            key_terms = [w for w in chunk_words[:10] if len(w) > 4][:4]  # Get meaningful words
            if len(key_terms) < 4:
                # Pad with related terms
                key_terms.extend(["procedure", "regulation", "principle", "requirement"][:4-len(key_terms)])
            
            correct = key_terms[0] if key_terms else "The described concept"
            # Create meaningful options based on chunk content
            options = [
                correct,
                f"Alternative approach to {key_terms[1] if len(key_terms) > 1 else 'the concept'}",
                f"Different method involving {key_terms[2] if len(key_terms) > 2 else 'related elements'}",
                f"Opposite or unrelated concept to {key_terms[3] if len(key_terms) > 3 else 'the main idea'}"
            ]
            secrets.SystemRandom().shuffle(options)
            
            questions.append({
                "question": question_text,
                "type": "MCQ",
                "options": options,
                "correct_answer": correct,
                "explanation": chunk[:150]
            })
        
        if len(questions) >= num_questions:
            break
    
    return questions[:num_questions]

def generate_default_dummy_questions(difficulty: str, num_questions: int, question_types: List[str]) -> List[Dict]:
    """Generate default dummy questions when no PDF/document is loaded"""
    # Pre-defined criminology questions pool
    dummy_questions_pool = [
        {
            "question": "What is the primary purpose of the Revised Penal Code in the Philippines?",
            "type": "MCQ",
            "options": ["To define criminal offenses and penalties", "To establish court procedures", "To regulate police operations", "To manage prison systems"],
            "correct_answer": "To define criminal offenses and penalties",
            "explanation": "The Revised Penal Code (Act No. 3815) is the main criminal law that defines offenses and prescribes penalties."
        },
        {
            "question": "True or False: A person is presumed innocent until proven guilty beyond reasonable doubt.",
            "type": "True/False",
            "options": ["True", "False"],
            "correct_answer": "True",
            "explanation": "This is a fundamental principle in criminal law - the presumption of innocence."
        },
        {
            "question": "What does 'nullum crimen, nulla poena sine lege' mean?",
            "type": "Identification",
            "options": [],
            "correct_answer": "No crime, no punishment without law",
            "explanation": "This Latin maxim means that no act is a crime unless it is clearly defined and penalized by law."
        },
        {
            "question": "Which of the following is NOT an element of a crime?",
            "type": "MCQ",
            "options": ["Human being", "Criminal act or omission", "Criminal intent", "Police report"],
            "correct_answer": "Police report",
            "explanation": "The elements of a crime are: human being, criminal act/omission, criminal intent, and the act must be punishable by law."
        },
        {
            "question": "True or False: Attempted felony is punishable under the Revised Penal Code.",
            "type": "True/False",
            "options": ["True", "False"],
            "correct_answer": "True",
            "explanation": "Attempted felony is one of the stages of a crime and is punishable, though with a lesser penalty than consummated felony."
        },
        {
            "question": "What is the maximum penalty for crimes punishable by reclusion perpetua?",
            "type": "MCQ",
            "options": ["20 years and 1 day to 40 years", "Life imprisonment", "6 years and 1 day to 12 years", "12 years and 1 day to 20 years"],
            "correct_answer": "20 years and 1 day to 40 years",
            "explanation": "Reclusion perpetua ranges from 20 years and 1 day to 40 years of imprisonment."
        },
        {
            "question": "Identify the term: The body of law that regulates the apprehension and prosecution of persons accused of crimes.",
            "type": "Identification",
            "options": [],
            "correct_answer": "Criminal procedure",
            "explanation": "Criminal procedure refers to the method prescribed by law for the apprehension and prosecution of persons accused of crimes."
        },
        {
            "question": "True or False: The accused has the right to remain silent and to have counsel during custodial investigation.",
            "type": "True/False",
            "options": ["True", "False"],
            "correct_answer": "True",
            "explanation": "This is a constitutional right under the Miranda doctrine and Philippine Constitution."
        },
        {
            "question": "What is the standard of proof required in criminal cases?",
            "type": "MCQ",
            "options": ["Preponderance of evidence", "Beyond reasonable doubt", "Substantial evidence", "Clear and convincing evidence"],
            "correct_answer": "Beyond reasonable doubt",
            "explanation": "Beyond reasonable doubt is the highest standard of proof required in criminal cases."
        },
        {
            "question": "True or False: A warrant of arrest is required for all arrests.",
            "type": "True/False",
            "options": ["True", "False"],
            "correct_answer": "False",
            "explanation": "Warrantless arrests are allowed in certain circumstances, such as when a crime is committed in the presence of the arresting officer."
        },
        {
            "question": "What is the term for a crime that is both a felony and a violation?",
            "type": "MCQ",
            "options": ["Complex crime", "Compound crime", "Continuing crime", "Composite crime"],
            "correct_answer": "Complex crime",
            "explanation": "A complex crime is a single act that constitutes two or more grave or less grave felonies, or an offense and a violation."
        },
        {
            "question": "Identify: The principle that no person shall be held to answer for a criminal offense without due process of law.",
            "type": "Identification",
            "options": [],
            "correct_answer": "Due process",
            "explanation": "Due process ensures that no person is deprived of life, liberty, or property without following proper legal procedures."
        },
        {
            "question": "True or False: Double jeopardy prevents a person from being tried twice for the same offense.",
            "type": "True/False",
            "options": ["True", "False"],
            "correct_answer": "True",
            "explanation": "Double jeopardy is a constitutional protection that prevents multiple prosecutions for the same offense."
        },
        {
            "question": "What is the minimum age of criminal responsibility in the Philippines?",
            "type": "MCQ",
            "options": ["12 years old", "15 years old", "18 years old", "9 years old"],
            "correct_answer": "15 years old",
            "explanation": "Under Republic Act No. 9344, as amended, the minimum age of criminal responsibility is 15 years old."
        },
        {
            "question": "True or False: Self-defense is a justifying circumstance that exempts a person from criminal liability.",
            "type": "True/False",
            "options": ["True", "False"],
            "correct_answer": "True",
            "explanation": "Self-defense is a justifying circumstance that exempts a person from criminal liability when all its elements are present."
        }
    ]
    
    # Filter by difficulty and question types
    selected_questions = []
    available_types = question_types if question_types else ["MCQ", "True/False", "Identification"]
    
    # Shuffle and select questions
    import random
    random.shuffle(dummy_questions_pool)
    
    for q in dummy_questions_pool:
        if q["type"] in available_types:
            selected_questions.append(q)
            if len(selected_questions) >= num_questions:
                break
    
    # If we need more questions, repeat with different variations
    while len(selected_questions) < num_questions:
        for q in dummy_questions_pool:
            if q["type"] in available_types and len(selected_questions) < num_questions:
                # Create a variation
                new_q = q.copy()
                new_q["question"] = q["question"] + " (Question " + str(len(selected_questions) + 1) + ")"
                selected_questions.append(new_q)
    
    return selected_questions[:num_questions]

def generate_questions(text: str, difficulty: str, num_questions: int, question_types: List[str], progress_callback=None) -> List[Dict]:
    """Generate questions using AI from documents. Only use dummy questions if no documents are available."""
    if progress_callback:
        progress_callback("Initializing question generation...")
    
    # Check if we have document text (from admin or user uploads)
    has_document_text = text and len(text.strip()) >= 50
    
    # If no document text, use default dummy questions
    if not has_document_text:
        if progress_callback:
            progress_callback("No document text found, using default questions...")
        questions = generate_default_dummy_questions(difficulty, num_questions, question_types)
        questions = deduplicate_questions(questions)
        if progress_callback:
            progress_callback(f"Generated {len(questions)} default questions!")
        return questions
    
    # We have document text - MUST use AI generation (no fallback to rule-based for document-based questions)
    if progress_callback:
        progress_callback("✓ Document text ready. Starting AI question generation...")
    
    # Try AI/LLM generation (REQUIRED for document-based questions)
    api_key = get_openai_api_key()
    if not api_key:
        if progress_callback:
            progress_callback("❌ AI generation unavailable")
        # Don't show API key errors to users - silently fall back
        return []
    
    # API key loaded (no user-facing message)
    if progress_callback:
        progress_callback("✓ AI generation ready")
    
    if progress_callback:
        progress_callback("🤖 Connecting to OpenAI API...")
    
    # Try AI generation with retries
    max_retries = 2
    for attempt in range(max_retries):
        try:
            if attempt > 0:
                if progress_callback:
                    progress_callback(f"🔄 Retrying AI generation (attempt {attempt + 1}/{max_retries})...")
            
            llm_questions = generate_questions_llm(text, difficulty, num_questions, question_types, progress_callback)
            
            if llm_questions and len(llm_questions) > 0:
                # Strictly limit to requested number
                llm_questions = llm_questions[:num_questions]
                if progress_callback:
                    progress_callback(f"✅ Generated {len(llm_questions)} AI-powered questions from your documents!")
                st.success(f"✅ Successfully generated {len(llm_questions)} AI questions from your documents!")
                return llm_questions
            else:
                if progress_callback:
                    progress_callback(f"⚠️ AI generation attempt {attempt + 1} returned no valid questions")
                if attempt < max_retries - 1:
                    continue  # Retry
                else:
                    # Last attempt failed
                    st.error("❌ AI generation failed after multiple attempts. Please check:")
                    st.markdown("""
                    - Is your OpenAI API key valid and has credits?
                    - Are the documents properly formatted and contain readable text?
                    - Try selecting fewer documents or reducing the number of questions
                    """)
                    return []
                    
        except Exception as e:
            error_msg = str(e)
            if progress_callback:
                progress_callback(f"❌ AI generation error (attempt {attempt + 1}): {error_msg}")
            
            if attempt < max_retries - 1:
                st.warning(f"AI generation attempt {attempt + 1} failed: {error_msg}. Retrying...")
                continue
            else:
                # Last attempt failed
                st.error(f"❌ AI generation failed after {max_retries} attempts: {error_msg}")
                st.info("""
                **Troubleshooting:**
                1. Please try again or contact support
                2. Verify you have API credits available
                3. Check your internet connection
                4. Try with fewer documents or fewer questions
                """)
                import traceback
                with st.expander("🔍 Detailed Error Information"):
                    st.code(traceback.format_exc())
                return []
    
    # Should never reach here, but just in case
    return []

# ============================================================================
# DOCUMENT LIBRARY MANAGEMENT
# ============================================================================

@st.cache_data(ttl=60, show_spinner=False)  # Cache for 60 seconds, no spinner
def get_document_library(user_email: Optional[str] = None, user_access_level: Optional[str] = None) -> List[Dict]:
    """Scan and return documents: Admin docs for all Advance/Premium users (existing and new); user docs filtered by email."""
    documents = []
    
    def scan_directory(directory: str, source: str, filter_email: Optional[str] = None) -> List[Dict]:
        """Helper to scan a directory for documents, optionally filtered by user email"""
        dir_docs = []
        if not os.path.exists(directory):
            return dir_docs
        try:
            for filename in os.listdir(directory):
                if filename.lower().endswith(ALL_DOC_AND_IMAGE_EXTENSIONS):
                    filepath = os.path.join(directory, filename)
                    try:
                        # Check if file belongs to the user (email in filename or check database)
                        if filter_email and source == "Uploaded":
                            # Check if filename contains user email or check database
                            user_email_clean = filter_email.lower().replace('@', '_').replace('.', '_')
                            if user_email_clean not in filename.lower():
                                # Check database for user association
                                table_name = get_table_name("pdf_resources")
                                cursor = execute_query(f"""
                                    SELECT uploaded_by FROM {table_name} 
                                    WHERE filepath = %s OR filename = %s
                                """, (filepath, filename))
                                result = cursor.fetchone()
                                cursor.close()
                                if result and result[0] and result[0].lower() != filter_email.lower():
                                    continue  # Skip files not uploaded by this user
                                elif not result:
                                    # If not in database, check filename pattern
                                    if user_email_clean not in filename.lower():
                                        continue
                        
                        file_size = os.path.getsize(filepath)
                        mtime = os.path.getmtime(filepath)
                        
                        # Check database for downloadable and premium_only status (Admin docs)
                        is_downloadable = True  # Default
                        is_premium_only = False  # Default
                        if source == "Admin":
                            columns_check = get_table_columns("pdf_resources")
                            has_downloadable_col = 'is_downloadable' in columns_check
                            table_name = get_table_name("pdf_resources")
                            if has_downloadable_col:
                                cursor_check = execute_query(f"""
                                    SELECT is_downloadable, is_premium_only FROM {table_name} 
                                    WHERE filepath = %s OR filename = %s
                                """, (filepath, filename))
                            else:
                                cursor_check = execute_query(f"""
                                    SELECT is_premium_only FROM {table_name} 
                                    WHERE filepath = %s OR filename = %s
                                """, (filepath, filename))
                            result_check = cursor_check.fetchone()
                            cursor_check.close()
                            if result_check:
                                if has_downloadable_col and len(result_check) >= 2:
                                    is_downloadable = bool(result_check[0])
                                    is_premium_only = bool(result_check[1])
                                elif has_downloadable_col:
                                    is_downloadable = bool(result_check[0])
                                else:
                                    is_premium_only = bool(result_check[0])
                        
                        dir_docs.append({
                            "filename": filename,
                            "filepath": filepath,
                            "source": source,
                            "date": datetime.fromtimestamp(mtime).strftime("%Y-%m-%d"),
                            "size": file_size,
                            "size_mb": round(file_size / (1024 * 1024), 2),
                            "type": _doc_type_from_filename(filename),
                            "is_downloadable": is_downloadable,
                            "is_premium_only": is_premium_only if source == "Admin" else False
                        })
                    except (OSError, ValueError):
                        continue  # Skip files that can't be accessed
        except (OSError, PermissionError):
            pass  # Skip directories that can't be accessed
        return dir_docs

def _doc_type_from_filename(filename: str) -> str:
    """Return display type: PDF, Word, or Image from filename."""
    if not filename:
        return "Document"
    low = filename.lower()
    if low.endswith('.pdf'):
        return "PDF"
    if low.endswith(('.docx', '.doc')):
        return "Word"
    if low.endswith(IMAGE_EXTENSIONS):
        return "Image"
    return "Document"

def _mime_for_doc(doc: dict) -> str:
    """Return MIME type for document download from doc dict (type + filename)."""
    t = doc.get("type", "Document")
    if t == "PDF":
        return "application/pdf"
    if t == "Word":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if t == "Image":
        fn = (doc.get("filename") or "").lower()
        if fn.endswith(".png"):
            return "image/png"
        return "image/jpeg"
    return "application/octet-stream"
    
    # 1. Admin-managed PDFs — always include for Advance and Premium (existing and new users)
    # Identify admin docs by uploader: uploaded_by = admin user (is_admin=1) or literal "admin"
    if user_access_level in ("Advance", "Premium"):
        os.makedirs(ADMIN_DOCS_DIR, exist_ok=True)
        admin_paths_seen = set()
        
        try:
            users_table = get_table_name("users")
            pdf_table = get_table_name("pdf_resources")
            # Get admin user emails (users with is_admin = 1)
            cursor = execute_query(f"""
                SELECT email FROM {users_table} WHERE is_admin = 1
            """)
            admin_emails = [row[0].strip().lower() for row in cursor.fetchall() if row and row[0]]
            cursor.close()
            admin_emails = list(set(admin_emails))
            admin_emails.append("admin")  # Legacy: uploads with uploaded_by = "admin"
            
            if not admin_emails:
                admin_emails = ["admin"]
            
            cols = get_table_columns("pdf_resources")
            has_dl_col = cols and "is_downloadable" in [c.lower() for c in cols]
            has_premium_col = cols and "is_premium_only" in [c.lower() for c in cols]
            has_uploaded_by = cols and "uploaded_by" in [c.lower() for c in cols]
            
            if not has_uploaded_by:
                # Fallback: filter by filepath containing admin_docs
                cursor = execute_query(f"""
                    SELECT filepath, filename, is_downloadable, is_premium_only FROM {pdf_table}
                """) if (has_dl_col and has_premium_col) else execute_query(f"""
                    SELECT filepath, filename FROM {pdf_table}
                """)
                all_rows = cursor.fetchall()
                cursor.close()
                admin_rows = [r for r in all_rows if r[0] and ("admin_docs" in (r[0] or "").lower())]
            else:
                # Primary: get pdf_resources where uploaded_by IN (admin emails)
                placeholders = ", ".join(["%s"] * len(admin_emails))
                if has_dl_col and has_premium_col:
                    cursor = execute_query(f"""
                        SELECT filepath, filename, is_downloadable, is_premium_only FROM {pdf_table}
                        WHERE LOWER(TRIM(COALESCE(uploaded_by, ''))) IN ({placeholders})
                    """, tuple(e.lower().strip() for e in admin_emails))
                else:
                    cursor = execute_query(f"""
                        SELECT filepath, filename FROM {pdf_table}
                        WHERE LOWER(TRIM(COALESCE(uploaded_by, ''))) IN ({placeholders})
                    """, tuple(e.lower().strip() for e in admin_emails))
                admin_rows = cursor.fetchall()
                cursor.close()
            
            for row in admin_rows:
                doc_path = (row[0] or "").strip()
                doc_filename = (row[1] or "").strip()
                is_downloadable = bool(row[2]) if len(row) > 2 and has_dl_col else True
                is_premium_only = bool(row[3]) if len(row) > 3 and has_premium_col else False
                if not doc_filename or not doc_filename.lower().endswith(ALL_DOC_AND_IMAGE_EXTENSIONS):
                    continue
                # Resolve path: stored path may be from another machine; try current app's admin_docs
                resolved_path = doc_path
                if not os.path.exists(resolved_path):
                    resolved_path = os.path.join(ADMIN_DOCS_DIR, os.path.basename(doc_path))
                if not os.path.exists(resolved_path):
                    resolved_path = os.path.join(ADMIN_DOCS_DIR, doc_filename)
                if os.path.exists(resolved_path) and resolved_path not in admin_paths_seen:
                    admin_paths_seen.add(resolved_path)
                    try:
                        file_size = os.path.getsize(resolved_path)
                        mtime = os.path.getmtime(resolved_path)
                        documents.append({
                            "filename": doc_filename,
                            "filepath": resolved_path,
                            "source": "Admin",
                            "date": datetime.fromtimestamp(mtime).strftime("%Y-%m-%d"),
                            "size": file_size,
                            "size_mb": round(file_size / (1024 * 1024), 2),
                            "type": _doc_type_from_filename(doc_filename),
                            "is_downloadable": is_downloadable,
                            "is_premium_only": is_premium_only
                        })
                    except (OSError, ValueError):
                        continue
        except Exception:
            pass
        
        # Also scan admin_docs folder for any files not already in list (e.g. uploaded but not in DB yet)
        admin_from_scan = scan_directory(ADMIN_DOCS_DIR, "Admin")
        for d in admin_from_scan:
            if d["filepath"] not in admin_paths_seen:
                documents.append(d)
    
    # 2. User-uploaded PDFs from /uploads/ (filtered by user email)
    if user_email:
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        # Get user's documents from database
        table_name = get_table_name("pdf_resources")
        cursor = execute_query(f"""
            SELECT filepath, filename FROM {table_name} 
            WHERE uploaded_by = %s
        """, (user_email.lower(),))
        user_docs = cursor.fetchall()
        cursor.close()
        
        # Add documents from database
        for doc_path, doc_filename in user_docs:
            if os.path.exists(doc_path):
                try:
                    file_size = os.path.getsize(doc_path)
                    mtime = os.path.getmtime(doc_path)
                    documents.append({
                        "filename": doc_filename,
                        "filepath": doc_path,
                        "source": "Uploaded",
                        "date": datetime.fromtimestamp(mtime).strftime("%Y-%m-%d"),
                        "size": file_size,
                        "size_mb": round(file_size / (1024 * 1024), 2),
                        "type": _doc_type_from_filename(doc_filename)
                    })
                except (OSError, ValueError):
                    continue
        
        # Also scan uploads directory for files with user email in filename
        user_email_clean = user_email.lower().replace('@', '_').replace('.', '_')
        if os.path.exists(UPLOADS_DIR):
            try:
                for filename in os.listdir(UPLOADS_DIR):
                    if filename.lower().endswith(ALL_DOC_AND_IMAGE_EXTENSIONS):
                        # Check if filename contains user email pattern
                        if user_email_clean in filename.lower():
                            filepath = os.path.join(UPLOADS_DIR, filename)
                            if filepath not in [d['filepath'] for d in documents]:  # Avoid duplicates
                                try:
                                    file_size = os.path.getsize(filepath)
                                    mtime = os.path.getmtime(filepath)
                                    documents.append({
                                        "filename": filename,
                                        "filepath": filepath,
                                        "source": "Uploaded",
                                        "date": datetime.fromtimestamp(mtime).strftime("%Y-%m-%d"),
                                        "size": file_size,
                                        "size_mb": round(file_size / (1024 * 1024), 2),
                                        "type": _doc_type_from_filename(filename)
                                    })
                                except (OSError, ValueError):
                                    continue
            except (OSError, PermissionError):
                pass
    
    return documents

def extract_text_from_documents(document_paths: List[str], max_pages_per_doc: int = None, max_total_chars: int = None, progress_callback=None) -> str:
    """Extract and combine text from multiple documents (no limits) - cloud-compatible"""
    combined_text = []
    total_chars = 0
    
    for idx, doc_path in enumerate(document_paths):
        # Remove character limit check - process all documents fully
        # if max_total_chars and total_chars >= max_total_chars:
        #     break
        
        # Normalize path for cloud compatibility
        if not os.path.isabs(doc_path):
            # Try relative to base directory first
            abs_path = os.path.join(BASE_DIR, doc_path)
            if os.path.exists(abs_path):
                doc_path = abs_path
            # Also try just the filename in uploads/admin_docs
            filename = os.path.basename(doc_path)
            for base_dir in [UPLOADS_DIR, ADMIN_DOCS_DIR]:
                test_path = os.path.join(base_dir, filename)
                if os.path.exists(test_path):
                    doc_path = test_path
                    break
            
        if not os.path.exists(doc_path):
            if progress_callback:
                progress_callback(f"⚠ Skipping {os.path.basename(doc_path)} (not found at {doc_path})...")
            print(f"[WARNING] File not found: {doc_path}")
            continue
        
        # Check file permissions
        if not os.access(doc_path, os.R_OK):
            if progress_callback:
                progress_callback(f"⚠ Skipping {os.path.basename(doc_path)} (permission denied)...")
            print(f"[WARNING] Permission denied: {doc_path}")
            continue
        
        try:
            if progress_callback:
                progress_callback(f"Extracting text from {os.path.basename(doc_path)} ({idx+1}/{len(document_paths)})...")
            
            if doc_path.lower().endswith('.pdf'):
                if PDF_AVAILABLE:
                    text, _ = extract_text_from_pdf(doc_path)
                    if text and text.strip():
                        # No character limit - extract full document
                        combined_text.append(text)
                        total_chars += len(text)
                        if progress_callback:
                            progress_callback(f"✓ Extracted {len(text)} chars from {os.path.basename(doc_path)}")
                    else:
                        # No text extracted - try OCR explicitly
                        if progress_callback:
                            progress_callback(f"⚠ No text found in {os.path.basename(doc_path)}, attempting OCR...")
                        # Force OCR attempt by calling extract_text_from_pdf again with OCR flag
                        # The function should automatically try OCR if initial extraction fails
                        if OCR_AVAILABLE:
                            if progress_callback:
                                progress_callback(f"🔍 Using OCR to extract text from {os.path.basename(doc_path)}...")
                            # Try extraction again - OCR should be attempted automatically
                            text, _ = extract_text_from_pdf(doc_path)
                            if text and text.strip():
                                combined_text.append(text)
                                total_chars += len(text)
                                if progress_callback:
                                    progress_callback(f"✓ OCR extracted {len(text)} chars from {os.path.basename(doc_path)}")
                            else:
                                if progress_callback:
                                    progress_callback(f"❌ Could not extract text from {os.path.basename(doc_path)} even with OCR")
                else:
                    if progress_callback:
                        progress_callback(f"⚠ PDF library not available for {os.path.basename(doc_path)}")
            elif doc_path.lower().endswith(('.docx', '.doc')):
                if DOCX_AVAILABLE:
                    text, _ = extract_text_from_docx(doc_path)
                    if text and text.strip():
                        # No character limit - extract full document
                        combined_text.append(text)
                        total_chars += len(text)
                        if progress_callback:
                            progress_callback(f"✓ Extracted {len(text)} chars from {os.path.basename(doc_path)}")
                    else:
                        # No text extracted - try OCR explicitly for Word documents
                        if progress_callback:
                            progress_callback(f"⚠ No text found in {os.path.basename(doc_path)}, attempting OCR...")
                        if OCR_AVAILABLE:
                            try:
                                # Try OCR on embedded images in Word document
                                ocr_text, _ = extract_text_from_docx_with_ocr(doc_path, os.path.basename(doc_path))
                                if ocr_text and ocr_text.strip():
                                    combined_text.append(ocr_text)
                                    total_chars += len(ocr_text)
                                    if progress_callback:
                                        progress_callback(f"✓ OCR extracted {len(ocr_text)} chars from {os.path.basename(doc_path)}")
                                else:
                                    if progress_callback:
                                        progress_callback(f"❌ Could not extract text from {os.path.basename(doc_path)} even with OCR")
                            except Exception as e:
                                if progress_callback:
                                    progress_callback(f"❌ OCR failed for {os.path.basename(doc_path)}: {str(e)}")
                        else:
                            if progress_callback:
                                progress_callback(f"⚠ No text found in {os.path.basename(doc_path)} - document may be empty or image-based (OCR not available)")
                else:
                    if progress_callback:
                        progress_callback(f"⚠ Word library not available for {os.path.basename(doc_path)}")
            elif doc_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                # Image files: extract text via OCR
                if progress_callback:
                    progress_callback(f"🔍 Reading image with OCR: {os.path.basename(doc_path)}...")
                text, _ = extract_text_from_image(doc_path, progress_callback=progress_callback)
                if text and text.strip():
                    combined_text.append(text)
                    total_chars += len(text)
                    if progress_callback:
                        progress_callback(f"✓ OCR extracted {len(text)} chars from {os.path.basename(doc_path)}")
                else:
                    if progress_callback:
                        progress_callback(f"⚠ No text found in image {os.path.basename(doc_path)} (OCR returned empty)")
        except Exception as e:
            error_msg = f"Error processing {os.path.basename(doc_path)}: {str(e)}"
            if progress_callback:
                progress_callback(f"❌ {error_msg}")
            st.warning(error_msg)
            continue
    
    result = " ".join(combined_text)
    if progress_callback:
        progress_callback(f"✓ Combined {len(result)} total characters from {len(combined_text)} document(s)")
    return result

# Cache for extracted text
@st.cache_data(ttl=3600)  # Cache for 1 hour
def get_cached_text(document_paths: Tuple[str]) -> str:
    """Get cached extracted text for documents"""
    return extract_text_from_documents(list(document_paths))

# ============================================================================
# OPENAI API SETTINGS
# ============================================================================

def get_openai_api_key() -> Optional[str]:
    """Get OpenAI API key from backend only (secrets.toml or environment variable)"""
    api_key = None
    try:
        # Try secrets.toml first (primary source)
        if hasattr(st, 'secrets') and st.secrets:
            api_key = st.secrets.get("OPENAI_API_KEY", None)
            if api_key and api_key.strip():
                # Store in session state for quick access
                st.session_state.openai_api_key = api_key.strip()
                return api_key.strip()
    except (AttributeError, KeyError, Exception) as e:
        # Secrets file might not exist or key not found - this is okay, try env var
        pass
    
    # Fallback: Try environment variable (but don't use if it has proxy settings)
    api_key = os.environ.get("OPENAI_API_KEY", None)
    if api_key and api_key.strip():
        # Check if it's a valid API key format (starts with sk-)
        if api_key.strip().startswith("sk-"):
            st.session_state.openai_api_key = api_key.strip()
            return api_key.strip()
    
    return None

def get_admin_password() -> Optional[str]:
    """Get admin password from secrets.toml or environment variables (required)"""
    # Check if already cached in session state
    if "admin_password" in st.session_state:
        return st.session_state.admin_password
    
    password = None
    try:
        # Try to get from Streamlit secrets first (primary source)
        if hasattr(st, 'secrets') and st.secrets:
            password = st.secrets.get("ADMIN_PASSWORD", None)
            if password and password.strip():
                # Store in session state for quick access
                st.session_state.admin_password = password.strip()
                return password.strip()
    except (AttributeError, KeyError, Exception):
        # Secrets file might not exist or key not found - try env var
        pass
    
    # Fallback: Try environment variable
    password = os.environ.get("ADMIN_PASSWORD", None)
    if password and password.strip():
        st.session_state.admin_password = password.strip()
        return password.strip()
    
    # No default password - must be set in secrets.toml or environment variable
    return None

def save_openai_settings(api_key: Optional[str], model: str, temperature: float):
    """Save OpenAI settings to session state (API key is backend-only, not saved from UI)"""
    # Only save model and temperature from UI
    # API key is always read from backend (secrets/env)
    if api_key:
        # Store in session state for current session (read from backend)
        st.session_state.openai_api_key = api_key
    st.session_state.openai_model = model
    st.session_state.openai_temperature = temperature

def get_openai_model() -> str:
    """Get selected OpenAI model (default: gpt-4)"""
    return st.session_state.get("openai_model", "gpt-4")

def get_openai_temperature() -> float:
    """Get OpenAI temperature setting (default: 0.3 for consistent exams)"""
    return st.session_state.get("openai_temperature", 0.3)

# ============================================================================
# EXPORT FUNCTIONS
# ============================================================================

def export_to_pdf(questions: List[Dict], exam_title: str = "Criminology Practice Exam") -> Optional[bytes]:
    """Export questions to PDF format - silently returns None if unavailable"""
    if not REPORTLAB_AVAILABLE:
        return None
    
    if not questions or len(questions) == 0:
        return None
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from io import BytesIO
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#1a2332',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Title
        story.append(Paragraph(exam_title, title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Instructions: Select the best answer for each question.", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Questions
        for i, q in enumerate(questions, 1):
            story.append(Paragraph(f"<b>{i}. {q.get('question', 'N/A')}</b>", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            
            # Options
            options = q.get('options', [])
            if options:
                for j, opt in enumerate(options, 1):
                    letter_opt = chr(64 + j)  # A, B, C, D
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{letter_opt}. {opt}", styles['Normal']))
            
            story.append(Spacer(1, 0.15*inch))
        
        # Answer key
        story.append(PageBreak())
        story.append(Paragraph("<b>ANSWER KEY</b>", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        for i, q in enumerate(questions, 1):
            correct = q.get('correct_answer', 'N/A')
            explanation = q.get('explanation', '')
            story.append(Paragraph(f"<b>{i}.</b> {correct}", styles['Normal']))
            if explanation:
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;<i>Explanation: {explanation}</i>", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        return None
    except Exception as e:
        return None

def export_to_docx(questions: List[Dict], exam_title: str = "Criminology Practice Exam") -> Optional[bytes]:
    """Export questions to DOCX format - silently returns None if unavailable"""
    if not DOCX_AVAILABLE or Document_class is None:
        return None
    
    if not questions or len(questions) == 0:
        return None
    
    try:
        from io import BytesIO
        doc = Document_class()
        
        # Title
        title = doc.add_heading(exam_title, 0)
        doc.add_paragraph("Instructions: Select the best answer for each question.")
        doc.add_paragraph()
        
        # Questions
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q.get('question', 'N/A')}", style='List Number')
            options = q.get('options', [])
            if options:
                for j, opt in enumerate(options, 1):
                    letter_opt = chr(64 + j)  # A, B, C, D
                    doc.add_paragraph(f"   {letter_opt}. {opt}", style='List Bullet')
            doc.add_paragraph()
        
        # Answer key
        doc.add_page_break()
        doc.add_heading("ANSWER KEY", level=1)
        doc.add_paragraph()
        
        for i, q in enumerate(questions, 1):
            correct = q.get('correct_answer', 'N/A')
            explanation = q.get('explanation', '')
            doc.add_paragraph(f"{i}. {correct}", style='List Number')
            if explanation:
                doc.add_paragraph(f"   Explanation: {explanation}")
            doc.add_paragraph()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        return None

# ============================================================================
# PREMIUM CODE MANAGEMENT
# ============================================================================

def generate_premium_code(length: int = 12) -> str:
    """Generate a random premium code"""
    alphabet = string.ascii_uppercase + string.digits
    return ''.join(secrets.choice(alphabet) for _ in range(length))

def create_premium_codes(num_codes: int, length: int, expiry_days: Optional[int], max_uses: int, created_by: str, access_level: str = "Premium") -> List[str]:
    """Create premium codes in database with specified access level"""
    codes = []
    table_name = get_table_name("premium_codes")
    
    for _ in range(num_codes):
        code = generate_premium_code(length)
        expiry = None
        if expiry_days:
            expiry = (datetime.now() + timedelta(days=expiry_days)).isoformat()
        
        # Check if access_level column exists
        columns = get_table_columns("premium_codes")
        has_access_level = 'access_level' in columns
        
        if has_access_level:
            cursor = execute_query(f"""
                INSERT INTO {table_name} (code, status, expiry_date, max_uses, uses_left, created_at, created_by, access_level)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (code, "active", expiry, max_uses, max_uses, datetime.now().isoformat(), created_by, access_level))
        else:
            # Fallback for databases without access_level column
            cursor = execute_query(f"""
                INSERT INTO {table_name} (code, status, expiry_date, max_uses, uses_left, created_at, created_by)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (code, "active", expiry, max_uses, max_uses, datetime.now().isoformat(), created_by))
        
        codes.append(code)
        cursor.close()
    
    if is_snowflake():
        db_conn.commit()
    return codes

def validate_premium_code(code: str) -> Tuple[bool, str, Optional[str]]:
    """Validate a premium code and return access level"""
    table_name = get_table_name("premium_codes")
    
    # Check if access_level column exists
    columns = get_table_columns("premium_codes")
    has_access_level = 'access_level' in columns
    
    if has_access_level:
        cursor = execute_query(f"""
            SELECT status, expiry_date, uses_left, max_uses, access_level
            FROM {table_name}
            WHERE code = %s
        """, (code,))
        
        result = cursor.fetchone()
        cursor.close()
        if not result:
            return False, "Code not found", None
        
        status, expiry_date, uses_left, max_uses, access_level = result
    else:
        cursor = execute_query(f"""
            SELECT status, expiry_date, uses_left, max_uses
            FROM {table_name}
            WHERE code = %s
        """, (code,))
        
        result = cursor.fetchone()
        cursor.close()
        if not result:
            return False, "Code not found", None
        
        status, expiry_date, uses_left, max_uses = result
        access_level = "Premium"  # Default for old codes
    
    if status != "active":
        return False, "Code is inactive", None
    
    if expiry_date:
        expiry = datetime.fromisoformat(expiry_date)
        if datetime.now() > expiry:
            return False, "Code has expired", None
    
    if uses_left <= 0:
        return False, "Code has no uses left", None
    
    return True, "Valid", access_level

def use_premium_code(code: str, session_id: str):
    """Mark a premium code as used"""
    premium_table = get_table_name("premium_codes")
    usage_table = get_table_name("code_usage")
    
    # Decrement uses_left
    cursor = execute_query(f"""
        UPDATE {premium_table}
        SET uses_left = uses_left - 1
        WHERE code = %s
    """, (code,))
    cursor.close()
    
    # Record usage
    cursor = execute_query(f"""
        INSERT INTO {usage_table} (code, used_at, session_id)
        VALUES (%s, %s, %s)
    """, (code, datetime.now().isoformat(), session_id))
    cursor.close()
    
    if is_snowflake():
        db_conn.commit()
    
    # Note: Cache will expire naturally (10s TTL), or can be manually cleared if needed
    # Streamlit cache_data doesn't support per-parameter clearing, so we rely on short TTL

# ============================================================================
# USER AUTHENTICATION & MANAGEMENT
# ============================================================================

def get_client_ip() -> Optional[str]:
    """Get client IP address from request context (Streamlit). Returns None if unavailable (e.g. localhost)."""
    try:
        if hasattr(st, "context") and hasattr(st.context, "ip_address"):
            ip = st.context.ip_address
            return ip if ip else None
    except Exception:
        pass
    return None

def login_user(email: str) -> Tuple[bool, Dict]:
    """Login user by email, create if doesn't exist. Tracks last login and client IP."""
    table_name = get_table_name("users")
    email = email.strip().lower()
    client_ip = get_client_ip()
    
    # Check if user exists
    cursor = execute_query(f"""
        SELECT email, access_level, questions_answered, premium_code_used, is_admin
        FROM {table_name}
        WHERE email = %s
    """, (email,))
    
    result = cursor.fetchone()
    cursor.close()
    
    if result:
        # User exists, update last login and last IP
        user_cols = get_table_columns("users")
        has_last_ip = user_cols and "last_ip" in [c.lower() for c in user_cols]
        if has_last_ip and client_ip:
            cursor = execute_query(f"""
                UPDATE {table_name}
                SET last_login = %s, last_ip = %s
                WHERE email = %s
            """, (datetime.now().isoformat(), client_ip, email))
        else:
            cursor = execute_query(f"""
                UPDATE {table_name}
                SET last_login = %s
                WHERE email = %s
            """, (datetime.now().isoformat(), email))
        cursor.close()
        
        if is_snowflake():
            db_conn.commit()
        
        return True, {
            "email": result[0],
            "access_level": result[1],
            "questions_answered": result[2],
            "premium_code_used": result[3],
            "is_admin": bool(result[4])
        }
    else:
        # Create new user with Free access
        user_cols = get_table_columns("users")
        has_last_ip = user_cols and "last_ip" in [c.lower() for c in user_cols]
        if has_last_ip:
            cursor = execute_query(f"""
                INSERT INTO {table_name} (email, access_level, questions_answered, created_at, last_login, is_admin, last_ip)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (email, "Free", 0, datetime.now().isoformat(), datetime.now().isoformat(), 0, client_ip or ""))
        else:
            cursor = execute_query(f"""
                INSERT INTO {table_name} (email, access_level, questions_answered, created_at, last_login, is_admin)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (email, "Free", 0, datetime.now().isoformat(), datetime.now().isoformat(), 0))
        cursor.close()
        
        if is_snowflake():
            db_conn.commit()
        
        return True, {
            "email": email,
            "access_level": "Free",
            "questions_answered": 0,
            "premium_code_used": None,
            "is_admin": False
        }

def update_user_access_level(email: str, access_level: str):
    """Update user access level (Free or Premium)"""
    table_name = get_table_name("users")
    cursor = execute_query(f"""
        UPDATE {table_name}
        SET access_level = %s
        WHERE email = %s
    """, (access_level, email))
    cursor.close()
    
    if is_snowflake():
        db_conn.commit()

def update_user_questions_answered(email: str, count: int):
    """Update user's questions answered count"""
    table_name = get_table_name("users")
    cursor = execute_query(f"""
        UPDATE {table_name}
        SET questions_answered = questions_answered + %s
        WHERE email = %s
    """, (count, email))
    cursor.close()
    
    if is_snowflake():
        db_conn.commit()
    
    # Also update session state
    cursor = execute_query(f"""
        SELECT questions_answered FROM {table_name} WHERE email = %s
    """, (email,))
    result = cursor.fetchone()
    cursor.close()
    if result:
        st.session_state.questions_answered = result[0]

@st.cache_data(ttl=60)  # Cache for 1 minute
def get_user_info(email: str) -> Optional[Dict]:
    """Get user information (cached)"""
    table_name = get_table_name("users")
    cursor = execute_query(f"""
        SELECT email, access_level, questions_answered, premium_code_used, is_admin
        FROM {table_name}
        WHERE email = %s
    """, (email,))
    # execute_query should normally return a valid cursor, but guard against None
    if cursor is None:
        return None

    result = cursor.fetchone()
    cursor.close()
    if result:
        return {
            "email": result[0],
            "access_level": result[1],
            "questions_answered": result[2],
            "premium_code_used": result[3],
            "is_admin": bool(result[4])
        }
    return None

# ============================================================================
# PDF RESOURCE MANAGEMENT
# ============================================================================

def save_pdf_resource(filename: str, filepath: str, is_premium_only: bool, use_for_ai: bool, uploaded_by: str, description: str = "", is_downloadable: bool = True):
    """Save PDF resource to database"""
    table_name = get_table_name("pdf_resources")
    # Check if is_downloadable column exists
    columns = get_table_columns("pdf_resources")
    has_downloadable = 'is_downloadable' in columns
    
    if has_downloadable:
        cursor = execute_query(f"""
            INSERT INTO {table_name} (filename, filepath, is_premium_only, use_for_ai_generation, uploaded_at, uploaded_by, description, is_downloadable)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (filename, filepath, 1 if is_premium_only else 0, 1 if use_for_ai else 0, 
              datetime.now().isoformat(), uploaded_by, description, 1 if is_downloadable else 0))
    else:
        cursor = execute_query(f"""
            INSERT INTO {table_name} (filename, filepath, is_premium_only, use_for_ai_generation, uploaded_at, uploaded_by, description)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (filename, filepath, 1 if is_premium_only else 0, 1 if use_for_ai else 0, 
              datetime.now().isoformat(), uploaded_by, description))
    cursor.close()
    
    if is_snowflake():
        db_conn.commit()

@st.cache_data(ttl=300)  # Cache for 5 minutes
def get_pdf_resources(premium_only: Optional[bool] = None) -> List[Dict]:
    """Get PDF resources, optionally filtered by premium status (cached)"""
    table_name = get_table_name("pdf_resources")
    # Check if is_downloadable column exists
    columns = get_table_columns("pdf_resources")
    has_downloadable = 'is_downloadable' in columns
    
    if premium_only is not None:
        if has_downloadable:
            cursor = execute_query(f"""
                SELECT id, filename, filepath, is_premium_only, use_for_ai_generation, uploaded_at, description, is_downloadable
                FROM {table_name}
                WHERE is_premium_only = %s
                ORDER BY uploaded_at DESC
            """, (1 if premium_only else 0,))
        else:
            cursor = execute_query(f"""
                SELECT id, filename, filepath, is_premium_only, use_for_ai_generation, uploaded_at, description
                FROM {table_name}
                WHERE is_premium_only = %s
                ORDER BY uploaded_at DESC
            """, (1 if premium_only else 0,))
    else:
        if has_downloadable:
            cursor = execute_query(f"""
                SELECT id, filename, filepath, is_premium_only, use_for_ai_generation, uploaded_at, description, is_downloadable
                FROM {table_name}
                ORDER BY uploaded_at DESC
            """)
        else:
            cursor = execute_query(f"""
                SELECT id, filename, filepath, is_premium_only, use_for_ai_generation, uploaded_at, description
                FROM {table_name}
                ORDER BY uploaded_at DESC
            """)
    
    results = cursor.fetchall()
    cursor.close()
    
    return [
        {
            "id": r[0],
            "filename": r[1],
            "filepath": r[2],
            "is_premium_only": bool(r[3]),
            "use_for_ai_generation": bool(r[4]),
            "uploaded_at": r[5],
            "description": r[6],
            "is_downloadable": bool(r[7]) if has_downloadable and len(r) > 7 else True  # Default to True if column doesn't exist
        }
        for r in results
    ]

def update_pdf_resource(resource_id: int, is_premium_only: Optional[bool] = None, use_for_ai: Optional[bool] = None):
    """Update PDF resource settings"""
    table_name = get_table_name("pdf_resources")
    updates = []
    params = []
    
    if is_premium_only is not None:
        updates.append("is_premium_only = %s")
        params.append(1 if is_premium_only else 0)
    
    if use_for_ai is not None:
        updates.append("use_for_ai_generation = %s")
        params.append(1 if use_for_ai else 0)
    
    if updates:
        params.append(resource_id)
        cursor = execute_query(f"""
            UPDATE {table_name}
            SET {', '.join(updates)}
            WHERE id = %s
        """, tuple(params))
        cursor.close()
        
        if is_snowflake():
            db_conn.commit()

def delete_pdf_resource(resource_id: int):
    """Delete PDF resource"""
    table_name = get_table_name("pdf_resources")
    cursor = execute_query(f"DELETE FROM {table_name} WHERE id = %s", (resource_id,))
    cursor.close()
    
    if is_snowflake():
        db_conn.commit()

# ============================================================================
# PAYMENT RECEIPT MANAGEMENT
# ============================================================================

def save_payment_receipt(name: str, email: str, reference: str, filename: str):
    """Save payment receipt to database"""
    table_name = get_table_name("payment_receipts")
    cursor = execute_query(f"""
        INSERT INTO {table_name} (full_name, email, gcash_reference, receipt_filename, status, submitted_at)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (name, email, reference, filename, "Pending", datetime.now().isoformat()))
    cursor.close()
    
    if is_snowflake():
        db_conn.commit()

# ============================================================================
# UI COMPONENTS - PNP THEME
# ============================================================================

def inject_pnp_theme_css():
    """Inject PNP/tactical theme CSS (cached per session)"""
    # Use session state to cache CSS injection
    if st.session_state.get("_css_injected", False):
        return
    
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=Poppins:wght@400;500;600;700;800&display=swap');
    
    * {
        font-family: 'Inter', 'Poppins', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Main background - Navy to Charcoal gradient */
    .stApp {
        background: linear-gradient(180deg, #1a2332 0%, #0f1419 100%);
        color: #e0e0e0;
    }
    
    /* Header banner */
    .header-banner {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d4a6b 50%, #1a2332 100%);
        padding: 1.5rem 2rem;
        border-radius: 0;
        border-bottom: 3px solid #d4af37;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.5);
        margin-bottom: 2rem;
    }
    
    .header-banner h1 {
        color: #ffffff;
        font-weight: 800;
        font-size: 2rem;
        text-transform: uppercase;
        letter-spacing: 2px;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.5);
    }
    
    .header-banner p {
        color: #d4af37;
        font-weight: 600;
        margin: 0.5rem 0 0 0;
        font-size: 1rem;
    }
    
    .header-badge {
        display: inline-block;
        background: rgba(212, 175, 55, 0.2);
        border: 2px solid #d4af37;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        color: #d4af37;
        font-weight: 700;
        font-size: 0.85rem;
        margin-left: 1rem;
    }
    
    /* Sidebar - Gunmetal panel */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1a2332 0%, #0f1419 100%);
        border-right: 2px solid #2d4a6b;
    }
    
    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] {
        color: #e0e0e0;
    }
    
    /* Cards/Panels */
    .pnp-card {
        background: rgba(30, 58, 95, 0.6);
        backdrop-filter: blur(10px);
        border: 1px solid #d4af37;
        border-radius: 16px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.4);
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }
    
    .pnp-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 12px 40px rgba(212, 175, 55, 0.3);
    }
    
    .pnp-card h3 {
        color: #d4af37;
        font-weight: 700;
        text-transform: uppercase;
        font-size: 1.2rem;
        margin-top: 0;
        border-bottom: 2px solid #d4af37;
        padding-bottom: 0.5rem;
    }
    
    /* Badges */
    .badge-easy {
        background: linear-gradient(135deg, #28a745 0%, #20c997 100%);
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 700;
        display: inline-block;
        box-shadow: 0 4px 15px rgba(40, 167, 69, 0.4);
    }
    
    .badge-average {
        background: linear-gradient(135deg, #ffc107 0%, #ff9800 100%);
        color: #1a2332;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 700;
        display: inline-block;
        box-shadow: 0 4px 15px rgba(255, 193, 7, 0.4);
    }
    
    .badge-difficult {
        background: linear-gradient(135deg, #dc3545 0%, #c82333 100%);
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 700;
        display: inline-block;
        box-shadow: 0 4px 15px rgba(220, 53, 69, 0.4);
    }
    
    .badge-premium {
        background: linear-gradient(135deg, #d4af37 0%, #f4d03f 100%);
        color: #1a2332;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: 700;
        display: inline-block;
        box-shadow: 0 4px 15px rgba(212, 175, 55, 0.4);
    }
    
    /* Buttons - Navy with gold hover */
    .stButton>button {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d4a6b 100%);
        color: #ffffff;
        border: 2px solid #d4af37;
        border-radius: 8px;
        font-weight: 600;
        padding: 0.75rem 1.5rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #2d4a6b 0%, #3d5a7b 100%);
        border-color: #f4d03f;
        box-shadow: 0 6px 20px rgba(212, 175, 55, 0.5);
        transform: translateY(-2px);
    }
    
    /* Alert panels */
    .alert-restricted {
        background: linear-gradient(135deg, #dc3545 0%, #c82333 100%);
        border: 3px solid #ff6b6b;
        border-radius: 12px;
        padding: 2rem;
        color: white;
        text-align: center;
        box-shadow: 0 8px 32px rgba(220, 53, 69, 0.5);
    }
    
    .alert-premium {
        background: linear-gradient(135deg, #d4af37 0%, #f4d03f 100%);
        border: 3px solid #d4af37;
        border-radius: 12px;
        padding: 2rem;
        color: #1a2332;
        text-align: center;
        box-shadow: 0 8px 32px rgba(212, 175, 55, 0.5);
    }
    
    /* Progress bar */
    .progress-container {
        background: #1a2332;
        border-radius: 10px;
        height: 20px;
        overflow: hidden;
        border: 1px solid #2d4a6b;
        margin: 1rem 0;
    }
    
    .progress-bar {
        height: 100%;
        background: linear-gradient(90deg, #d4af37 0%, #f4d03f 100%);
        border-radius: 10px;
        transition: width 0.6s ease;
        box-shadow: 0 0 10px rgba(212, 175, 55, 0.6);
    }
    
    /* System status box */
    .system-status {
        background: rgba(30, 58, 95, 0.8);
        border: 1px solid #2d4a6b;
        border-radius: 12px;
        padding: 1rem;
        margin-top: 2rem;
    }
    
    .system-status h4 {
        color: #d4af37;
        font-weight: 700;
        text-transform: uppercase;
        margin-top: 0;
        font-size: 0.9rem;
        border-bottom: 1px solid #2d4a6b;
        padding-bottom: 0.5rem;
    }
    
    .system-status p {
        color: #e0e0e0;
        margin: 0.5rem 0;
        font-size: 0.85rem;
    }
    
    /* Text colors */
    .stMarkdown {
        color: #e0e0e0;
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background: rgba(30, 58, 95, 0.6);
        color: #d4af37;
        font-weight: 600;
    }
    
    /* Mobile Responsive Design */
    @media (max-width: 768px) {
        .header-banner h1 {
            font-size: 1.3rem;
            letter-spacing: 1px;
        }
        
        .header-banner {
            padding: 1rem 1rem;
        }
        
        .pnp-card {
            padding: 1rem;
            margin: 0.5rem 0;
        }
        
        .stButton>button {
            padding: 0.5rem 1rem;
            font-size: 0.9rem;
        }
        
        /* Stack columns on mobile */
        [data-testid="column"] {
            width: 100% !important;
            margin-bottom: 1rem;
        }
    }
    
    /* High contrast text for dark theme */
    .stMarkdown, .stText, .stSelectbox label, .stRadio label {
        color: #e0e0e0 !important;
    }
    
    /* Clear CTA buttons */
    .cta-button {
        background: linear-gradient(135deg, #d4af37 0%, #f4d03f 100%);
        color: #1a2332;
        font-weight: 700;
        padding: 1rem 2rem;
        border-radius: 8px;
        border: none;
        box-shadow: 0 4px 15px rgba(212, 175, 55, 0.5);
        transition: all 0.3s ease;
    }
    
    .cta-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(212, 175, 55, 0.7);
    }
    
    /* Progress feedback container */
    .progress-feedback {
        background: rgba(30, 58, 95, 0.8);
        border: 1px solid #d4af37;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    .progress-step {
        display: flex;
        align-items: center;
        margin: 0.5rem 0;
        color: #e0e0e0;
    }
    
    .progress-step.active {
        color: #d4af37;
        font-weight: 600;
    }
    
    .progress-step.completed {
        color: #28a745;
    }
    
    /* Document library card */
    .doc-card {
        background: rgba(30, 58, 95, 0.6);
        border: 1px solid #2d4a6b;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .doc-card.selected {
        border-color: #d4af37;
        background: rgba(212, 175, 55, 0.1);
    }
    </style>
    """, unsafe_allow_html=True)

def render_header():
    """Render PNP-style header banner"""
    if st.session_state.user_logged_in:
        mode = st.session_state.user_access_level.upper()
    else:
        mode = "PREMIUM" if st.session_state.premium_active else "FREE"
    badge_html = f'<span class="header-badge">MODE: {mode}</span>'
    
    st.markdown(f"""
    <div class="header-banner">
        <h1>🛡️ PH CRIMINOLOGY EXAM REVIEWER {badge_html}</h1>
    </div>
    """, unsafe_allow_html=True)

def render_premium_lock(feature_name: str, upgrade_message: str = "Upgrade to Premium to unlock this feature"):
    """Render a premium lock indicator"""
    if st.session_state.user_access_level != "Premium":
        st.markdown(f"""
        <div style="background: rgba(220, 53, 69, 0.2); border: 2px solid #dc3545; border-radius: 12px; padding: 1.5rem; text-align: center; margin: 1rem 0;">
            <h3 style="color: #dc3545; margin: 0 0 0.5rem 0;">🔒 {feature_name}</h3>
            <p style="color: #e0e0e0; margin: 0;">{upgrade_message}</p>
            <p style="color: #d4af37; margin: 0.5rem 0 0 0; font-weight: 600;">Go to Premium Access or Payment page to upgrade</p>
        </div>
        """, unsafe_allow_html=True)
        return True
    return False

def render_card(title: str, content_html: str):
    """Render a PNP-style card"""
    st.markdown(f"""
    <div class="pnp-card">
        <h3>{title}</h3>
        <div>{content_html}</div>
    </div>
    """, unsafe_allow_html=True)

def render_badge(text: str, color: str):
    """Render a difficulty badge"""
    badge_class = f"badge-{color.lower()}"
    return f'<span class="{badge_class}">{text}</span>'

# Helper function for admin panel - get all users (cached)
@st.cache_data(ttl=120)  # Cache for 2 minutes
@st.cache_data(ttl=300)
def get_all_users_cached():
    """Get all users for admin panel (cached). Includes last_ip when column exists."""
    table_name = get_table_name("users")
    user_cols = get_table_columns("users")
    has_last_ip = user_cols and "last_ip" in [c.lower() for c in user_cols]
    if has_last_ip:
        cursor = execute_query(f"""
            SELECT email, access_level, questions_answered, created_at, last_login, premium_code_used, is_admin, last_ip
            FROM {table_name}
            ORDER BY created_at DESC
        """)
    else:
        cursor = execute_query(f"""
            SELECT email, access_level, questions_answered, created_at, last_login, premium_code_used, is_admin
            FROM {table_name}
            ORDER BY created_at DESC
        """)
    results = cursor.fetchall()
    cursor.close()
    out = []
    for r in results:
        row = {
            "email": r[0],
            "access_level": r[1],
            "questions_answered": r[2],
            "created_at": r[3],
            "last_login": r[4],
            "premium_code_used": r[5] if len(r) > 5 else None,
            "is_admin": bool(r[6]) if len(r) > 6 else False
        }
        if has_last_ip and len(r) > 7:
            row["last_ip"] = r[7] or ""
        else:
            row["last_ip"] = ""
        out.append(row)
    return out

# Inject theme
inject_pnp_theme_css()
render_header()

# ============================================================================
# SIDEBAR NAVIGATION (page selector only when not logged in; after login use main-page cards)
# ============================================================================

with st.sidebar:
    st.markdown("### 🧭 NAVIGATION")
    
    if st.session_state.user_logged_in:
        page = st.session_state.get("current_page", "🏠 Home")
        st.caption("Use the **cards on the main page** to switch sections.")
    else:
        page = st.radio(
            "Select Page",
            ["🏠 Home", "📄 Upload Reviewer", "🧠 Practice Exam", "💳 Payment", "🛠️ Admin Panel"],
            label_visibility="collapsed"
        )
    
    st.markdown("---")
    
    # User Status
    if st.session_state.user_logged_in:
        st.markdown("""
        <div class="system-status">
            <h4>USER STATUS</h4>
            <p>👤 Email: {}</p>
            <p>🔑 Access: {}</p>
            <p>📊 Questions: {}/{}</p>
            <p>📄 PDF Loaded: {}</p>
        </div>
        """.format(
            st.session_state.user_email[:20] + "..." if len(st.session_state.user_email) > 20 else st.session_state.user_email,
            st.session_state.user_access_level,
            st.session_state.questions_answered,
            999999 if st.session_state.user_access_level == "Premium" else (ADVANCE_QUESTION_LIMIT if st.session_state.user_access_level == "Advance" else FREE_QUESTION_LIMIT),
            "Yes" if st.session_state.pdf_text else "No"
        ), unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="system-status">
            <h4>SYSTEM STATUS</h4>
            <p>📄 PDFs Loaded: {}</p>
            <p>❓ Questions Generated: {}</p>
            <p>📊 Remaining: {}/{}</p>
            <p>🔑 Premium: {}</p>
        </div>
        """.format(
            "Yes" if st.session_state.pdf_text else "No",
            len(st.session_state.current_questions),
            st.session_state.questions_answered,
            999999 if st.session_state.premium_active else (ADVANCE_QUESTION_LIMIT if st.session_state.user_access_level == "Advance" else FREE_QUESTION_LIMIT),
            "ON" if st.session_state.premium_active else "OFF"
        ), unsafe_allow_html=True)
    
    st.markdown("---")
    st.caption("For review purposes only. Verify with latest PH laws.")

# ============================================================================
# CARD-STYLE NAVIGATION (main page, shown after login)
# ============================================================================

if st.session_state.user_logged_in:
    nav_pages = [
        ("🏠", "Home", "🏠 Home"),
        ("📄", "Upload Reviewer", "📄 Upload Reviewer"),
        ("🧠", "Practice Exam", "🧠 Practice Exam"),
        ("💳", "Payment", "💳 Payment"),
    ]
    if st.session_state.get("admin_logged_in"):
        nav_pages.append(("🛠️", "Admin Panel", "🛠️ Admin Panel"))
    
    cols = st.columns(len(nav_pages))
    for i, (icon, label, page_value) in enumerate(nav_pages):
        with cols[i]:
            is_current = (page == page_value)
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, rgba(0, 51, 102, 0.25) 0%, rgba(0, 51, 102, 0.15) 100%);
                border: 2px solid {'#d4af37' if is_current else 'rgba(212, 175, 55, 0.5)'};
                border-radius: 12px;
                padding: 1rem;
                text-align: center;
                margin-bottom: 0.5rem;
                box-shadow: 0 2px 8px rgba(0,0,0,0.2);
            ">
                <div style="font-size: 1.8rem; margin-bottom: 0.25rem;">{icon}</div>
                <div style="font-size: 0.85rem; font-weight: 600; color: #e0e0e0;">{label}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Go to {label}", key=f"nav_{page_value}", use_container_width=True):
                st.session_state.current_page = page_value
                st.rerun()
    st.markdown("---")

# ============================================================================
# PAGE: HOME
# ============================================================================

if page == "🏠 Home":
    st.markdown("# 🏠 Welcome to PH Criminology Exam Reviewer")
    
    # User Login Section
    if not st.session_state.user_logged_in:
        render_card("🔐 Login Required", """
        <p>Please login with your email address to access the reviewer.</p>
        <p><strong>New users:</strong> Enter your email to automatically create a Free account.</p>
        """)
        
        with st.form("login_form"):
            email_input = st.text_input("Email Address", placeholder="your.email@example.com", help="Enter your email to login or create an account")
            login_submitted = st.form_submit_button("🔓 Login / Sign Up", type="primary", use_container_width=True)
            
            if login_submitted:
                if email_input and "@" in email_input and "." in email_input.split("@")[1]:
                    success, user_info = login_user(email_input)
                    if success:
                        st.session_state.user_logged_in = True
                        st.session_state.user_email = user_info["email"]
                        st.session_state.user_access_level = user_info["access_level"]
                        st.session_state.questions_answered = user_info["questions_answered"]
                        st.session_state.premium_active = (user_info["access_level"] == "Premium")
                        st.session_state.premium_code_used = user_info.get("premium_code_used")
                        if user_info.get("is_admin"):
                            st.session_state.admin_logged_in = True
                        st.session_state.show_credential_warning = True  # Show warning on next page load
                        
                        st.success(f"✅ Welcome, {user_info['email']}! Access Level: {user_info['access_level']}")
                        st.rerun()
                    else:
                        st.error("❌ Login failed. Please try again.")
                else:
                    st.error("❌ Please enter a valid email address.")
        
        st.stop()  # Stop here if not logged in
    
    # User is logged in - show dashboard
    user_info = get_user_info(st.session_state.user_email)
    if user_info:
        st.session_state.user_access_level = user_info["access_level"]
        st.session_state.questions_answered = user_info["questions_answered"]
        st.session_state.premium_active = (user_info["access_level"] == "Premium")
    
    # Show credential-sharing warning once after login
    if st.session_state.get("show_credential_warning"):
        st.warning("""
        **⚠️ Keep your account secure**  
        **Do not share your email or login credentials with anyone.**  
        Access is tied to your account. Sharing credentials can lead to your access being blocked or restricted.  
        Use only your own account from your own device.
        """)
        st.session_state.show_credential_warning = False
    
    # User profile card
    if st.session_state.user_access_level == "Premium":
        access_badge = "👑 PREMIUM"
    elif st.session_state.user_access_level == "Advance":
        access_badge = "⚡ ADVANCE"
    else:
        access_badge = "🆓 FREE"
    render_card(f"👤 User Profile - {access_badge}", f"""
    <p><strong>Email:</strong> {st.session_state.user_email}</p>
    <p><strong>Access Level:</strong> {st.session_state.user_access_level}</p>
    <p><strong>Questions Answered:</strong> {st.session_state.questions_answered}</p>
    """)
    
    if st.button("🚪 Logout", type="secondary"):
        st.session_state.user_logged_in = False
        st.session_state.user_email = None
        st.session_state.user_access_level = "Free"
        st.session_state.admin_logged_in = False
        st.rerun()
    
    st.markdown("---")
    
    render_card("📋 About This App", """
    <p>This application helps you prepare for the Philippine Criminology Exam by:</p>
    <ul>
        <li>📄 Uploading or selecting reviewer PDFs</li>
        <li>🧠 Generating practice questions automatically</li>
        <li>📊 Tracking your progress and scores</li>
        <li>🔑 Unlocking premium features with codes</li>
    </ul>
    """)
    
    render_card("🎯 Getting Started", """
    <ol>
        <li><strong>Upload/Select Document:</strong> Go to "Upload Reviewer" page and select or upload your review materials</li>
        <li><strong>Generate Questions:</strong> Choose difficulty level and question types</li>
        <li><strong>Practice:</strong> Answer questions and track your progress</li>
        <li><strong>Upgrade:</strong> Unlock 100 questions with premium code or payment</li>
    </ol>
    """)
    
    render_card("⚠️ Important Notice", """
    <p><strong>For review purposes only.</strong> Always verify with official references and latest PH laws.</p>
    <p>This app generates questions based on uploaded PDF content. Always cross-reference with official materials.</p>
    """)
    
    # Quick stats
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Questions Answered", st.session_state.questions_answered)
    with col2:
        st.metric("Current Mode", st.session_state.user_access_level.upper())
    with col3:
        # Revised tiering system
        if st.session_state.user_access_level == "Premium":
            max_q = 999999  # Unlimited
        elif st.session_state.user_access_level == "Advance":
            max_q = ADVANCE_QUESTION_LIMIT
        else:
            max_q = FREE_QUESTION_LIMIT
        remaining = max_q - st.session_state.questions_answered
        st.metric("Remaining", max(0, remaining))

# ============================================================================
# PAGE: UPLOAD REVIEWER
# ============================================================================

elif page == "📄 Upload Reviewer":
    st.markdown("# 📄 Document Library & Upload")
    
    # Check if user is logged in
    if not st.session_state.user_logged_in:
        st.warning("⚠️ Please login first on the Home page to access this feature.")
        st.stop()
    
    # Document Library Section
    with st.container():
        st.markdown("### 📚 Document Library")
        st.markdown("Select documents to include in exam generation. Content from selected documents will be combined.")
        if st.session_state.user_access_level in ("Advance", "Premium"):
            st.caption("📄 **Admin-uploaded documents** are visible to all Advance and Premium users (existing and new). Select the documents you want to use for question generation.")
        
        # Get documents: all admin-uploaded docs for Advance/Premium; user docs by email
        current_user_email = st.session_state.user_email if st.session_state.user_logged_in else None
        current_access_level = st.session_state.user_access_level if st.session_state.user_logged_in else None
        # Clear cache for Advance/Premium so admin docs from DB are always fresh when opening this page
        if current_access_level in ("Advance", "Premium"):
            get_document_library.clear()
        all_documents = get_document_library(user_email=current_user_email, user_access_level=current_access_level)
        
        if not all_documents:
            if st.session_state.user_access_level == "Free":
                st.info("📭 **Document Library** is available for **Advance** and **Premium** users. Upgrade to access admin-uploaded documents and upload your own.")
            else:
                st.info("📭 No documents found. Upload documents below or ask admin to add documents.")
        else:
            # Initialize selected documents in session state
            if "document_selections" not in st.session_state:
                st.session_state.document_selections = {}
            
            # Ensure all current documents are in the selections dictionary
            for doc in all_documents:
                if doc['filepath'] not in st.session_state.document_selections:
                    st.session_state.document_selections[doc['filepath']] = True  # Default to selected
            
            # Remove selections for documents that no longer exist
            existing_paths = {doc['filepath'] for doc in all_documents}
            st.session_state.document_selections = {
                path: selected for path, selected in st.session_state.document_selections.items()
                if path in existing_paths
            }
            
            selected_count = sum(1 for path, selected in st.session_state.document_selections.items() if selected and path in existing_paths)
            
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{selected_count} of {len(all_documents)} documents selected**")
            with col2:
                if st.button("Select All"):
                    st.session_state.document_selections = {doc['filepath']: True for doc in all_documents}
                    st.rerun()
                if st.button("Deselect All"):
                    st.session_state.document_selections = {doc['filepath']: False for doc in all_documents}
                    st.rerun()
            
            st.markdown("---")
            
            # Display documents with checkboxes — status label must match checkbox (render status after checkbox)
            for doc in all_documents:
                doc_path = doc['filepath']
                is_selected = st.session_state.document_selections.get(doc_path, True)
                
                # Source badge
                source_colors = {
                    "Dummy": "#28a745",
                    "Admin": "#d4af37",
                    "Uploaded": "#007bff"
                }
                source_color = source_colors.get(doc['source'], "#6c757d")
                
                with st.container():
                    col1, col2, col3, col4 = st.columns([0.5, 3, 2, 1])
                    with col1:
                        # Checkbox first so we get current value; then status uses that value
                        doc_index = all_documents.index(doc)
                        doc_hash = hashlib.md5(doc_path.encode()).hexdigest()[:8]
                        unique_doc_key = f"doc_{doc_index}_{doc_hash}_{doc.get('source', 'unknown')}"
                        checked = st.checkbox("", value=is_selected, key=unique_doc_key, label_visibility="collapsed")
                        st.session_state.document_selections[doc_path] = checked
                    with col2:
                        st.markdown(f"**{doc['filename']}**")
                        st.caption(f"📅 {doc['date']} | 📦 {doc['size_mb']} MB | {doc['type']}")
                    with col3:
                        st.markdown(f'<span style="background: {source_color}; color: white; padding: 0.2rem 0.5rem; border-radius: 4px; font-size: 0.8rem;">{doc["source"]}</span>', unsafe_allow_html=True)
                    with col4:
                        # Admin docs: download only for Premium; Advance can view/use but not download
                        if doc['source'] == "Admin":
                            can_download_admin = (
                                st.session_state.user_access_level == "Premium"
                                and doc.get('is_downloadable', True)
                            )
                            if not can_download_admin:
                                if st.session_state.user_access_level == "Advance":
                                    st.markdown("👁️ View only", help="Advance users can view and use; download is Premium only")
                                else:
                                    st.markdown("🔒 Premium", help="Upgrade to Premium to download")
                            elif os.path.exists(doc_path):
                                with open(doc_path, "rb") as f:
                                    file_data = f.read()
                                    dl_key = f"dl_{unique_doc_key}"
                                    st.download_button(
                                        "📥",
                                        data=file_data,
                                        file_name=doc['filename'],
                                        mime=_mime_for_doc(doc),
                                        key=dl_key,
                                        help="Download this file"
                                    )
                        else:
                            # User-uploaded: owner can always download
                            if os.path.exists(doc_path):
                                with open(doc_path, "rb") as f:
                                    file_data = f.read()
                                    dl_key = f"dl_{unique_doc_key}"
                                    st.download_button(
                                        "📥",
                                        data=file_data,
                                        file_name=doc['filename'],
                                        mime=_mime_for_doc(doc),
                                        key=dl_key,
                                        help="Download this file"
                                    )
                
                # Status label after checkbox so it always matches: checked = Selected, unchecked = Not selected
                display_selected = st.session_state.document_selections.get(doc_path, True)
                border_color = "#28a745" if display_selected else "#6c757d"
                bg_color = "rgba(40, 167, 69, 0.12)" if display_selected else "rgba(108, 117, 125, 0.08)"
                status_text = "✓ Selected" if display_selected else "○ Not selected"
                status_color = "#28a745" if display_selected else "#6c757d"
                st.markdown(f"""
                <div style="
                    border-left: 4px solid {border_color};
                    background: {bg_color};
                    border-radius: 6px;
                    padding: 0.4rem 0.6rem;
                    margin-bottom: 0.5rem;
                ">
                    <span style="color: {status_color}; font-weight: 600; font-size: 0.85rem;">{status_text}</span>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("---")
            
            # Preview combined text from selected documents
            if selected_count > 0:
                # Mark documents as loaded even before preview (so status shows "Yes")
                if not st.session_state.pdf_text or st.session_state.pdf_text == "UPLOADED_DOCUMENT_NO_EXTRACTED_TEXT":
                    selected_doc_names = [doc['filename'] for doc in all_documents if st.session_state.document_selections.get(doc['filepath'], False)]
                    if selected_doc_names:
                        st.session_state.pdf_text = "SELECTED_DOCUMENTS_LOADED"
                        st.session_state.pdf_name = f"{selected_count} document(s) selected"
                
                if st.button("🔍 Preview Selected Documents", use_container_width=True):
                    selected_paths = [doc['filepath'] for doc in all_documents if st.session_state.document_selections.get(doc['filepath'], False)]
                    if selected_paths:
                        # Verify files exist
                        existing_paths = [p for p in selected_paths if os.path.exists(p)]
                        if not existing_paths:
                            st.error("❌ Selected documents not found on server. Please re-select documents.")
                        else:
                            progress_placeholder = st.empty()
                            
                            def update_preview_progress(msg):
                                progress_placeholder.info(f"🔄 {msg}")
                            
                            try:
                                update_preview_progress("Extracting text from selected documents...")
                                combined_text = extract_text_from_documents(
                                    existing_paths,
                                    max_pages_per_doc=None,
                                    max_total_chars=None,
                                    progress_callback=update_preview_progress
                                )
                                
                                progress_placeholder.empty()
                                
                                if combined_text and combined_text.strip():
                                    st.session_state.pdf_text = combined_text
                                    st.session_state.pdf_name = f"Combined from {selected_count} document(s)"
                                    st.success(f"✅ Loaded content from {selected_count} document(s)!")
                                    with st.expander("📖 Preview (First 1000 characters)"):
                                        st.text(combined_text[:1000] + "..." if len(combined_text) > 1000 else combined_text)
                                    st.metric("Total Characters", f"{len(combined_text):,}")

                                    # If text is very short, warn that AI generation may be limited
                                    if len(combined_text.strip()) < 50:
                                        st.warning("⚠️ Only a small amount of text was extracted. AI question generation may be limited.")
                                else:
                                    # No text extracted - try OCR explicitly on each document
                                    update_preview_progress("No text found, attempting OCR on all documents...")
                                    ocr_texts = []
                                    for doc_path in existing_paths:
                                        try:
                                            if doc_path.lower().endswith('.pdf'):
                                                update_preview_progress(f"Running OCR on {os.path.basename(doc_path)}...")
                                                ocr_text, _ = extract_text_from_pdf(doc_path)
                                                if ocr_text and ocr_text.strip():
                                                    ocr_texts.append(ocr_text)
                                                    st.info(f"✓ OCR extracted {len(ocr_text)} chars from {os.path.basename(doc_path)}")
                                            elif doc_path.lower().endswith(('.docx', '.doc')):
                                                update_preview_progress(f"Running OCR on {os.path.basename(doc_path)}...")
                                                ocr_text, _ = extract_text_from_docx_with_ocr(doc_path, os.path.basename(doc_path))
                                                if ocr_text and ocr_text.strip():
                                                    ocr_texts.append(ocr_text)
                                                    st.info(f"✓ OCR extracted {len(ocr_text)} chars from {os.path.basename(doc_path)}")
                                            elif doc_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                                                update_preview_progress(f"Reading image with OCR: {os.path.basename(doc_path)}...")
                                                ocr_text, _ = extract_text_from_image(doc_path)
                                                if ocr_text and ocr_text.strip():
                                                    ocr_texts.append(ocr_text)
                                                    st.info(f"✓ OCR extracted {len(ocr_text)} chars from {os.path.basename(doc_path)}")
                                        except Exception as doc_error:
                                            st.warning(f"⚠ OCR failed for {os.path.basename(doc_path)}: {str(doc_error)}")
                                    
                                    progress_placeholder.empty()
                                    
                                    if ocr_texts:
                                        combined_text = "\n".join(ocr_texts)
                                        st.session_state.pdf_text = combined_text
                                        st.session_state.pdf_name = f"Combined from {selected_count} document(s) (OCR)"
                                        st.success(f"✅ OCR extracted {len(combined_text)} characters from {selected_count} document(s)!")
                                        with st.expander("📖 Preview (First 1000 characters)"):
                                            st.text(combined_text[:1000] + "..." if len(combined_text) > 1000 else combined_text)
                                        st.metric("Total Characters", f"{len(combined_text):,}")
                                    else:
                                        # Even if extraction fails, mark documents as selected/loaded
                                        st.session_state.pdf_text = "SELECTED_DOCUMENTS_NO_TEXT"
                                        st.session_state.pdf_name = f"{selected_count} document(s) selected (text extraction limited)"
                                        st.error("❌ **No readable text could be extracted** from the selected documents, even with OCR.")
                                        st.warning("""
                                        **Possible reasons:**
                                        - Documents are password-protected or encrypted
                                        - Documents contain only images with no readable text
                                        - Documents are corrupted or in an unsupported format
                                        - OCR models may not be loaded (first-time setup can take 2-5 minutes)
                                        
                                        **Solutions:**
                                        - Wait a few minutes and try again (EasyOCR downloads models on first use)
                                        - Ensure documents are not password-protected
                                        - Try uploading documents with machine-readable text
                                        - Contact admin for assistance
                                        """)
                            except Exception as e:
                                progress_placeholder.empty()
                                # Even on error, mark documents as selected
                                st.session_state.pdf_text = "SELECTED_DOCUMENTS_ERROR"
                                st.session_state.pdf_name = f"{selected_count} document(s) selected"
                                st.error(f"❌ **Error processing documents:** {str(e)}")
                                import traceback
                                with st.expander("🔍 Error Details"):
                                    st.code(traceback.format_exc())
                                st.info("💡 Documents are available in your library. Please check the error details above.")
    
    st.markdown("---")
    
    # Upload New Document Section
    with st.container():
        st.markdown("### 📤 Upload New Document")
        
        # Check if user is in Free mode - disable uploads
        if st.session_state.user_access_level == "Free":
            st.markdown("""
            <div style="background: rgba(220, 53, 69, 0.2); border: 2px solid #dc3545; border-radius: 12px; padding: 1.5rem; text-align: center; margin: 1rem 0;">
                <h3 style="color: #dc3545; margin: 0 0 0.5rem 0;">🔒 Document Upload Locked</h3>
                <p style="color: #e0e0e0; margin: 0.5rem 0;">Document upload and AI-based question generation are only available in <strong>Advance</strong> or <strong>Premium</strong> mode.</p>
                <p style="color: #d4af37; margin: 0.5rem 0 0 0; font-weight: 600;">Upgrade to unlock:</p>
                <ul style="text-align: left; display: inline-block; margin: 1rem 0; color: #e0e0e0;">
                    <li>📤 Upload your own review documents</li>
                    <li>🤖 AI-powered question generation from your documents</li>
                    <li>📚 Personalized practice exams</li>
                </ul>
                <p style="color: #d4af37; margin: 1rem 0 0 0; font-weight: 600;">Go to <strong>💳 Payment</strong> page to upgrade!</p>
            </div>
            """, unsafe_allow_html=True)
            st.file_uploader("Choose document or image", type=["pdf", "docx", "doc", "png", "jpg", "jpeg"], disabled=True, help="Upgrade to Advance or Premium to unlock document upload")
            uploaded_file = None  # No file upload for Free users
        else:
            uploaded_file = st.file_uploader(
                "Choose PDF, Word, or image (PNG/JPG)",
                type=["pdf", "docx", "doc", "png", "jpg", "jpeg"],
                help="Upload reviewer PDF, Word doc, or images (text in images is read via OCR for question generation)",
                key=f"user_upload_{st.session_state.user_email or 'guest'}",
            )

            if uploaded_file:
                # Avoid re-processing the same file on rerun
                file_id = f"{uploaded_file.name}_{uploaded_file.size}"
                last_id = st.session_state.get("last_processed_upload_id")
                if last_id == file_id:
                    # Already processed this exact file in this session; skip to prevent upload loop
                    pass
                else:
                    st.session_state.last_processed_upload_id = file_id

                    # Save to uploads directory with user email in filename
                    os.makedirs(UPLOADS_DIR, exist_ok=True)

                    # Generate unique filename with timestamp and user email
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    file_ext = os.path.splitext(uploaded_file.name)[1]
                    user_email_clean = ""
                    if st.session_state.user_logged_in and st.session_state.user_email:
                        user_email_clean = (
                            st.session_state.user_email.lower()
                            .replace("@", "_")
                            .replace(".", "_")
                            + "_"
                        )

                    # Prevent duplicate uploads for the same user and original filename
                    is_duplicate = False
                    if st.session_state.user_logged_in and st.session_state.user_email:
                        try:
                            table_name = get_table_name("pdf_resources")
                            cursor = execute_query(
                                f"""
                                SELECT COUNT(*)
                                FROM {table_name}
                                WHERE uploaded_by = %s
                                  AND filename LIKE %s
                                """,
                                (
                                    st.session_state.user_email,
                                    f"%_{uploaded_file.name}",
                                ),
                            )
                            count_row = cursor.fetchone()
                            cursor.close()
                            if count_row and count_row[0] > 0:
                                is_duplicate = True
                        except Exception:
                            # If duplicate check fails, fall back to allowing upload
                            is_duplicate = False

                    if is_duplicate:
                        st.warning(
                            "⚠️ This file (same name) has already been uploaded. "
                            "Please rename the file or upload a different document."
                        )
                    else:
                        filename = f"{user_email_clean}{timestamp}_{uploaded_file.name}"
                        file_path = os.path.join(UPLOADS_DIR, filename)

                        with open(file_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())

                        # Save to database with user email
                        if st.session_state.user_logged_in and st.session_state.user_email:
                            try:
                                save_pdf_resource(
                                    filename,
                                    file_path,
                                    False,
                                    True,
                                    st.session_state.user_email,
                                    f"Uploaded by {st.session_state.user_email}",
                                )
                            except Exception:
                                pass  # Continue even if database save fails

                        # Extract text - try lightweight in-memory helper first,
                        # then fall back to the more advanced pipelines (with OCR).
                        file_ext = uploaded_file.name.lower().split(".")[-1] if uploaded_file.name else ""
                        text = ""
                        name = uploaded_file.name or "Uploaded Document"
                        
                        with st.spinner("📄 Extracting text from document... This may take a moment."):
                            # First attempt: simple extraction directly from the uploaded file object
                            try:
                                quick_text = extract_text(uploaded_file)
                            except Exception:
                                quick_text = ""

                            if quick_text and quick_text.strip():
                                text = quick_text.strip()
                                name = uploaded_file.name or "Uploaded Document"

                            # If simple extraction failed or was too short, use the existing pipelines
                            if (not text or len(text.strip()) < 50) and file_ext == "pdf":
                                # Try extracting from file object first
                                uploaded_file.seek(0)
                                text, name = extract_text_from_pdf(uploaded_file)
                                
                                # If that fails, try extracting from saved file path
                                if not text or not text.strip():
                                    try:
                                        text, name = extract_text_from_pdf(file_path)
                                    except Exception:
                                        pass
                                
                            elif file_ext in ["docx", "doc"]:
                                # Try extracting from file object first
                                uploaded_file.seek(0)
                                text, name = extract_text_from_docx(uploaded_file)
                                
                                # If that fails, try extracting from saved file path
                                if not text or not text.strip():
                                    try:
                                        text, name = extract_text_from_docx(file_path)
                                    except Exception:
                                        pass

                            elif file_ext in ["png", "jpg", "jpeg"]:
                                # Images: extract text via OCR (file path is more reliable after save)
                                with st.spinner("🔍 Reading image with OCR... This may take a moment."):
                                    try:
                                        text, name = extract_text_from_image(file_path)
                                        if not text and uploaded_file:
                                            uploaded_file.seek(0)
                                            text, name = extract_text_from_image(uploaded_file)
                                    except Exception:
                                        pass

                            # If still no text, try OCR for PDFs
                            if (not text or not text.strip()) and file_ext == "pdf" and OCR_AVAILABLE:
                                try:
                                    with st.spinner("🔍 Attempting OCR extraction (this may take longer)..."):
                                        text, name = extract_text_from_pdf(file_path)
                                except Exception:
                                    pass

                        # User-facing feedback similar to the standalone implementation
                        if file_ext == "doc":
                            st.error("This is a .DOC (old Word format). For best results, please convert it to a .DOCX and re-upload.")
                        elif not text or len(text.strip()) < 50:
                            st.warning("Text extraction failed or is too short. If this is a scanned PDF or image-heavy document, OCR may be required and results can vary.")
                        else:
                            st.success("Text extracted successfully from your document.")

                        if text and text.strip():
                            st.session_state.pdf_text = text
                            st.session_state.pdf_name = name
                        else:
                            # Even if text extraction fails, mark that the user has uploaded a document
                            # so status shows PDF Loaded = Yes and Practice Exam knows a doc exists.
                            st.session_state.pdf_text = "UPLOADED_DOCUMENT_NO_EXTRACTED_TEXT"
                            st.session_state.pdf_name = uploaded_file.name or "Uploaded Document"

                        st.success(f"✅ File uploaded and added to your document library: {st.session_state.pdf_name}")

                        # Also update document selections
                        if "document_selections" not in st.session_state:
                            st.session_state.document_selections = {}
                        st.session_state.document_selections[file_path] = True

                        # Clear cache to refresh document library
                        get_document_library.clear()

                        # Preview only when we have actual text
                        if text:
                            with st.expander("📖 Document Preview (First 1000 characters)"):
                                st.text(text[:1000] + "..." if len(text) > 1000 else text)

                            # Stats
                            word_count = len(text.split())
                            char_count = len(text)
                            col1, col2 = st.columns(2)
                            with col1:
                                st.metric("Word Count", f"{word_count:,}")
                            with col2:
                                st.metric("Character Count", f"{char_count:,}")
                        else:
                            # More graceful error handling when no text extracted
                            st.warning("⚠️ Text extraction was limited or unsuccessful for this file.")
                            st.markdown(
                                """
                                - It may be a scanned/image-based document requiring OCR
                                - It may be protected/encrypted
                                - It may contain very little machine-readable text
                                """
                            )
                            st.info(
                                "💡 The file has still been uploaded and is available in your document library. "
                                "You can use it for preview, but AI question generation may be limited."
                            )
    
    st.markdown("---")
    
    # Use Default Questions Option
    with st.container():
        st.markdown("### 🎲 Use Default Questions")
        st.markdown("Practice with pre-loaded criminology questions without uploading a document.")
        if st.button("🎯 Use Default Questions", type="secondary", use_container_width=True):
            st.session_state.pdf_text = "DEFAULT_DUMMY_QUESTIONS"
            st.session_state.pdf_name = "Default Criminology Questions"
            st.success("✅ Default questions enabled! You can now go to Practice Exam page.")
            st.info("💡 Default questions will be used automatically when generating practice exams.")

# ============================================================================
# PAGE: PRACTICE EXAM
# ============================================================================

elif page == "🧠 Practice Exam":
    st.markdown("# 🧠 Practice Exam")
    
    # Check if user is logged in
    if not st.session_state.user_logged_in:
        st.warning("⚠️ Please login first on the Home page to access Practice Exam.")
        st.stop()
    
    # Check for paused exam and offer resume (only show if not currently in a question)
    if st.session_state.get("exam_paused", False) and st.session_state.get("current_questions") and len(st.session_state.get("current_questions", [])) > 0 and st.session_state.current_question_index == 0:
        st.markdown("### ⏸️ Paused Exam Detected")
        paused_at = st.session_state.get("paused_at_index", 0)
        total_questions = len(st.session_state.current_questions)
        st.info(f"📊 You have a paused exam. You were on question {paused_at + 1} of {total_questions}.")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("▶️ Resume Exam", type="primary", use_container_width=True):
                st.session_state.current_question_index = paused_at
                st.session_state.exam_paused = False
                st.rerun()
        with col2:
            if st.button("🔄 Start New Exam", type="secondary", use_container_width=True):
                st.session_state.current_questions = []
                st.session_state.current_question_index = 0
                st.session_state.answers = {}
                st.session_state.exam_paused = False
                st.rerun()
        st.markdown("---")
        st.stop()  # Stop here until user resumes or starts new
    
    # Check if PDF is loaded or default questions enabled
    # Documents are considered "loaded" if pdf_text exists (even if it's a placeholder)
    has_actual_text = (
        st.session_state.pdf_text and 
        st.session_state.pdf_text not in ["", "DEFAULT_DUMMY_QUESTIONS"] and
        st.session_state.pdf_text not in ["SELECTED_DOCUMENTS_LOADED", "SELECTED_DOCUMENTS_NO_TEXT", "SELECTED_DOCUMENTS_ERROR", "UPLOADED_DOCUMENT_NO_EXTRACTED_TEXT"]
    )
    
    if not st.session_state.pdf_text:
        st.info("💡 No document loaded. Please upload or select documents to generate questions.")
    elif not has_actual_text:
        # Documents are selected/uploaded but text extraction was limited - still allow question generation
        if st.session_state.pdf_text.startswith("SELECTED_DOCUMENTS"):
            st.info("💡 Documents are selected. Question generation will proceed with available text.")
        elif st.session_state.pdf_text == "UPLOADED_DOCUMENT_NO_EXTRACTED_TEXT":
            st.info("💡 Document uploaded. Question generation will proceed with available text.")
        else:
            st.info("💡 Document loaded. Question generation will proceed with available text.")
    # Allow question generation even if text extraction was limited - no dummy fallback
    
    # Check question limit based on user access level
    # Revised tiering: Free (15), Advance (90), Premium (unlimited)
    if st.session_state.user_access_level == "Premium":
        max_questions = 999999  # Unlimited for Premium
    elif st.session_state.user_access_level == "Advance":
        max_questions = ADVANCE_QUESTION_LIMIT
    else:
        max_questions = FREE_QUESTION_LIMIT
    remaining = max_questions - st.session_state.questions_answered
    
    if remaining <= 0:
        # Show paywall
        st.markdown("""
        <div class="alert-restricted">
            <h2>🚫 ACCESS RESTRICTED</h2>
            <p style="font-size: 1.2rem; margin: 1rem 0;">You reached the Free Limit (15 Questions).</p>
            <p style="font-size: 1rem; margin: 0.5rem 0;">Upgrade to continue practicing:</p>
            <ul style="text-align: left; display: inline-block; margin: 1rem auto;">
                <li><strong>Advance Mode (₱{ADVANCE_PAYMENT_AMOUNT}):</strong> Get 75 additional questions (Total: 90 questions)</li>
                <li><strong>Premium Mode (₱{PREMIUM_PAYMENT_AMOUNT}):</strong> Unlimited questions for 1 month</li>
            </ul>
            <p style="margin-top: 1rem; font-size: 0.9rem; opacity: 0.9;">Pay via GCash: <strong>{GCASH_NUMBER}</strong> | Email receipt to: <strong>{RECEIPT_EMAIL}</strong></p>
            <p style="margin-top: 0.5rem; font-size: 0.9rem; opacity: 0.9;">Access will be activated after receipt validation (12-24 hours).</p>
        </div>
        """.format(ADVANCE_PAYMENT_AMOUNT=ADVANCE_PAYMENT_AMOUNT, PREMIUM_PAYMENT_AMOUNT=PREMIUM_PAYMENT_AMOUNT, GCASH_NUMBER=GCASH_NUMBER, GCASH_NAME=GCASH_NAME, RECEIPT_EMAIL=RECEIPT_EMAIL), unsafe_allow_html=True)
        
        st.markdown("### 📍 Next Steps")
        st.info("👆 Use the sidebar navigation to go to **💳 Payment** page to upgrade.")
        st.stop()
    
    # Question generation form
    if not st.session_state.current_questions:
        render_card("⚙️ Generate Questions", """
        <p>Configure your practice exam settings.</p>
        """)
        
        col1, col2 = st.columns(2)
        with col1:
            difficulty = st.selectbox("Difficulty Level", ["Easy", "Average", "Difficult"])
            num_questions = st.number_input("Number of Questions", min_value=1, max_value=min(remaining, 6), value=min(6, remaining), help="Maximum 6 questions per practice exam")
        
        st.info("ℹ️ **Note:** You can generate up to 6 questions per practice exam. All questions will be Multiple Choice (MCQ) format only, based on your selected documents.")
        st.warning("⚠️ **Disclaimer:** Questions are generated for review purposes only. Always verify with official references and latest PH laws. Maximum 6 questions per practice exam session.")
        
        with col2:
            question_types = st.multiselect(
                "Question Types",
                ["Situational", "Recall/Definition", "Decision-making/Application", "Ethics/Procedure", "MCQ"],
                default=["Situational", "MCQ"],
                help="All questions will be Multiple Choice (MCQ) format only"
            )
        
        # Progress feedback container
        progress_container = st.empty()
        
        def update_progress(message: str):
            """Update progress feedback"""
            st.session_state.generation_progress = message
            try:
                with progress_container.container():
                    st.markdown(f"""
                    <div class="progress-feedback">
                        <div class="progress-step active">🔄 {message}</div>
                    </div>
                    """, unsafe_allow_html=True)
            except Exception:
                # Fallback if container update fails
                st.info(f"🔄 {message}")
        
        if st.button("🎯 Generate Questions", type="primary", use_container_width=True):
            if not question_types:
                st.error("Please select at least one question type.")
            else:
                # Check user access level
                user_access = st.session_state.user_access_level
                
                # FREE MODE: Use dummy questions (no document access)
                if user_access == "Free":
                    update_progress("📚 Generating practice questions for Free mode...")
                    try:
                        with st.spinner("Generating practice questions... Please wait..."):
                            questions = generate_default_dummy_questions(difficulty, num_questions, question_types)
                            questions = deduplicate_questions(questions)
                            # Limit to requested number
                            questions = questions[:num_questions]
                            
                            if len(questions) > 0:
                                update_progress(f"✅ Successfully generated {len(questions)} practice questions!")
                                # Store questions in session state
                                st.session_state.current_questions = questions
                                st.session_state.current_question_index = 0
                                st.session_state.answers = {}
                                st.session_state.exam_completed = False
                                st.session_state.exam_paused = False
                                progress_container.empty()
                                st.success(f"✅ Generated {len(questions)} practice questions!")
                                st.balloons()
                                st.rerun()
                            else:
                                update_progress("❌ Failed to generate questions")
                                progress_container.empty()
                                st.error("Failed to generate questions. Please try again.")
                    except Exception as e:
                        update_progress(f"❌ Error: {str(e)}")
                        progress_container.empty()
                        st.error(f"Error generating questions: {str(e)}")
                
                # ADVANCE/PREMIUM MODE: Use AI from selected documents (NO dummy questions)
                else:
                    # Get text from selected documents
                    text_for_generation = ""
                    
                    # Check if documents are selected in library (REQUIRED for Advance/Premium)
                    selected_paths = []
                    if "document_selections" in st.session_state:
                        current_user_email = st.session_state.user_email if st.session_state.user_logged_in else None
                        current_access_level = st.session_state.user_access_level if st.session_state.user_logged_in else None
                        all_docs = get_document_library(user_email=current_user_email, user_access_level=current_access_level)
                        selected_paths = [doc['filepath'] for doc in all_docs if st.session_state.document_selections.get(doc['filepath'], False)]
                    
                    # CRITICAL: For Advance/Premium, require document selection - NO dummy questions
                    if not selected_paths:
                        update_progress("❌ No documents selected.")
                        progress_container.empty()
                        st.error("❌ **No documents selected.** Please go to **📄 Upload Reviewer** page and select documents from the Document Library to generate AI-powered questions.")
                        st.info("💡 **Advance** and **Premium** users must select documents to generate questions. No dummy questions are available.")
                    else:
                        # Verify documents exist before extraction
                        existing_paths = [p for p in selected_paths if os.path.exists(p)]
                        if not existing_paths:
                            update_progress("❌ Selected documents not found on server.")
                            progress_container.empty()
                            st.error("❌ **Selected documents not found.** The files may have been deleted or moved. Please re-select documents from the Document Library.")
                        elif len(existing_paths) < len(selected_paths):
                            update_progress(f"⚠ {len(selected_paths) - len(existing_paths)} document(s) not found, processing {len(existing_paths)} available document(s)...")
                        
                        # Extract text from selected documents
                        update_progress(f"📄 Reading {len(existing_paths)} selected document(s)...")
                        try:
                            text_for_generation = extract_text_from_documents(
                                existing_paths,
                                max_pages_per_doc=None,  # No limit
                                max_total_chars=None,  # No limit
                                progress_callback=update_progress,
                            )
                            # Use any text that was extracted, no minimum requirement
                            if text_for_generation and text_for_generation.strip():
                                update_progress(f"✓ Extracted {len(text_for_generation)} characters from {len(selected_paths)} document(s)")
                                # Update session state to reflect documents are loaded
                                st.session_state.pdf_text = text_for_generation
                                st.session_state.pdf_name = f"Combined from {len(selected_paths)} document(s)"
                            else:
                                # No text extracted - try OCR explicitly on each document
                                update_progress(f"⚠ No text found, attempting OCR on {len(selected_paths)} document(s)...")
                                ocr_texts = []
                                for doc_path in selected_paths:
                                    if os.path.exists(doc_path):
                                        try:
                                            if doc_path.lower().endswith('.pdf'):
                                                update_progress(f"🔍 Running OCR on {os.path.basename(doc_path)}...")
                                                ocr_text, _ = extract_text_from_pdf(doc_path)
                                                if ocr_text and ocr_text.strip():
                                                    ocr_texts.append(ocr_text)
                                                    update_progress(f"✓ OCR extracted {len(ocr_text)} chars from {os.path.basename(doc_path)}")
                                            elif doc_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                                                update_progress(f"🔍 Reading image with OCR: {os.path.basename(doc_path)}...")
                                                ocr_text, _ = extract_text_from_image(doc_path)
                                                if ocr_text and ocr_text.strip():
                                                    ocr_texts.append(ocr_text)
                                                    update_progress(f"✓ OCR extracted {len(ocr_text)} chars from {os.path.basename(doc_path)}")
                                        except Exception as e:
                                            update_progress(f"❌ OCR failed for {os.path.basename(doc_path)}: {str(e)}")
                                
                                if ocr_texts:
                                    text_for_generation = "\n".join(ocr_texts)
                                    st.session_state.pdf_text = text_for_generation
                                    st.session_state.pdf_name = f"Combined from {len(selected_paths)} document(s) (OCR)"
                                    update_progress(f"✓ OCR extracted {len(text_for_generation)} total characters")
                                else:
                                    # Still no text - mark documents as selected but no text
                                    update_progress(f"❌ No readable text extracted even with OCR")
                                    st.session_state.pdf_text = "SELECTED_DOCUMENTS_NO_TEXT"
                                    st.session_state.pdf_name = f"{len(selected_paths)} document(s) selected"
                                    text_for_generation = ""
                        except Exception as e:
                            st.error(f"Error extracting documents: {str(e)}")
                            update_progress(f"❌ Error: {str(e)}")
                            text_for_generation = ""
                            import traceback
                            with st.expander("🔍 Error Details"):
                                st.code(traceback.format_exc())
                        
                        # Generate questions with AI from selected documents (NO dummy fallback)
                        # Proceed with question generation if we have any text at all
                        if not text_for_generation or not text_for_generation.strip():
                            update_progress("❌ No readable text could be extracted from the selected documents.")
                            progress_container.empty()
                            st.error("❌ **Cannot generate questions:** No readable text was extracted from your selected documents.")
                            
                            # Diagnostic information
                            with st.expander("🔍 Diagnostic Information"):
                                st.write(f"**Documents selected:** {len(selected_paths)}")
                                st.write(f"**Documents found:** {len([p for p in selected_paths if os.path.exists(p)])}")
                                st.write(f"**OCR Available:** {OCR_AVAILABLE}")
                                st.write(f"**EasyOCR Available:** {EASYOCR_AVAILABLE}")
                                st.write(f"**Tesseract Available:** {TESSERACT_AVAILABLE}")
                                st.write(f"**PDF Library:** {PDF_LIB}")
                                st.write(f"**Word Library:** {DOCX_AVAILABLE}")
                                
                                # Check if packages are actually installed
                                st.write("**Package Installation Status:**")
                                packages_to_check = {
                                    "PyMuPDF": "fitz",
                                    "pdfplumber": "pdfplumber",
                                    "python-docx": "docx",
                                    "EasyOCR": "easyocr",
                                    "pytesseract": "pytesseract",
                                    "Pillow": "PIL"
                                }
                                
                                for package_name, import_name in packages_to_check.items():
                                    try:
                                        __import__(import_name)
                                        st.write(f"- {package_name}: ✓ Installed")
                                    except ImportError as e:
                                        st.write(f"- {package_name}: ✗ Not installed ({str(e)[:50]})")
                                    except Exception as e:
                                        st.write(f"- {package_name}: ⚠ Error ({str(e)[:50]})")
                                
                                # List document details
                                st.write("**Document Details:**")
                                for doc_path in selected_paths:
                                    exists = os.path.exists(doc_path)
                                    size = os.path.getsize(doc_path) if exists else 0
                                    st.write(f"- {os.path.basename(doc_path)}: {'✓ Exists' if exists else '✗ Not found'} ({size} bytes)")
                                    
                                    # Try to extract text directly to see what happens
                                    if exists:
                                        try:
                                            if doc_path.lower().endswith('.pdf'):
                                                if PDF_AVAILABLE:
                                                    test_text, _ = extract_text_from_pdf(doc_path)
                                                    if test_text:
                                                        st.write(f"  - PDF extraction test: ✓ Success ({len(test_text)} chars)")
                                                    else:
                                                        st.write(f"  - PDF extraction test: ⚠ No text extracted")
                                                else:
                                                    st.write(f"  - PDF extraction test: ✗ PDF library not available")
                                            elif doc_path.lower().endswith(('.docx', '.doc')):
                                                if DOCX_AVAILABLE:
                                                    test_text, _ = extract_text_from_docx(doc_path)
                                                    if test_text:
                                                        st.write(f"  - Word extraction test: ✓ Success ({len(test_text)} chars)")
                                                    else:
                                                        st.write(f"  - Word extraction test: ⚠ No text extracted")
                                                else:
                                                    st.write(f"  - Word extraction test: ✗ Word library not available")
                                            elif doc_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                                                if OCR_AVAILABLE:
                                                    test_text, _ = extract_text_from_image(doc_path)
                                                    if test_text:
                                                        st.write(f"  - Image OCR test: ✓ Success ({len(test_text)} chars)")
                                                    else:
                                                        st.write(f"  - Image OCR test: ⚠ No text extracted")
                                                else:
                                                    st.write(f"  - Image OCR test: ✗ OCR not available")
                                        except Exception as e:
                                            st.write(f"  - Extraction test error: {str(e)[:100]}")
                            
                            st.warning("""
                            **Possible reasons:**
                            - Documents are scanned/image-based (require OCR)
                            - Documents are password-protected or encrypted
                            - Documents contain only images with no text layer
                            - OCR models may not be loaded (first-time setup can take several minutes)
                            
                            **Solutions:**
                            - Wait a few minutes and try again (EasyOCR downloads models on first use)
                            - Try uploading documents with machine-readable text
                            - Ensure documents are not password-protected
                            - Contact admin for assistance with document processing
                            """)
                            st.info("💡 **No dummy questions will be used.** Please fix your documents and try again.")
                        else:
                            # We have text - proceed with generation
                            update_progress(f"✓ Ready to generate from {len(text_for_generation)} characters of text from selected documents")
                            
                            update_progress("🚀 Starting AI question generation from selected documents...")
                            try:
                                # Use spinner for long-running operation
                                with st.spinner("Generating questions from your selected documents... This may take 30-60 seconds. Please wait..."):
                                    questions = generate_questions(
                                        text_for_generation,
                                        difficulty,
                                        num_questions,
                                        question_types,
                                        progress_callback=update_progress
                                    )
                                    
                                    # Process AI-generated questions
                                    if questions is None:
                                        update_progress("❌ Question generation returned None")
                                        progress_container.empty()
                                        st.error("Question generation returned None. This is a bug - please report this error.")
                                    elif isinstance(questions, list):
                                        if len(questions) > 0:
                                            # Ensure we have the full requested count - if less, show warning
                                            if len(questions) < num_questions:
                                                st.warning(f"⚠️ Generated {len(questions)} questions (requested {num_questions}). This may be due to document content limitations.")
                                            # Use all available questions up to requested count
                                            questions = questions[:num_questions]
                                            update_progress(f"✅ Successfully generated {len(questions)} questions from selected documents!")
                                            # Store questions in session state
                                            st.session_state.current_questions = questions
                                            st.session_state.current_question_index = 0
                                            st.session_state.answers = {}
                                            st.session_state.exam_completed = False
                                            st.session_state.exam_paused = False
                                            progress_container.empty()
                                            st.success(f"✅ Generated {len(questions)} unique questions from your selected documents!")
                                            st.balloons()
                                            # Force rerun to display questions immediately
                                            st.rerun()
                                        else:
                                            update_progress("❌ Question generation returned empty list")
                                            progress_container.empty()
                                            st.error("Failed to generate questions - returned empty list. Please check:")
                                            st.markdown("""
                                            - Are documents properly selected in the Document Library?
                                            - Please try again or contact support
                                            - Do you have API credits available?
                                            - Try selecting fewer documents or reducing the number of questions
                                            """)
                                            # Show debug info
                                            with st.expander("🔍 Debug Information"):
                                                text_len = len(text_for_generation) if text_for_generation else 0
                                                st.write(f"Text length: {text_len}")
                                                st.write(f"Question types: {question_types}")
                                                st.write(f"Difficulty: {difficulty}")
                                                st.write(f"Number requested: {num_questions}")
                                                st.write("Questions returned: 0 (empty list)")
                                    else:
                                        update_progress("❌ No questions generated. Please check your documents.")
                                        progress_container.empty()
                                        st.error("Failed to generate questions. Please check:")
                                        st.markdown("""
                                        - Are documents properly selected in the Document Library?
                                        - Please try again or contact support
                                        - Try selecting fewer documents or reducing the number of questions
                                        """)
                                        # Show debug info
                                        with st.expander("🔍 Debug Information"):
                                            text_len = len(text_for_generation) if text_for_generation else 0
                                            q_count = len(questions) if questions else 0
                                            st.write(f"Text length: {text_len}")
                                            st.write(f"Question types: {question_types}")
                                            st.write(f"Difficulty: {difficulty}")
                                            st.write(f"Number requested: {num_questions}")
                                            st.write(f"Questions returned: {q_count}")
                            except Exception as e:
                                update_progress(f"❌ Error: {str(e)}")
                                progress_container.empty()
                                st.error(f"Error generating questions: {str(e)}")
                                import traceback
                                with st.expander("🔍 Error Details"):
                                    st.code(traceback.format_exc())
    
    # Display current question
    if st.session_state.current_questions and len(st.session_state.current_questions) > 0:
        questions = st.session_state.current_questions
        current_idx = st.session_state.current_question_index
        
        if current_idx < len(questions):
            q = questions[current_idx]
            
            # Progress
            progress = (current_idx + 1) / len(questions)
            st.markdown(f"""
            <div style="margin: 1rem 0;">
                <div style="display: flex; justify-content: space-between; margin-bottom: 0.5rem;">
                    <span style="color: #d4af37; font-weight: 600;">TRAINING PROGRESS</span>
                    <span style="color: #d4af37; font-weight: 700;">Question {current_idx + 1}/{len(questions)}</span>
                </div>
                <div class="progress-container">
                    <div class="progress-bar" style="width: {progress * 100}%;"></div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Question card
            render_card(f"Question {current_idx + 1}", f"""
            <p style="font-size: 1.1rem; margin-bottom: 1rem;">{q['question']}</p>
            """)
            
            # Navigation buttons at top
            nav_col1, nav_col2, nav_col3 = st.columns([1, 1, 2])
            with nav_col1:
                if st.button("🏠 Back to Home", type="secondary", use_container_width=True):
                    st.session_state.current_questions = []
                    st.session_state.current_question_index = 0
                    st.session_state.answers = {}
                    st.rerun()
            with nav_col2:
                if st.button("⏸️ Pause Exam", type="secondary", use_container_width=True):
                    # Store current progress
                    st.session_state.exam_paused = True
                    st.session_state.paused_at_index = current_idx
                    st.success("⏸️ Exam paused. Your progress has been saved. Answer selection is now disabled until you resume.")
                    st.rerun()
            
            st.markdown("---")
            
            # Check if exam is paused - disable answer selection
            is_paused = st.session_state.get("exam_paused", False)
            
            if is_paused:
                st.warning("⏸️ Exam is paused. Click 'Resume Exam' to continue answering questions.")
                st.info("💡 Your progress has been saved. Use the navigation buttons above to resume or go home.")
                # Don't show answer options when paused
                answer = None
            else:
                # Answer options - All questions are Multiple Choice only
                # Use unique key to prevent duplicate key errors
                unique_key = f"answer_radio_q{current_idx}_{q.get('type', 'mcq')}"
                
                # Ensure we have options
                if not q.get('options') or len(q['options']) < 4:
                    st.error("⚠️ This question has incomplete options. Please contact support.")
                    st.stop()
                
                # Validate options are not generic
                has_generic = False
                for opt in q['options']:
                    opt_lower = opt.lower().strip()
                    if (len(opt_lower.split()) <= 2 and any(p in opt_lower for p in ['option a', 'option b', 'option c', 'option d'])) or \
                       opt_lower in ['a', 'b', 'c', 'd', 'a.', 'b.', 'c.', 'd.'] or \
                       ('option' in opt_lower and len(opt_lower.split()) <= 3):
                        has_generic = True
                        break
                
                if has_generic:
                    st.error("⚠️ This question has generic options. Regenerating...")
                    # Clear this question and regenerate
                    st.session_state.current_questions = []
                    st.session_state.current_question_index = 0
                    st.rerun()
                
                # Format options with A, B, C, D labels and display radio button
                formatted_options = [f"{chr(65+i)}. {opt}" for i, opt in enumerate(q['options'])]
                answer_idx = st.radio("Select your answer:", formatted_options, key=unique_key)
                # Extract the actual answer text
                answer = q['options'][formatted_options.index(answer_idx)]
            
            st.markdown("---")
            
            # Navigation buttons for questions (Previous, Submit, Skip, Next)
            nav_btn_col1, nav_btn_col2, nav_btn_col3, nav_btn_col4 = st.columns([1, 1, 1, 1])
            
            with nav_btn_col1:
                # Previous button - allow going back to previous questions (including skipped ones)
                if current_idx > 0:
                    if st.button("⬅️ Back", type="secondary", use_container_width=True, disabled=is_paused, key=f"prev_{current_idx}"):
                        if not is_paused:
                            st.session_state.current_question_index = current_idx - 1
                            st.rerun()
                else:
                    st.button("⬅️ Back", type="secondary", use_container_width=True, disabled=True, key=f"prev_{current_idx}")
            
            with nav_btn_col2:
                # Submit Answer button
                if st.button("✅ Submit", type="primary", use_container_width=True, disabled=is_paused, key=f"submit_{current_idx}"):
                    if not is_paused and answer is not None:
                        st.session_state.answers[current_idx] = answer
                        st.session_state.current_question_index += 1
                        # Only count when exam is fully completed
                        if st.session_state.current_question_index >= len(questions):
                            if "exam_completed" not in st.session_state or not st.session_state.get("exam_completed", False):
                                st.session_state.questions_answered += len(questions)
                                # Update user's question count in database
                                if st.session_state.user_logged_in and st.session_state.user_email:
                                    update_user_questions_answered(st.session_state.user_email, len(questions))
                                st.session_state.exam_completed = True
                        st.rerun()
                    elif is_paused:
                        st.warning("⏸️ Exam is paused. Please resume to continue.")
            
            with nav_btn_col3:
                # Skip button - allows skipping and going back later
                if st.button("⏭️ Skip", type="secondary", use_container_width=True, disabled=is_paused, key=f"skip_{current_idx}"):
                    if not is_paused:
                        # Mark as skipped (don't store answer, but allow going back)
                        if current_idx not in st.session_state.answers:
                            st.session_state.answers[current_idx] = None  # Mark as skipped
                        st.session_state.current_question_index += 1
                        # Only count when exam is fully completed
                        if st.session_state.current_question_index >= len(questions):
                            if "exam_completed" not in st.session_state or not st.session_state.get("exam_completed", False):
                                st.session_state.questions_answered += len(questions)
                                # Update user's question count in database
                                if st.session_state.user_logged_in and st.session_state.user_email:
                                    update_user_questions_answered(st.session_state.user_email, len(questions))
                                st.session_state.exam_completed = True
                        st.rerun()
                    elif is_paused:
                        st.warning("⏸️ Exam is paused. Please resume to continue.")
            
            with nav_btn_col4:
                # Next button (if not last question)
                if current_idx < len(questions) - 1:
                    if st.button("➡️ Next", type="secondary", use_container_width=True, disabled=is_paused, key=f"next_{current_idx}"):
                        if not is_paused:
                            # Save current answer if selected
                            if answer is not None:
                                st.session_state.answers[current_idx] = answer
                            st.session_state.current_question_index += 1
                            st.rerun()
                else:
                    st.button("➡️ Next", type="secondary", use_container_width=True, disabled=True, key=f"next_{current_idx}")
        else:
            # Show results
            st.markdown("# 📊 Exam Results")
            
            score = 0
            total = len(questions)
            wrong_answers = []
            
            for idx, q in enumerate(questions):
                user_answer = st.session_state.answers.get(idx, "Not answered")
                correct = q['correct_answer']
                is_correct = str(user_answer).strip().lower() == str(correct).strip().lower()
                
                if is_correct:
                    score += 1
                else:
                    wrong_answers.append({
                        "question": q['question'],
                        "your_answer": user_answer,
                        "correct_answer": correct,
                        "explanation": q.get('explanation', '')
                    })
            
            # Score display
            percentage = (score / total * 100) if total > 0 else 0
            st.markdown(f"""
            <div style="text-align: center; padding: 2rem; background: rgba(30, 58, 95, 0.6); border-radius: 16px; border: 2px solid #d4af37;">
                <h2 style="color: #d4af37; font-size: 3rem; margin: 0;">{score}/{total}</h2>
                <p style="color: #e0e0e0; font-size: 1.5rem; margin: 0.5rem 0;">{percentage:.1f}%</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Performance Insights
            st.markdown("### 📈 Performance Summary")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("✅ Correct", score)
            with col2:
                st.metric("❌ Incorrect", total - score)
            with col3:
                st.metric("📊 Accuracy", f"{percentage:.1f}%")
            
            # Performance message
            if percentage >= 90:
                perf_msg = "🌟 Outstanding! You're excelling in your preparation!"
                perf_color = "#28a745"
            elif percentage >= 75:
                perf_msg = "👍 Great job! You're on the right track!"
                perf_color = "#20c997"
            elif percentage >= 60:
                perf_msg = "💪 Good start! Keep practicing to improve!"
                perf_color = "#ffc107"
            else:
                perf_msg = "📚 Keep studying! Review the materials and try again!"
                perf_color = "#dc3545"
            
            st.markdown(f"""
            <div style="background: rgba(30, 58, 95, 0.6); border-left: 4px solid {perf_color}; padding: 1rem; border-radius: 8px; margin: 1rem 0;">
                <p style="color: {perf_color}; font-weight: 600; font-size: 1.1rem; margin: 0;">{perf_msg}</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Export options
            st.markdown("### 📥 Export Exam")
            col1, col2 = st.columns(2)
            with col1:
                exam_title = st.text_input("Exam Title", value=f"Criminology Practice Exam - {datetime.now().strftime('%Y-%m-%d')}")
            
            # Only show export buttons if libraries are available
            export_cols = []
            if REPORTLAB_AVAILABLE:
                export_cols.append(1)
            if DOCX_AVAILABLE:
                export_cols.append(2)
            
            if export_cols:
                cols = st.columns(len(export_cols))
                col_idx = 0
                
                if REPORTLAB_AVAILABLE:
                    with cols[col_idx]:
                        try:
                            pdf_data = export_to_pdf(questions, exam_title)
                            if pdf_data:
                                st.download_button(
                                    "📄 Download as PDF",
                                    data=pdf_data,
                                    file_name=f"{exam_title.replace(' ', '_')}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                        except Exception:
                            # Silently fail - don't show error to user
                            pass
                    col_idx += 1
                
                if DOCX_AVAILABLE:
                    with cols[col_idx] if len(export_cols) > 1 else st.container():
                        try:
                            docx_data = export_to_docx(questions, exam_title)
                            if docx_data:
                                st.download_button(
                                    "📝 Download as DOCX",
                                    data=docx_data,
                                    file_name=f"{exam_title.replace(' ', '_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                        except Exception:
                            # Silently fail - don't show error to user
                            pass
            else:
                # Only show info message if no export options are available
                st.info("💡 Export functionality requires additional libraries. For local development, install with: `pip install reportlab python-docx`")
            
            # Regenerate button
            st.markdown("---")
            if st.button("🔄 Regenerate Exam (New Questions)", type="secondary", use_container_width=True):
                st.session_state.current_questions = []
                st.session_state.current_question_index = 0
                st.session_state.answers = {}
                st.rerun()
            
            # Review wrong answers
            if wrong_answers:
                st.markdown("### ❌ Review Wrong Answers")
                for i, wa in enumerate(wrong_answers):
                    with st.expander(f"Question {i+1}: {wa['question'][:50]}..."):
                        st.write(f"**Your Answer:** {wa['your_answer']}")
                        st.write(f"**Correct Answer:** {wa['correct_answer']}")
                        st.write(f"**Explanation:** {wa['explanation']}")
            
            # Check if Free user reached 15 questions
            is_free_user = st.session_state.user_access_level == "Free"
            reached_free_limit = st.session_state.questions_answered >= FREE_QUESTION_LIMIT
            
            if is_free_user and reached_free_limit:
                st.markdown("---")
                # Access Restriction Card with Premium Upsell
                st.markdown("""
                <div style="background: linear-gradient(135deg, #dc3545 0%, #c82333 100%); border: 3px solid #ff6b6b; border-radius: 16px; padding: 2.5rem; text-align: center; margin: 2rem 0; box-shadow: 0 8px 32px rgba(220, 53, 69, 0.5);">
                    <h2 style="color: white; font-size: 2rem; margin: 0 0 1rem 0;">🚫 ACCESS RESTRICTED</h2>
                    <p style="color: white; font-size: 1.2rem; margin: 1rem 0;">You've completed your free 15 questions!</p>
                </div>
                """, unsafe_allow_html=True)
                
                st.markdown("""
                <div style="background: linear-gradient(135deg, #d4af37 0%, #f4d03f 100%); border: 3px solid #d4af37; border-radius: 16px; padding: 2rem; margin: 1rem 0; box-shadow: 0 8px 32px rgba(212, 175, 55, 0.5);">
                    <h2 style="color: #1a2332; font-size: 1.8rem; margin: 0 0 1rem 0;">🔑 UNLOCK PREMIUM ACCESS</h2>
                    <p style="color: #1a2332; font-size: 1.1rem; margin: 0.5rem 0; font-weight: 600;">Get 100+ additional questions and advanced features:</p>
                    <ul style="text-align: left; color: #1a2332; font-size: 1rem; margin: 1rem auto; display: inline-block;">
                        <li>✅ 100+ additional practice questions</li>
                        <li>✅ Access to advanced situational scenarios</li>
                        <li>✅ Better preparation with deeper insights</li>
                        <li>✅ Real-world readiness training</li>
                        <li>✅ Full access to all premium PDF resources</li>
                    </ul>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("💳 Go to Payment", type="primary", use_container_width=True):
                        st.info("💡 Please use the sidebar to navigate to the Payment page to upgrade.")
                with col2:
                    if st.button("💳 Go to Payment", type="primary", use_container_width=True):
                        st.session_state.page = "💳 Payment"
                        st.rerun()
            
            # Reset or new set (only if not at limit)
            if not (is_free_user and reached_free_limit):
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🔄 Generate New Set", type="primary", use_container_width=True):
                        st.session_state.current_questions = []
                        st.session_state.current_question_index = 0
                        st.session_state.answers = {}
                        st.session_state.exam_completed = False
                        st.rerun()
                with col2:
                    if st.button("🏠 Back to Home", use_container_width=True):
                        st.session_state.current_questions = []
                        st.session_state.current_question_index = 0
                        st.session_state.answers = {}
                        st.session_state.exam_completed = False
                        st.rerun()
            else:
                st.info("💡 Upgrade to Premium to continue practicing with more questions!")

# ============================================================================
# PAGE: PREMIUM ACCESS
# ============================================================================

# ============================================================================
# PAGE: PREMIUM ACCESS - REMOVED
# ============================================================================
# Premium Access page has been removed. Users can upgrade via Payment page.

# ============================================================================
# PAGE: PAYMENT
# ============================================================================

elif page == "💳 Payment":
    st.markdown("# 💳 Payment & Receipt Upload")
    
    render_card("💰 Payment Tiers", f"""
    <h3 style="color: #d4af37; margin-top: 0;">Choose Your Plan:</h3>
    
    <div style="background: rgba(30, 58, 95, 0.6); border: 2px solid #2d4a6b; border-radius: 12px; padding: 1.5rem; margin: 1rem 0;">
        <h4 style="color: #d4af37; margin-top: 0;">🆓 FREE MODE</h4>
        <p style="margin: 0.5rem 0;"><strong>15 questions</strong> - No payment required</p>
    </div>
    
    <div style="background: rgba(30, 58, 95, 0.6); border: 2px solid #d4af37; border-radius: 12px; padding: 1.5rem; margin: 1rem 0;">
        <h4 style="color: #d4af37; margin-top: 0;">⚡ ADVANCE MODE</h4>
        <p style="margin: 0.5rem 0;"><strong>₱{ADVANCE_PAYMENT_AMOUNT}</strong> - Get <strong>75 additional questions</strong> (Total: 90 questions)</p>
        <p style="margin: 0.5rem 0; font-size: 0.9rem; color: #b0b0b0;">Perfect for focused practice sessions</p>
    </div>
    
    <div style="background: rgba(212, 175, 55, 0.2); border: 2px solid #d4af37; border-radius: 12px; padding: 1.5rem; margin: 1rem 0;">
        <h4 style="color: #d4af37; margin-top: 0;">👑 PREMIUM MODE</h4>
        <p style="margin: 0.5rem 0;"><strong>₱{PREMIUM_PAYMENT_AMOUNT}</strong> - <strong>Unlimited questions</strong> for 1 month</p>
        <p style="margin: 0.5rem 0; font-size: 0.9rem; color: #b0b0b0;">Best value for intensive exam preparation</p>
    </div>
    
    <h3 style="color: #d4af37; margin-top: 1.5rem;">💳 Payment Method:</h3>
    <ul>
        <li>Pay via GCash</li>
        <li>GCash Number: <strong>{GCASH_NUMBER}</strong></li>
        <li>Account Name: <strong>{GCASH_NAME}</strong></li>
    </ul>
    
    <h3 style="color: #d4af37; margin-top: 1rem;">📩 Receipt Submission:</h3>
    <p>Email your payment receipt to: <strong>{RECEIPT_EMAIL}</strong></p>
    <p style="font-size: 0.9rem; color: #b0b0b0;">Please specify in your email which plan you're purchasing (Advance or Premium)</p>
    
    <div style="background: rgba(212, 175, 55, 0.2); border-left: 4px solid #d4af37; padding: 1rem; border-radius: 8px; margin-top: 1rem;">
        <p style="color: #d4af37; font-weight: 600; margin: 0;">⚠️ Important: Access will be activated after receipt validation.</p>
        <p style="color: #e0e0e0; margin: 0.5rem 0 0 0; font-size: 0.9rem;">Please allow 12-24 hours for processing after submitting your receipt.</p>
    </div>
    """)
    
    with st.form("payment_form"):
        st.markdown("### 📝 Payment Details")
        full_name = st.text_input("Full Name *", placeholder="Juan dela Cruz")
        email = st.text_input("Email Address *", placeholder="juan@example.com")
        gcash_ref = st.text_input("GCash Reference Number (Optional)", placeholder="Reference number from GCash")
        receipt_file = st.file_uploader("Upload Receipt (Image/PDF)", type=["png", "jpg", "jpeg", "pdf"])
        
        submitted = st.form_submit_button("📤 Submit Payment Request", type="primary", use_container_width=True)
        
        if submitted:
            if not full_name or not email:
                st.error("Please fill in all required fields (marked with *)")
            elif not receipt_file:
                st.warning("⚠️ Receipt upload is recommended for faster processing.")
                if st.button("Submit Anyway"):
                    save_payment_receipt(full_name, email, gcash_ref, "")
                    st.success("✅ Payment request submitted! Please send receipt to " + RECEIPT_EMAIL)
            else:
                # Save receipt file
                os.makedirs(RECEIPTS_DIR, exist_ok=True)
                receipt_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{receipt_file.name}"
                receipt_path = os.path.join(RECEIPTS_DIR, receipt_filename)
                
                with open(receipt_path, "wb") as f:
                    f.write(receipt_file.getbuffer())
                
                save_payment_receipt(full_name, email, gcash_ref, receipt_filename)
                st.success("✅ Payment request and receipt submitted! We'll review it within 12-24 hours.")

# ============================================================================
# PAGE: ADMIN PANEL
# ============================================================================

elif page == "🛠️ Admin Panel":
    st.markdown("# 🛠️ Admin Command Center")
    
    # Admin login
    if not st.session_state.admin_logged_in:
        render_card("🔐 Admin Login", """
        <p>Enter admin password to access the command center.</p>
        """)
        
        admin_password = st.text_input("Admin Password", type="password", key="admin_pw")
        
        if st.button("🔓 Login", type="primary", use_container_width=True):
            # Get admin password from secrets.toml (required)
            correct_password = get_admin_password()
            
            if not correct_password:
                st.error("❌ Admin password not configured. Please set ADMIN_PASSWORD in secrets.toml or environment variables.")
                st.stop()
            
            if admin_password == correct_password:
                st.session_state.admin_logged_in = True
                st.success("✅ Admin access granted!")
                st.rerun()
            else:
                st.error("❌ Invalid password.")
        st.stop()
    
    # Admin logged in - show tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["🔑 Generate Codes", "📋 Code Database", "💳 Receipt Validation", "📄 PDF Management", "👥 User Management"])
    
    with tab1:
        st.markdown("### 🔑 Generate Access Codes")
        
        with st.form("generate_codes"):
            col1, col2 = st.columns(2)
            with col1:
                num_codes = st.number_input("Number of Codes", min_value=1, max_value=100, value=5)
                code_length = st.number_input("Code Length", min_value=8, max_value=20, value=12)
                access_level = st.selectbox(
                    "Access Level",
                    ["Advance", "Premium"],
                    help="Advance: +75 questions (total 90). Premium: Unlimited questions for 1 month"
                )
            with col2:
                expiry_days = st.number_input("Expiry (Days, 0 = No expiry)", min_value=0, max_value=365, value=30)
                max_uses = st.number_input("Max Uses per Code", min_value=1, max_value=100, value=1)
            
            if st.form_submit_button("🎯 Generate Codes", type="primary", use_container_width=True):
                expiry = expiry_days if expiry_days > 0 else None
                codes = create_premium_codes(num_codes, code_length, expiry, max_uses, "admin", access_level)
                st.session_state.generated_codes = codes
                st.session_state.generated_codes_access_level = access_level
                st.success(f"✅ Generated {len(codes)} {access_level} access codes!")
                st.rerun()
        
        # Display codes and download button outside form
        if "generated_codes" in st.session_state and st.session_state.generated_codes:
            access_level = st.session_state.get("generated_codes_access_level", "Premium")
            codes_df = pd.DataFrame({
                "Code": st.session_state.generated_codes,
                "Access Level": [access_level] * len(st.session_state.generated_codes)
            })
            st.dataframe(codes_df, use_container_width=True, hide_index=True)
            
            # Export CSV
            csv = codes_df.to_csv(index=False)
            st.download_button(
                "⬇️ Download Codes as CSV",
                data=csv,
                file_name=f"{access_level.lower()}_codes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    
    with tab2:
        st.markdown("### 📋 Code Database")
        
        # Check if access_level column exists
        columns = get_table_columns("premium_codes")
        has_access_level = 'access_level' in columns
        
        table_name = get_table_name("premium_codes")
        if has_access_level:
            cursor = execute_query(f"""
                SELECT code, status, expiry_date, max_uses, uses_left, created_at, access_level
                FROM {table_name}
                ORDER BY created_at DESC
            """)
            codes_data = cursor.fetchall()
            cursor.close()
            if codes_data:
                df = pd.DataFrame(codes_data, columns=["Code", "Status", "Expiry", "Max Uses", "Uses Left", "Created At", "Access Level"])
            else:
                df = None
        else:
            cursor = execute_query(f"""
                SELECT code, status, expiry_date, max_uses, uses_left, created_at
                FROM {table_name}
                ORDER BY created_at DESC
            """)
            codes_data = cursor.fetchall()
            cursor.close()
            if codes_data:
                df = pd.DataFrame(codes_data, columns=["Code", "Status", "Expiry", "Max Uses", "Uses Left", "Created At"])
                df["Access Level"] = "Premium"  # Default for old codes
            else:
                df = None
        
        if df is not None and not df.empty:
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            # Search and filter
            search_code = st.text_input("🔍 Search Code", placeholder="Enter code to search")
            if search_code:
                filtered = df[df["Code"].str.contains(search_code.upper(), na=False)]
                st.dataframe(filtered, use_container_width=True, hide_index=True)
        else:
            st.info("No codes generated yet.")
    
    with tab3:
        st.markdown("### 💳 Receipt Validation")
        
        table_name = get_table_name("payment_receipts")
        cursor = execute_query(f"""
            SELECT id, full_name, email, gcash_reference, receipt_filename, status, submitted_at
            FROM {table_name}
            ORDER BY submitted_at DESC
        """)
        
        receipts_data = cursor.fetchall()
        cursor.close()
        
        if receipts_data:
            df = pd.DataFrame(receipts_data, columns=["ID", "Name", "Email", "GCash Ref", "Receipt", "Status", "Submitted"])
            
            # Filter by status
            status_filter = st.selectbox("Filter by Status", ["All", "Pending", "Approved", "Rejected"])
            if status_filter != "All":
                df = df[df["Status"] == status_filter]
            
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            # Display receipt images
            st.markdown("### 📷 View Receipt")
            receipt_id_view = st.number_input("Receipt ID to View", min_value=1, key="receipt_view_id")
            
            if receipt_id_view:
                table_name = get_table_name("payment_receipts")
                cursor = execute_query(f"""
                    SELECT receipt_filename, full_name, email, status
                    FROM {table_name}
                    WHERE id = %s
                """, (receipt_id_view,))
                receipt_info = cursor.fetchone()
                cursor.close()
                
                if receipt_info:
                    receipt_filename, full_name, email, status = receipt_info
                    st.markdown(f"**Name:** {full_name} | **Email:** {email} | **Status:** {status}")
                    
                    if receipt_filename:
                        receipt_path = os.path.join(RECEIPTS_DIR, receipt_filename)
                        if os.path.exists(receipt_path):
                            # Check if it's an image
                            if receipt_filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                                st.image(receipt_path, caption=f"Receipt for {full_name}", use_container_width=True)
                            elif receipt_filename.lower().endswith('.pdf'):
                                st.info("📄 PDF receipt file. Download to view.")
                                with open(receipt_path, "rb") as f:
                                    st.download_button(
                                        "📥 Download Receipt PDF",
                                        data=f.read(),
                                        file_name=receipt_filename,
                                        mime="application/pdf"
                                    )
                            else:
                                st.warning(f"⚠️ Receipt file format not supported for preview: {receipt_filename}")
                        else:
                            st.error(f"❌ Receipt file not found: {receipt_filename}")
                    else:
                        st.info("No receipt file uploaded for this submission.")
                else:
                    st.warning(f"⚠️ Receipt ID {receipt_id_view} not found.")
            
            # Approve/Reject actions
            st.markdown("### ⚡ Actions")
            receipt_id = st.number_input("Receipt ID to Update", min_value=1, key="receipt_update_id")
            new_status = st.selectbox("New Status", ["Pending", "Approved", "Rejected"])
            admin_notes = st.text_area("Admin Notes")
            
            if st.button("💾 Update Status", type="primary"):
                # Get receipt info before updating
                table_name = get_table_name("payment_receipts")
                cursor = execute_query(f"""
                    SELECT email, full_name
                    FROM {table_name}
                    WHERE id = %s
                """, (receipt_id,))
                receipt_info = cursor.fetchone()
                cursor.close()
                
                # Update receipt status
                cursor = execute_query(f"""
                    UPDATE {table_name}
                    SET status = %s, reviewed_at = %s, reviewed_by = %s, notes = %s
                    WHERE id = %s
                """, (new_status, datetime.now().isoformat(), "admin", admin_notes, receipt_id))
                cursor.close()
                
                if is_snowflake():
                    db_conn.commit()
                
                # If approved, update user access level
                if new_status == "Approved" and receipt_info:
                    receipt_email = receipt_info[0]
                    receipt_name = receipt_info[1]
                    
                    # Determine access level based on payment amount or notes
                    # Check if user exists, create if not
                    table_name = get_table_name("users")
                    cursor = execute_query(f"SELECT email FROM {table_name} WHERE email = %s", (receipt_email.lower(),))
                    user_exists = cursor.fetchone()
                    
                    cursor.close()
                    if not user_exists:
                        # Create user account
                        table_name = get_table_name("users")
                        cursor = execute_query(f"""
                            INSERT INTO {table_name} (email, access_level, questions_answered, created_at, last_login, is_admin)
                            VALUES (%s, %s, %s, %s, %s, %s)
                        """, (receipt_email.lower(), "Free", 0, datetime.now().isoformat(), datetime.now().isoformat(), 0))
                        cursor.close()
                        if is_snowflake():
                            db_conn.commit()
                    
                    # Determine access level - check notes for "Advance" or "Premium", default to Premium
                    access_level = "Premium"  # Default for approved receipts
                    if admin_notes and "advance" in admin_notes.lower():
                        access_level = "Advance"
                    elif admin_notes and "premium" in admin_notes.lower():
                        access_level = "Premium"
                    else:
                        # Try to infer from payment amount (if stored) or default to Premium
                        access_level = "Premium"
                    
                    # Update user access level
                    update_user_access_level(receipt_email.lower(), access_level)
                    st.success(f"✅ Receipt #{receipt_id} updated to {new_status} and user {receipt_email} access set to {access_level}")
                else:
                    st.success(f"✅ Receipt #{receipt_id} updated to {new_status}")
                
                # Clear caches - use try/except in case function not defined yet
                try:
                    get_all_users_cached.clear()
                except:
                    pass
                st.rerun()
        else:
            st.info("No payment receipts submitted yet.")
    
    with tab4:
        st.markdown("### 📄 PDF Resource Management")

        # Admin reset: remove all documents and start fresh
        with st.expander("🧹 Reset Document Library (Danger Zone)"):
            st.warning(
                "This will delete **all** PDF records and uploaded files from the server. "
                "Use only if you want to completely reset the document library."
            )
            confirm_reset = st.checkbox("I understand this will permanently delete all documents.", key="confirm_reset_docs")
            if st.button("🧹 Delete All Documents", type="secondary", disabled=not confirm_reset):
                try:
                    # Delete all rows from pdf_resources table
                    table_name = get_table_name("pdf_resources")
                    cursor = execute_query(f"DELETE FROM {table_name}")
                    cursor.close()
                    if is_snowflake():
                        db_conn.commit()

                    # Remove files from admin_docs and uploads directories
                    for folder in [ADMIN_DOCS_DIR, UPLOADS_DIR]:
                        if os.path.exists(folder):
                            for fname in os.listdir(folder):
                                fpath = os.path.join(folder, fname)
                                try:
                                    if os.path.isfile(fpath):
                                        os.remove(fpath)
                                except Exception:
                                    pass
                    get_pdf_resources.clear()
                    get_document_library.clear()
                    st.success("✅ Document library reset. All documents have been removed.")
                    st.experimental_rerun()
                except Exception as e:
                    st.error(f"❌ Failed to reset document library: {str(e)}")

        # Upload new PDF
        st.markdown("#### 📤 Upload New PDF")
        with st.form("upload_pdf_admin_tab4"):
            uploaded_pdf = st.file_uploader("Upload PDF, Word, or image", type=["pdf", "docx", "doc", "png", "jpg", "jpeg"], key="admin_pdf_upload_tab4")
            is_premium_only = st.checkbox("Premium Only", value=False, help="Only Premium users can download this PDF")
            use_for_ai = st.checkbox("Use for AI Question Generation", value=True, help="Include this PDF in AI question generation")
            is_downloadable = st.checkbox("Allow Download", value=True, help="✅ Downloadable: Users can download this file | ❌ Preview-only: Users can only preview, not download")
            description = st.text_area("Description (optional)", placeholder="Brief description of this PDF")
            
            if st.form_submit_button("📤 Upload PDF", type="primary", use_container_width=True):
                if uploaded_pdf:
                    # Save file to admin_docs directory
                    os.makedirs(ADMIN_DOCS_DIR, exist_ok=True)
                    filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uploaded_pdf.name}"
                    filepath = os.path.join(ADMIN_DOCS_DIR, filename)
                    
                    with open(filepath, "wb") as f:
                        f.write(uploaded_pdf.getbuffer())
                    
                    save_pdf_resource(filename, filepath, is_premium_only, use_for_ai, st.session_state.user_email or "admin", description, is_downloadable)
                    st.success(f"✅ PDF uploaded successfully: {filename}")
                    get_document_library.clear()  # Clear cache
                    st.rerun()
                else:
                    st.error("Please select a PDF file to upload.")
        
        st.markdown("---")
        
        # List PDFs
        st.markdown("#### 📋 Manage PDFs")
        pdf_resources = get_pdf_resources()
        
        if pdf_resources:
            for pdf in pdf_resources:
                download_status = "📥 Downloadable" if pdf.get('is_downloadable', True) else "👁️ Preview Only"
                with st.expander(f"📄 {pdf['filename']} - {'🔒 Premium' if pdf['is_premium_only'] else '🆓 Free'} - {'🤖 AI Enabled' if pdf['use_for_ai_generation'] else '❌ AI Disabled'} - {download_status}"):
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.write(f"**Description:** {pdf.get('description', 'No description')}")
                        st.write(f"**Uploaded:** {pdf['uploaded_at']}")
                        st.write(f"**Uploaded by:** {pdf.get('uploaded_by', 'Unknown')}")
                        st.write(f"**Download Status:** {'✅ Downloadable' if pdf.get('is_downloadable', True) else '👁️ Preview Only'}")
                    with col2:
                        premium_toggle = st.checkbox("Premium Only", value=pdf['is_premium_only'], key=f"premium_{pdf['id']}")
                        ai_toggle = st.checkbox("Use for AI", value=pdf['use_for_ai_generation'], key=f"ai_{pdf['id']}")
                        downloadable_toggle = st.checkbox(
                            "Allow Download",
                            value=pdf.get('is_downloadable', True),
                            key=f"download_{pdf['id']}",
                            help="✅ Downloadable: Users can download | ❌ Preview-only: Users can only preview",
                        )

                        if st.button("💾 Update Settings", key=f"update_{pdf['id']}", use_container_width=True):
                            table_name = get_table_name("pdf_resources")
                            # Check if is_downloadable column exists
                            columns = get_table_columns("pdf_resources")
                            has_downloadable = 'is_downloadable' in columns

                            if has_downloadable:
                                cursor = execute_query(f"""
                                    UPDATE {table_name}
                                    SET is_premium_only = %s, use_for_ai_generation = %s, is_downloadable = %s
                                    WHERE id = %s
                                """, (1 if premium_toggle else 0, 1 if ai_toggle else 0, 1 if downloadable_toggle else 0, pdf['id']))
                            else:
                                cursor = execute_query(f"""
                                    UPDATE {table_name}
                                    SET is_premium_only = %s, use_for_ai_generation = %s
                                    WHERE id = %s
                                """, (1 if premium_toggle else 0, 1 if ai_toggle else 0, pdf['id']))
                            cursor.close()

                            if is_snowflake():
                                db_conn.commit()
                            get_pdf_resources.clear()  # Clear cache
                            get_document_library.clear()  # Clear cache
                            st.success(f"✅ Settings updated for {pdf['filename']}")
                            st.rerun()
                        
                        if st.button("🗑️ Delete", key=f"delete_{pdf['id']}"):
                            # Delete file from disk if present
                            if pdf.get('filepath') and os.path.exists(pdf['filepath']):
                                try:
                                    os.remove(pdf['filepath'])
                                except Exception as e:
                                    st.warning(f"⚠️ File could not be removed from disk: {e}")
                            # Delete record from database
                            delete_pdf_resource(pdf['id'])
                            if is_snowflake():
                                db_conn.commit()
                            # Clear caches and confirm
                            get_pdf_resources.clear()
                            get_document_library.clear()
                            st.success("✅ PDF deleted from library.")
                            st.rerun()
        else:
            st.info("No PDFs uploaded yet.")
    
    with tab5:
        st.markdown("### 👥 User Management")
        
        users_data = get_all_users_cached()
        
        # Get approved receipts
        table_name = get_table_name("payment_receipts")
        cursor = execute_query(f"""
            SELECT email, full_name, status, reviewed_at, notes
            FROM {table_name}
            WHERE status = 'Approved'
            ORDER BY reviewed_at DESC
        """)
        approved_receipts = cursor.fetchall()
        cursor.close()
        
        # Add new user section
        st.markdown("#### ➕ Add New User")
        with st.form("add_user_form"):
            col1, col2 = st.columns(2)
            with col1:
                new_user_email = st.text_input("Email Address *", placeholder="user@example.com", key="new_user_email")
                new_user_access = st.selectbox("Initial Access Level", ["Free", "Advance", "Premium"], key="new_user_access")
            with col2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.form_submit_button("➕ Add User", type="primary", use_container_width=True):
                    if new_user_email and "@" in new_user_email and "." in new_user_email.split("@")[1]:
                        # Check if user already exists
                        table_name = get_table_name("users")
                        cursor = execute_query(f"SELECT email FROM {table_name} WHERE email = %s", (new_user_email.lower(),))
                        if cursor.fetchone():
                            st.error(f"❌ User with email {new_user_email} already exists.")
                        else:
                            # Create new user (include last_ip if column exists)
                            user_cols = get_table_columns("users")
                            has_last_ip = user_cols and "last_ip" in [c.lower() for c in user_cols]
                            if has_last_ip:
                                cursor = execute_query(f"""
                                    INSERT INTO {table_name} (email, access_level, questions_answered, created_at, last_login, is_admin, last_ip)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                                """, (new_user_email.lower(), new_user_access, 0, datetime.now().isoformat(), datetime.now().isoformat(), 0, ""))
                            else:
                                cursor = execute_query(f"""
                                    INSERT INTO {table_name} (email, access_level, questions_answered, created_at, last_login, is_admin)
                                    VALUES (%s, %s, %s, %s, %s, %s)
                                """, (new_user_email.lower(), new_user_access, 0, datetime.now().isoformat(), datetime.now().isoformat(), 0))
                            cursor.close()
                            if is_snowflake():
                                db_conn.commit()
                            get_all_users_cached.clear()
                            st.success(f"✅ User {new_user_email} added with {new_user_access} access!")
                            st.rerun()
                    else:
                        st.error("❌ Please enter a valid email address.")
        
        st.markdown("---")
        
        # Display approved receipts
        if approved_receipts:
            st.markdown("#### ✅ Approved Receipts (Users)")
            receipts_df = pd.DataFrame(approved_receipts, columns=["Email", "Name", "Status", "Approved At", "Notes"])
            st.dataframe(receipts_df, use_container_width=True, hide_index=True)
            st.markdown("---")
        
        if users_data:
            st.markdown("#### 📊 All Users")
            # Build DataFrame from dict keys, then use friendly column titles (includes Last IP)
            df = pd.DataFrame(users_data)
            df = df.rename(columns={
                "email": "Email",
                "access_level": "Access Level",
                "questions_answered": "Questions Answered",
                "created_at": "Created At",
                "last_login": "Last Login",
                "last_ip": "Last IP",
            })
            # Ensure Last IP column is visible (reorder so it appears after Last Login)
            if "Last IP" in df.columns:
                cols = [c for c in df.columns if c != "Last IP"]
                idx = cols.index("Last Login") + 1 if "Last Login" in cols else len(cols)
                cols.insert(idx, "Last IP")
                df = df[[c for c in cols if c in df.columns]]
            
            # Search
            search_email = st.text_input("🔍 Search User by Email", placeholder="Enter email to search")
            if search_email:
                df = df[df["Email"].str.contains(search_email.lower(), na=False, case=False)]
            
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            st.markdown("---")
            st.markdown("#### ⚡ Update User Access")
            
            col1, col2 = st.columns(2)
            with col1:
                user_email_update = st.text_input("User Email", placeholder="user@example.com", key="update_user_email")
                new_access_level = st.selectbox("New Access Level", ["Free", "Advance", "Premium"], key="user_access_update")
            with col2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("💾 Update Access Level", type="primary", use_container_width=True):
                    if user_email_update:
                        update_user_access_level(user_email_update.lower(), new_access_level)
                        # Clear relevant caches
                        get_all_users_cached.clear()
                        get_user_info.clear(user_email_update.lower())
                        st.success(f"✅ User {user_email_update} access updated to {new_access_level}")
                        st.rerun()
                    else:
                        st.error("Please enter a user email.")
        else:
            st.info("No users registered yet.")

# ============================================================================
# FOOTER
# ============================================================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #6c757d; padding: 1rem;">
    <p><strong>For review purposes only.</strong> Always verify with official references and latest PH laws and regulations.</p>
    <p>© 2024 PH Criminology Exam Reviewer</p>
</div>
""", unsafe_allow_html=True)
